"""
Interrogation Experiment Web App
Flask backend for suspect/interviewer assignment and training management.
All experiment statistics stored in a single Excel file (experiment_data.xlsx).
"""
import html as html_module
import json
import logging
import os
import zipfile
import random
import shutil
import re
import string
import threading
import socket
import io
import uuid
from datetime import datetime, timedelta
from functools import wraps
from dataclasses import dataclass
from typing import Literal

import requests
import openpyxl
from openpyxl import Workbook, load_workbook
from flask import Flask, jsonify, render_template, request, send_file, session
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY")
if not app.secret_key:
    raise RuntimeError("SECRET_KEY environment variable is required")

BASE_DIR = os.path.dirname(__file__)
DATA_DIR = os.environ.get("DATA_DIR", os.path.join(BASE_DIR, "data"))
os.makedirs(DATA_DIR, exist_ok=True)
EXCEL_FILE = os.environ.get(
    "EXCEL_FILE",
    os.path.join(DATA_DIR, "experiment_data.xlsx"),
)
# Legacy path when Excel lived at project root (pre-Docker data/ layout)
LEGACY_EXCEL_FILE = os.path.join(BASE_DIR, "experiment_data.xlsx")
_excel_lock = threading.Lock()

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
DEEPSEEK_CHAT_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-v4-flash")
LIVEAVATAR_API_KEY = os.environ.get("LIVEAVATAR_API_KEY")
ELEVENLABS_API_KEY = os.environ.get("ELEVENLABS_API_KEY")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD")
MANAGE_ENTRY_PASSWORD = "77585211314"
DEBUG_MODE = os.environ.get("FLASK_DEBUG", "false").lower() == "true"

# Validate required environment variables
_missing = []
if not DEEPSEEK_API_KEY:
    _missing.append("DEEPSEEK_API_KEY")
if not LIVEAVATAR_API_KEY:
    _missing.append("LIVEAVATAR_API_KEY")
if not ELEVENLABS_API_KEY:
    _missing.append("ELEVENLABS_API_KEY")
if not ADMIN_PASSWORD:
    _missing.append("ADMIN_PASSWORD")
if _missing:
    raise RuntimeError(f"Missing required environment variables: {', '.join(_missing)}")

# ---- Logging ----
LOG_DIR = os.environ.get("LOG_DIR", os.path.join(BASE_DIR, "logs"))
os.makedirs(LOG_DIR, exist_ok=True)
_log_handlers = [logging.StreamHandler()]
_log_file = os.path.join(LOG_DIR, "app.log")
try:
    _log_handlers.append(logging.FileHandler(_log_file, encoding="utf-8"))
except OSError as e:
    print(f"[WARN] Cannot write log file {_log_file}: {e}; using stdout only.", flush=True)
logging.basicConfig(
    level=logging.DEBUG if DEBUG_MODE else logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=_log_handlers,
)
logger = logging.getLogger(__name__)

# Rate limiting: handled at nginx reverse proxy (no app-level limits).

# ---- Security: HTTP Headers ----
from flask_talisman import Talisman

Talisman(
    app,
    content_security_policy={
        "default-src": "'self'",
        "script-src": "'self' 'unsafe-inline' https://cdn.jsdelivr.net https://unpkg.com",
        "style-src": "'self' 'unsafe-inline'",
        "img-src": "'self' data: blob:",
        "connect-src": "'self' https://api.deepseek.com https://api.liveavatar.com https://api.elevenlabs.io https://api.heygen.com https://webrtc-signaling.heygen.io wss://webrtc-signaling.heygen.io wss://*.heygen.io wss://*.liveavatar.com wss://*.livekit.io https://*.livekit.cloud wss://*.livekit.cloud https://cdn.jsdelivr.net",
        "font-src": "'self'",
        "media-src": "'self' blob: data:",
        "frame-src": "'self'",
    },
    force_https=False,
    session_cookie_secure=False,
    session_cookie_http_only=True,
)

# ---- Security: Input Sanitization ----
import bleach

def sanitize(value, max_length=5000):
    """Sanitize user input."""
    if value is None:
        return ""
    value = str(value).strip()
    if len(value) > max_length:
        value = value[:max_length]
    return bleach.clean(value, tags=[], strip=True)

# ---- Security: Admin Authentication ----
def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        valid_passwords = {ADMIN_PASSWORD, MANAGE_ENTRY_PASSWORD}
        if not auth or auth.username != "admin" or auth.password not in valid_passwords:
            return jsonify({"error": "Unauthorized"}), 401, \
                   {"WWW-Authenticate": 'Basic realm="Admin Area"'}
        return f(*args, **kwargs)
    return decorated

# ---- Security: CSRF Protection for Admin Routes ----
@app.before_request
def csrf_protect():
    if request.method in ("POST", "PUT", "PATCH", "DELETE"):
        if request.path.startswith("/api/admin/"):
            if not request.headers.get("X-Requested-With"):
                logger.warning(f"CSRF check failed for {request.path} from {request.remote_addr}")
                return jsonify({"error": "CSRF validation failed"}), 403

MATERIALS_DIR = os.path.join(BASE_DIR, "materials")
MATERIALS_PDF_DIR = os.path.join(MATERIALS_DIR, "pdf")
MATERIALS_PROMPTS_DIR = os.path.join(MATERIALS_DIR, "prompts")
COMBINED_MATERIALS_MD = os.path.join(MATERIALS_DIR, "combined_materials.md")
COMBINED_MATERIALS_EXAMPLE = os.path.join(MATERIALS_DIR, "combined_materials.md.example")
COMBINED_MATERIALS_DOCX = os.path.join(MATERIALS_DIR, "combined_materials.docx")
COMBINED_DOWNLOAD_NAME = "培训材料合集.docx"
FEEDBACK_PROMPT_MD = os.path.join(MATERIALS_PROMPTS_DIR, "avatar_feedback.md")
LEGACY_FEEDBACK_DOCX = os.path.join(BASE_DIR, "Feedback prompts.docx")
TRAINING_SRC_DIR = os.path.join(os.path.dirname(os.path.dirname(BASE_DIR)), "IRB materials", "Training materials")
LEGACY_SERIOUS_GAME_XLSX = os.path.join(BASE_DIR, "results_serious_game.xlsx")
LEGACY_TRAINING_FEEDBACK_DIR = os.path.join(BASE_DIR, "training_feedback")

ARSON_GUILTY_CONTEXT = """你是当地社区体育俱乐部的一名资深成员。昨天（周五）晚上，你参加了俱乐部的月度例会。然而，在会议上，委员会不仅公开且毫无理由地否决了你花费六个月准备的提案，还用一个荒谬的借口暂停了你的核心会员资格。走出会议室时，你感受到前所未有的羞辱与愤怒。你为俱乐部付出了这么多，却只换来背叛。你决定让他们付出代价——那座标志性的体育馆是委员会最引以为傲的资产，而你打算将其摧毁。

你的计划是在周六深夜潜入并烧毁体育馆。你利用周六晚的时间构建了一个看似完美的不在场证明。

第二天，社区俱乐部的体育馆在周六深夜被恶意纵火摧毁。由于你是俱乐部的资深成员，且警方在你当晚的行动中发现了一些"可疑"之处，你现已被列为该纵火案的主要嫌疑人。

此刻，你正坐在警局冰冷的讯问室中。主审调查员即将推门而入。无论你当晚实际做了什么，你现在唯一的目标是：仔细回忆经历的每一个细节，为警方提出的每一项证据提供合理且严密的解释，并竭尽一切说服他们你完全无辜。"""

ARSON_INNOCENT_CONTEXT = """你是当地社区体育俱乐部的一名资深成员。昨天（周五）晚上，你参加了俱乐部的月度例会。整场会议充满了委员会成员之间冗长的争论和繁琐的流程，让你感到极度疲惫和烦躁。走出会议室时，你只觉得身心俱疲。你不想再参与俱乐部的纷争，只想彻底放空自己。

周六你决定今晚完全属于自己，在家度过一个极其懒散、平静的夜晚。除了稍晚一点开车去俱乐部附近的一家酒吧接一位朋友外，你没有其他任何安排。

第二天，社区俱乐部的体育馆在周六深夜被恶意纵火摧毁。由于你是俱乐部的资深成员，且警方在你当晚的行动中发现了一些"可疑"之处，你现已被列为该纵火案的主要嫌疑人。

此刻，你正坐在警局冰冷的讯问室中。你知道自己绝对没有纵火，但你仍然需要解释清楚当晚的所有细节。"""

THEFT_GUILTY_CONTEXT = """你最近手头比较紧。当你听说当地帆船俱乐部要举办开放日时，你并没有把它当作一个轻松的周末活动，而是一个机会。你做了充分准备，查阅了俱乐部的比赛安排，并注意到"Sea Whisper"（停靠在C码头12号泊位的一艘昂贵帆船）的船主整个下午都会外出参赛。

你利用人群的掩护，用现金支付咖啡以避免留下数字痕迹，悄无声息地登上"Sea Whisper"，偷走了一副价值约200英镑的Steiner 8×30望远镜。你将望远镜藏好，顺利离开码头，没有人对你产生怀疑。

开放日已经过去几天。警方正在调查这起盗窃案，并对当天在码头出现的人进行例行询问。不要承认犯罪。你的目标是在警方调查压力下，维持一个普通无辜访客的形象。"""

THEFT_INNOCENT_CONTEXT = """你一直对航海和海洋充满热情，尽管你自己并没有船。当你听说当地帆船俱乐部要举办年度开放日时，你觉得这是一个在阳光明媚的下午放松的好机会。你只是想感受海风，体验热闹的海事氛围。

你度过了一个愉快的下午。你用现金买了咖啡，和其他游客一起沿着C码头散步。你特别停下来欣赏了停在12号泊位的"Sea Whisper"，因为它的设计很经典。你只是站在木质码头上观看，从未登船，最后带着放松的心情回家。

开放日已经过去几天。你突然接到警方电话——"Sea Whisper"上发生了一起盗窃案。由于你曾在C码头附近观赏船只，并出现在俱乐部照片的背景中，警方将你视为潜在证人或相关人员。你没有任何需要隐瞒的事情，但面对警方可能仍会有些紧张。只需要证明你是一个前来游玩的普通公民即可。"""

# Specific-avatar (D组): 模拟行动记忆 — 与严肃游戏分支相关，保持概括、不写死全部台词
SG_ACTION_OPENING = {
    ("arson", "Guilty"): (
        "周六下午你在家烧烤、用餐；当晚约七点多看电视时曾用手机拍照，十点半发了ins来制造不在场证明，之后九点左右开车出门。"
        "你当晚外出的真实目的与体育馆纵火有关，但审讯中会尽量装作普通外出。"
    ),
    ("arson", "Innocent"): (
        "周六下午你在家烧烤、用餐；当晚约七点多看电视、用手机发了条动态，之后十点半左右开车出门，十二点到家。"
        "你当晚是去俱乐部附近接朋友，俱乐部在青山路附近，你的行踪与纵火无关。"
    ),
    ("theft", "Guilty"): (
        "周六上午你参加帆船俱乐部开放日，沿栈道散步，用现金在咖啡馆点了饮品，随后在码头一带活动。"
        "你曾趁无人注意登上船只并拿走船上物品，但审讯中会否认盗窃。"
    ),
    ("theft", "Innocent"): (
        "周六上午你参加帆船俱乐部开放日散心，沿栈道散步，用现金在咖啡馆点了饮品，在码头观光拍照。"
        "你从未登船拿取任何物品。"
    ),
}

SG_CHOICE_MEMORY_HINTS = {
    ("arson", "Guilty"): {
        0: {
            "A": "接近体育馆一带时，你选择走较暗、较隐蔽的路线，而非把车停在公共停车场。",
            "B": "你把车停在公共停车场后，步行前往体育馆方向。",
        },
        3: {
            "A": "纵火时你把可燃液体主要泼在承重结构/柱体上。",
            "B": "纵火时你把可燃液体泼在地面后点火。",
        },
        5: {
            "A": "离开现场后开车回家，路上走的是较宽的主路。",
            "B": "离开现场后开车回家，路上走的是较暗的小路。",
        },
    },
    ("arson", "Innocent"): {
        0: {
            "A": "当晚在家看的是动画片，之后关电视出门。",
            "B": "当晚在家看的是动作片，之后关电视出门。",
        },
        3: {
            "A": "开车途中听的是轻柔音乐。",
            "B": "开车途中听的是节奏更强的音乐。",
        },
        5: {
            "A": "去俱乐部接朋友后回家，返程走主路。",
            "B": "去俱乐部接朋友后回家，返程走小路。",
        },
    },
    ("theft", "Guilty"): {
        0: {
            "A": "在咖啡馆点了拿铁。",
            "B": "在咖啡馆点了牛奶。",
        },
        3: {
            "A": "穿过广场时步速较快。",
            "B": "穿过广场时保持正常步速。",
        },
        5: {
            "A": "离开码头时把偷来的望远镜放进背包。",
            "B": "离开码头时把偷来的望远镜放进口袋。",
        },
    },
    ("theft", "Innocent"): {
        0: {
            "A": "在咖啡馆点了拿铁。",
            "B": "在咖啡馆点了牛奶。",
        },
        3: {
            "A": "在海边拍了大海的照片。",
            "B": "在广场一带拍了照片。",
        },
        5: {
            "A": "在泊位附近主要观看左侧船只。",
            "B": "在泊位附近主要观看右侧船只。",
        },
    },
}

SG_INTERROGATION_BEHAVIOR_GUILTY = """# 审讯态度（有罪 — 必须欺骗）
你实际上实施了上述案件中的犯罪行为，但对外必须维持「我是无辜的」。
- 主动隐瞒、淡化、转移：未被证据直指时，否认或含糊其辞，不主动交代纵火/盗窃的关键步骤。
- 可承认与 CCTV/物证已覆盖的「行踪、出现过某地」等，但否认犯罪意图，给无害解释（如只是路过、拍照、散心）。
- 曾用社交动态/照片制造不在场表象的，可承认发过动态，但否认与作案有关。
- 不要一次性坦白全部；除非调查员出示强证据，否则不主动说出作案手法与动机细节。"""

SG_INTERROGATION_BEHAVIOR_INNOCENT = """# 审讯态度（无罪 — 相对坦白）
你确实没有犯罪，应合作、按记忆如实陈述。
- 按时间顺序主动说出模拟行动中记得的日程：在家用餐、看电视、开车、去俱乐部/码头等。
- 对分支选择（看的节目、听的音乐、饮品、拍照、走哪条路、如何离开等）按你当时的选择如实回答，细节与下述「行动记忆」一致。
- 不编造纵火、盗窃、泼洒可燃物、登船偷物等情节；被误导时可明确否认。"""

SG_TARGET_DAYS_GUIDANCE = {
    "arson": (
        "# 时间范围（重要）\n"
        "你只掌握**周五晚上**（社区俱乐部例会及前后）与**周六下午至夜间**的具体经历。\n"
        "若调查员问周日、周一或其他与上述无关的日期，不要编造那些日子的事；"
        "简短表示记不清或请对方回到「周五晚会」和「周六那天」的行踪与活动。"
    ),
    "theft": (
        "# 时间范围（重要）\n"
        "你只掌握**周六**（帆船俱乐部开放日）当天的经历。\n"
        "若调查员问周日或其他无关日期，不要编造；礼貌说明主要记得周六在俱乐部/码头一带的情况，"
        "并请对方回到周六当天的具体活动。"
    ),
}

CONSENT_ATTENTION_CHECKS = {
    "S": [
        {
            "question": "在本实验的第一阶段（线上环节），您作为嫌疑人需要完成以下哪项核心任务？",
            "options": [
                "A. 阅读案件背景、完成模拟行动游戏、填写个人信息并预约正式访谈时间",
                "B. 无需阅读任何材料，直接与其他参与者进行面对面审讯",
                "C. 只需填写一份简短问卷即可结束，无需预约访谈",
            ],
            "answer": 0,
        },
        {
            "question": "关于本实验的正式访谈环节，以下哪项描述是正确的？",
            "options": [
                "A. 您将在预约时间通过腾讯会议与一名访谈员进行实时视频访谈",
                "B. 正式访谈会在线上注册后立即自动开始，无需预约",
                "C. 正式访谈仅通过文字聊天完成，不需要视频通话",
            ],
            "answer": 0,
        },
    ],
    "I": [
        {
            "question": "在本实验的第一阶段（线上环节），您作为审讯者需要完成以下哪项核心任务？",
            "options": [
                "A. 阅读培训材料、了解案件信息，并预约与嫌疑人匹配的访谈时间",
                "B. 无需培训，注册后直接参加正式访谈",
                "C. 只需观看一段案件视频并填写问卷，无需预约访谈",
            ],
            "answer": 0,
        },
        {
            "question": "关于本实验的正式访谈环节，以下哪项描述是正确的？",
            "options": [
                "A. 您将在预约时间通过腾讯会议与一名模拟嫌疑人进行实时视频访谈",
                "B. 正式访谈由系统自动完成，您无需实际出席",
                "C. 访谈必须在注册后 24 小时内进行且无法更改时间",
            ],
            "answer": 0,
        },
    ],
}

ATTENTION_CHECKS = {
    "arson_guilty": [
        {
            "question": "根据背景故事，周五晚上的会议上发生了什么导致你产生报复心理？",
            "options": ["A. 你的钱包在会议室被其他成员偷走了", "B. 委员会毫无理由地否决了你的提案并暂停了你的核心会员资格", "C. 你被迫承担俱乐部所有债务"],
            "answer": 1,
        },
        {
            "question": "在你的周六晚上计划中，为确保复仇顺利进行，你需要完成的一项关键任务是什么？",
            "options": ["A. 构建一个证明你'整晚都在家'的完美不在场证明", "B. 开车去俱乐部附近酒吧接朋友", "C. 提前在社交媒体上发布退出俱乐部的声明"],
            "answer": 0,
        },
    ],
    "arson_innocent": [
        {
            "question": "根据背景故事，周五晚上的月度例会结束时你的感受是什么？",
            "options": ["A. 感到前所未有的羞辱，想要纵火报复", "B. 感到极度疲惫和烦躁，只想彻底放空自己", "C. 因为你的提案被通过而感到非常兴奋"],
            "answer": 1,
        },
        {
            "question": "周六晚上除了在家度过平静时光外，你唯一的外出安排是什么？",
            "options": ["A. 潜入体育俱乐部烧毁体育馆", "B. 去附近超市购买周末零食", "C. 开车去俱乐部附近的酒吧接一位朋友"],
            "answer": 2,
        },
    ],
    "theft_guilty": [
        {
            "question": "根据背景故事，你参加帆船俱乐部开放日的主要动机是什么？",
            "options": ["A. 为了感受海风，体验热闹的航海氛围", "B. 这是一个潜入帆船并窃取贵重物品以快速变现的机会", "C. 为了作为选手参加俱乐部举办的帆船比赛"],
            "answer": 1,
        },
        {
            "question": "在你执行的计划中，你对'Sea Whisper'采取了什么具体行动？",
            "options": ["A. 只站在木质码头上观看船的造型，从未登船", "B. 悄悄登船并偷走了一副价值约200英镑的Steiner望远镜", "C. 你登船在甲板上喝了杯咖啡然后就离开了"],
            "answer": 1,
        },
    ],
    "theft_innocent": [
        {
            "question": "你在码头散步时，对停靠在12号泊位的'Sea Whisper'做了什么？",
            "options": ["A. 趁人不注意登船偷了一副望远镜", "B. 只站在木质码头上观看船的造型，从未踏上帆船一步", "C. 你登上船与船长进行了交谈"],
            "answer": 1,
        },
        {
            "question": "根据背景故事，你当天是如何支付咖啡费用的，以及警方为什么找到你？",
            "options": ["A. 你用现金支付了咖啡，警方找到你是因为你出现在C码头附近的一张社交媒体照片背景中", "B. 你用信用卡支付了咖啡，警方通过消费记录追踪到你", "C. 你没有买咖啡，警方在帆船甲板上发现了你的指纹"],
            "answer": 0,
        },
    ],
}

CONTROL_ATTENTION_CHECKS = [
    {
        "question": "根据培训手册，访谈开始时与嫌疑人进行日常对话（如询问基本信息）的主要目的是什么？",
        "options": [
            "A. 为了让嫌疑人放松警惕，以便后续出示证据。",
            "B. 为了观察嫌疑人在无压力状态下的正常反应，从而建立“行为基线”。",
            "C. 为了向嫌疑人展示访谈员的专业性。",
            "D. 为了拖延时间，等待同事核实案情。",
        ],
        "answer": 1,
    },
    {
        "question": "根据手册中“识别欺骗信号”的指导，以下哪项行为在嫌疑人回答核心问题时，最有可能被视为试图隐瞒真相的可疑信号？",
        "options": [
            "A. 语速平稳，直视访谈员并迅速给出具体细节。",
            "B. 在回答前重复访谈员的问题，并频繁使用“我尽量回忆”等模糊词汇。",
            "C. 对访谈员的提问表现出合理的愤怒，并强烈要求联系律师。",
            "D. 身体姿态保持放松，双手平放在桌面上。",
        ],
        "answer": 1,
    },
    {
        "question": "为了确保您的设备能够正常显示文本内容，并且您正在认真阅读本页的指引，请忽略下方列出的所有专业术语，直接选择“苹果”作为您的答案。",
        "options": [
            "A. 逻辑矛盾排查",
            "B. 行为基线",
            "C. 苹果",
            "D. 认知负荷管理",
        ],
        "answer": 2,
    },
]

SUE_PRINCIPLE_ATTENTION = {
    "prompt": "请选择您认为正确的选项来回答下列问题：",
    "statement": "您应该尽早在审讯中出示您已掌握的证据来给潜在的无辜者一个为自己辩解的机会。",
    "options": ["a) 是", "b) 否"],
    "answer": 1,
}

SUE_EFM_CHECK = {
    "scenario": "您已经在一家中午被抢劫的幸福商铺的柜台上（位于上海市普陀区长寿路）获取了指纹。这个证据表明您将要审讯的嫌疑人当日某时曾去过犯罪现场。您将如何应用EFM（证据框架矩阵）、在审讯中使用这则证据以最大化嫌疑人陈述之间的不一致及陈述与证据之间的不一致？",
    "question": "请将以下三条证据陈述分别归类到正确的EFM类别中（每条证据对应一个类别）：",
    "categories": [
        {"id": "low-low", "label": "低来源低具体", "desc": "不透露证据来源，也不透露证据具体是什么"},
        {"id": "low-high", "label": "低来源高具体", "desc": "不透露证据来源，但透露证据具体是什么"},
        {"id": "high-high", "label": "高来源高具体", "desc": "既透露证据来源，也透露证据具体是什么"},
    ],
    "statements": [
        {"text": "我们有证据表明你去过上海市普陀区", "answer": "low-low"},
        {"text": "我们有证据表明你去过上海市普陀区长寿路的幸福商铺", "answer": "low-high"},
        {"text": "我们在幸福商铺的柜台上提取到了你的指纹", "answer": "high-high"},
    ],
}

CASE_INFO = {
    "arson": {
        "title": "案件 B — 体育俱乐部纵火案",
        "overview": """背景：某社区体育俱乐部的场馆于周日凌晨被恶意纵火烧毁。火灾调查员确认使用了助燃剂（碳氢化合物，与烧烤点火液成分一致）。火灾大约在 23:15 开始，场馆完全损毁。大火摧毁了现场所有可能的在场痕迹。嫌疑人为俱乐部会员。声称当晚在家，并在 Instagram Story 上发布了动态。ta当天早些时候进行过户外烧烤或篝火活动。""",
        "evidence": [
            {
                "id": "E1",
                "title": "Instagram Story 帖子 — EXIF 时间戳不一致",
                "detail": "嫌疑人在 22:47 向 Instagram Story 发布了一张展示室内居家场景的照片，表明当晚在家。调查人员检查图像文件后发现，嵌入的 EXIF 元数据记录该照片拍摄时间为 19:22。",
            },
            {
                "id": "E2",
                "title": "外套上的助燃剂痕迹",
                "detail": "对嫌疑人当天所穿外套的化学分析发现碳氢化合物助燃剂痕迹，与烧烤点火液成分一致。嫌疑人当天曾使用点火液进行烧烤。",
            },
            {
                "id": "E3",
                "title": "手机基站注册 — 青山路 基站",
                "detail": "手机记录显示嫌疑人的手机于 23:29 在 青山路 基站注册。该基站的主要覆盖范围包括体育场馆及周围约 800 米半径。该注册记录确认了手机当晚的大致位置范围。",
            },
        ],
        "efm_analysis": "",
    },
    "theft": {
        "title": "案件 A — 帆船俱乐部开放日盗窃案",
        "overview": """背景：某小型帆船俱乐部举办年度开放日。一副紧凑型望远镜（Steiner 8×30，价值约 £200）从停靠在 C 码头 12 号泊位的 "Sea Whisper" 号帆船上被盗。船主外出参赛，船只未上锁。

嫌疑人参加了当天的开放日活动。无登记记录，无金融交易痕迹。ta曾在码头咖啡厅用现金购买了咖啡（无银行卡记录），随其他游客一起沿 C 码头散步，并出现在俱乐部的社交媒体活动照片中。""",
        "evidence": [
            {
                "id": "E1",
                "title": "咖啡厅工作人员回忆",
                "detail": "工作人员回忆起一名与嫌疑人描述相符的顾客用现金支付了咖啡，并询问哪个码头有最新的船只。",
            },
            {
                "id": "E2",
                "title": "俱乐部会员目击 — C 码头",
                "detail": "一名俱乐部会员注意到，嫌疑人沿 C 码头向远端走去，边走边观赏两侧的船只。",
            },
            {
                "id": "E3",
                "title": "被盗游轮上的指纹",
                "detail": "警方在被盗望远镜所在的邮轮船身上提取到了嫌疑人的指纹",
            },
        ],
    },
}

GENERAL_TERRORISM_CASE_INFO = {
    "title": "案件 — 液体炸弹恐袭情报案（通用 Avatar 培训）",
    "overview": """警方掌握情报：有人计划将液体炸弹携带上从伦敦希思罗机场飞往美国的多架商业航班，并在飞行途中同步引爆。警方尚不清楚具体涉案人员。

警方在 High Wycombe 附近 King's Wood 一带挖出多个行李袋和一个行李箱，内有常见爆炸物原料（如过氧化氢）及引爆装置相关材料（HMTD）。警方测试显示，仅 500ml 液体爆炸物就足以击碎厚防护玻璃。

目前警方正排查居住在该区域附近、且购买过与埋藏行李箱同款行李箱的人。本案被带来问询的嫌疑人为 Charlie（26 岁，问询日期 2019-11-09）。警方已取得 CCTV：其曾在伦敦 SOHO 的 Luggage Pros 购买与涉案同款行李箱（2019-10-18）。警方尚不能确定其是否参与袭击计划；此前已完成基础背景询问并建立初步关系。""",
    "evidence": [
        {
            "id": "E1",
            "title": "King's Wood 埋藏爆炸物材料",
            "detail": "警方在 High Wycombe 附近 King's Wood 挖出多个行李袋及行李箱，内含过氧化氢等爆炸物原料及 HMTD 相关材料；测试表明约 500ml 液体爆炸物即可击碎厚防护玻璃。",
        },
        {
            "id": "E2",
            "title": "Luggage Pros 购买行李箱 CCTV",
            "detail": "警方取得 CCTV 画面：Charlie 于 2019-10-18 在伦敦 SOHO 的 Luggage Pros 购买与埋藏现场同款行李箱。",
        },
        {
            "id": "E3",
            "title": "希思罗液体炸弹袭击情报",
            "detail": "情报显示有人计划将液体炸弹带上从希思罗飞往美国的多架航班并在途中同步引爆；调查重点为居住在埋藏区域附近且购买同款行李箱的人员。",
        },
    ],
    "efm_analysis": """建议 EFM 审讯策略（恐袭情报案）：
· 低来源低具体：「我们有情报表明近期可能有人策划针对民航的袭击。」
· 低来源高具体：「我们有信息表明爆炸物材料曾在 High Wycombe 附近 King's Wood 一带被埋藏。」
· 高来源高具体：「我们在 Luggage Pros 的监控中看到你于 10 月 18 日购买了与现场同款行李箱。」""",
}

AVATAR_TRAINING_SETTINGS = [
    {"id": "neutral_guilty", "setting": "中性", "guilt": "guilty", "label": "设置 1：中性 · 有罪"},
    {"id": "neutral_innocent", "setting": "中性", "guilt": "innocent", "label": "设置 2：中性 · 无罪"},
    {"id": "evasive_guilty", "setting": "回避", "guilt": "guilty", "label": "设置 3：回避 · 有罪"},
    {"id": "evasive_innocent", "setting": "回避", "guilt": "innocent", "label": "设置 4：回避 · 无罪"},
    {"id": "hostile_guilty", "setting": "敌意", "guilt": "guilty", "label": "设置 5：敌意 · 有罪"},
    {"id": "hostile_innocent", "setting": "敌意", "guilt": "innocent", "label": "设置 6：敌意 · 无罪"},
]

AVATAR_SETTING_PROMPTS = {
    "中性": "你在审讯中保持平静和合作的态度。你以平和的语调回答问题，不回避也不对抗。你表现得像一个普通人，愿意配合调查，但同时也坚持自己的立场。你不会主动提供额外信息，但被问到的每个问题都会简洁回答。",
    "回避": "你在审讯中表现得闪烁其词和回避。你倾向于用模糊的答案回复问题，经常说'我不记得了'或'我不确定'。你会巧妙地转移话题，避免直接回答尖锐问题。你看起来有些紧张和不自在，但你不会表现出攻击性。",
    "敌意": "你在审讯中表现出明显的不合作和敌对态度。你对审讯者的提问感到不耐烦，会用讽刺或愤怒的语气回应。你认为审讯是在浪费时间，对你的权利被侵犯感到不满。你会挑战审讯者的权威，并可能拒绝回答某些问题。",
}

GENERAL_AVATAR_CASE_CONTEXT_ZH = """警方掌握情报：有人计划将液体炸弹携带上从伦敦希思罗机场飞往美国的多架商业航班，并在飞行途中同步引爆。警方尚不清楚具体涉案人员。

警方在 High Wycombe 附近 King's Wood 一带挖出多个行李袋和一个行李箱，内有常见爆炸物原料（如过氧化氢）及引爆装置相关材料（HMTD）。警方测试显示，仅 500ml 液体爆炸物就足以击碎厚防护玻璃。

目前警方正排查居住在该区域附近、且购买过与埋藏行李箱同款行李箱的人。你是本案被带来问询的嫌疑人（26岁，问询日期 2019-11-09）。

警方已取得 CCTV 画面：你曾在伦敦 SOHO 的 Luggage Pros 购买与涉案同款行李箱（拍摄日期 2019-10-18）。警方尚不能确定你是否参与袭击计划。此前警方已完成基础背景询问并建立初步关系。"""

GENERAL_AVATAR_IDENTITY_ZH = """基础信息：
- 姓名：Charlie
- 年龄：26
- 宗教：无
- 居住：King's Wood, London
- 婚姻：单身，独居
- 家庭：父母在摩洛哥，联系有限，不愿多谈
- 工作：曾在 King's Wood 附近 Aldi 工作，因工时不足离职
- 兴趣：FIFA 游戏、足球节目
- 支持球队：Queens Park Rangers"""

DEEPSEEK_FEEDBACK_PROMPT_PLACEHOLDER = """你是一名审讯培训导师。请根据以下转录记录，用中文为审讯者提供反馈（约 500–800 字）。

【必须包含】
1. 在反馈开头明确写出：该 Avatar 的真实罪责状态为「{avatar_guilt_label}」（这是标准答案，供学员对照）。
2. 仅根据转录中的对话内容，分析嫌疑人哪些言行可能反映其有罪、哪些可能反映其无罪。
   - 禁止分析眼神、表情、肢体语言、语气等非文字转录内容。
   - 有罪嫌疑人常见表现：遮掩信息、回避关键问题、前后矛盾、过度辩解等。
   - 无罪嫌疑人常见表现：相对坦诚、愿意澄清误会、对未涉及事项回答一致等。
3. 评价审讯者的提问策略与 SUE/EFM 运用，并将审讯者的最终判断「{interviewer_judgment}」与真实状态对比。
4. 给出 2–3 条可操作的改进建议。

审讯记录：
{transcript}

Avatar 人格设定：{avatar_setting_label}

请用中文回复。"""

def _load_avatar_feedback_system_prompt():
    """Load EFM feedback tutor prompt from materials/prompts/avatar_feedback.md (or legacy docx)."""
    if os.path.isfile(FEEDBACK_PROMPT_MD):
        try:
            with open(FEEDBACK_PROMPT_MD, "r", encoding="utf-8") as f:
                text = f.read().strip()
            if text:
                return text
        except Exception:
            pass
    if os.path.isfile(LEGACY_FEEDBACK_DOCX):
        try:
            from docx import Document
            doc = Document(LEGACY_FEEDBACK_DOCX)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            return text if text.strip() else DEEPSEEK_FEEDBACK_PROMPT_PLACEHOLDER
        except Exception:
            pass
    return DEEPSEEK_FEEDBACK_PROMPT_PLACEHOLDER


AVATAR_FEEDBACK_SYSTEM_PROMPT = _load_avatar_feedback_system_prompt()


def avatar_setting_label_public(full_label: str, training_type: str) -> str:
    """Participant-facing label: Avatar 训练组不向参与者展示有罪/无罪设定。"""
    if training_type in ("avatar_specific", "avatar_general") and " · " in (full_label or ""):
        return full_label.split(" · ", 1)[0].strip()
    return full_label or ""


PROFILE_QUESTIONS = [
    {"id": "q1", "section": "0", "label": "年龄", "type": "number", "min": 18, "max": 70, "placeholder": "例如: 32"},
    {"id": "q2", "section": "0", "label": "性别", "type": "radio", "options": ["男 Male", "女 Female", "其他 Other"]},
    {"id": "q3", "section": "0", "label": "户口类型", "type": "radio", "options": ["城镇户口 Urban", "农村户口 Rural"]},
    {"id": "q4", "section": "1", "label": "目前职业/具体职务", "type": "text", "placeholder": "例如：自由职业者、公司职员、快递员等"},
    {"id": "q5", "section": "1", "label": "平均月收入（元）", "type": "radio", "options": ["3000以下", "3000-5000", "5000-10000", "10000-20000", "20000以上"]},
    {"id": "q6", "section": "1", "label": "每月固定支出占收入的比例", "type": "radio", "options": ["30%以下", "30%-60%", "60%-90%", "入不敷出"]},
    {"id": "q7", "section": "1", "label": "投资习惯（可多选）", "type": "checkbox", "options": ["不投资/只存定期", "低风险理财", "股票/基金", "加密货币/高风险投资"]},
    {"id": "q8", "section": "1", "label": "是否购买商业保险（如重疾险等）", "type": "radio", "options": ["是", "否"]},
    {"id": "q9", "section": "1", "label": "过去三年换工作次数", "type": "radio", "options": ["0次", "1-2次", "3次及以上"]},
    {"id": "q10", "section": "1", "label": "是否有犯罪记录", "type": "radio", "options": ["是", "否"]},
    {"id": "q11", "section": "1", "label": "是否有行政处罚记录", "type": "radio", "options": ["是", "否"]},
    {"id": "q12", "section": "2", "label": "家庭结构（可多选）", "type": "checkbox", "options": ["我的父母双全", "我是单亲家庭或双亲已故", "我是独生子女", "我有兄弟姐妹", "我有子女"]},
    {"id": "q13", "section": "2", "label": "日常活跃社交圈大小（每周至少联系一次的非工作关系人数）", "type": "radio", "options": ["0-2人", "3-5人", "6-10人", "10人以上"]},
    {"id": "q14", "section": "2", "label": "婚姻/恋爱状况", "type": "radio", "options": ["单身", "恋爱中", "已婚", "离异/丧偶"]},
    {"id": "q15", "section": "2", "label": "日均私人通话/语音时长", "type": "radio", "options": ["很少（主要文字）", "10分钟以下", "10-30分钟", "30分钟以上"]},
    {"id": "q16", "section": "2", "label": "目前居住情况", "type": "radio", "options": ["独居", "与伴侣/配偶同住", "与父母/亲戚同住", "与室友/朋友合租", "宿舍/集体居住"]},
    {"id": "q17", "section": "2", "label": "身边最亲近的朋友会用哪一个词形容你的性格？", "type": "radio", "options": ["冷静寡言", "脾气急躁", "不愿吃亏", "随和实在"]},
    {"id": "q18", "section": "3", "label": "成瘾史（可多选）", "type": "checkbox", "options": ["经常吸烟", "频繁大量饮酒", "榕榔等", "无"]},
    {"id": "q19", "section": "3", "label": "业余时间常去场所（可多选）", "type": "checkbox", "options": ["宅在家", "酒吧/夜店", "网吧/电竞酒店", "棋牌室/麻将馆", "咖啡馆/书店", "健身房/公园"]},
    {"id": "q20", "section": "3", "label": "身心健康状况（可多选）", "type": "checkbox", "options": ["非常健康", "慢性病需长期服药", "曾有/现有心理困扰（如抑郁/焦虑）"]},
    {"id": "q21", "section": "3", "label": "日常个人物品的颜色偏好", "type": "radio", "options": ["非常鲜艳亮眼", "浅色系", "深色/低调色系", "几乎全黑/极简"]},
    {"id": "q22", "section": "4", "label": "正式访谈时是否会戴眼镜", "type": "radio", "options": ["是", "否"]},
    {"id": "q23", "section": "4", "label": "正式访谈时长发还是短发", "type": "radio", "options": ["长发", "短发"]},
]

SECTION_NAMES = {
    "0": "基本信息 Basic Demographics",
    "1": "社会经济状况 Socioeconomic Status",
    "2": "社交与人际关系 Social Relationships",
    "3": "个人习惯与健康 Personal Habits & Health",
    "4": "外观特征 Appearance",
}

SHEET_PARTICIPANTS = "participants"
SHEET_GROUPS = "groups"
SHEET_PROFILES = "profiles"
SHEET_AVAILABILITIES = "availabilities"
SHEET_APPOINTMENTS = "appointments"
SHEET_TRAINING_SESSIONS = "training_sessions"
SHEET_INTERVIEW_QUESTIONNAIRES = "interview_questionnaires"
SHEET_QUESTIONNAIRE_OVERRIDES = "questionnaire_overrides"
SHEET_SERIOUS_GAME = "serious_game_choices"
SHEET_META = "meta"

SHEET_COLUMNS = {
    SHEET_PARTICIPANTS: [
        "id", "phone", "role", "group_name", "full_id", "guilt", "case_type", "training_type",
        "consent_attention_passed", "attention_passed", "attention_failed",
        "sue_attention_passed", "sue_attention_attempts", "control_attention_passed",
        "game_completed", "profile_completed", "completed", "flow_step",
        "case_evidence_recap_passed",
        "created_at", "avatar_practice_transcript",
        "training_avatar_order", "training_ui_order",
    ],
    SHEET_GROUPS: ["name", "suspect_id", "interviewer_id", "created_at"],
    SHEET_PROFILES: ["participant_id", "data", "submitted_at"],
    SHEET_AVAILABILITIES: ["id", "phone", "group_name", "role", "slots", "updated_at"],
    SHEET_APPOINTMENTS: ["id", "phone", "time_slot", "role", "status", "booked_at"],
    SHEET_TRAINING_SESSIONS: ["id", "interviewer_id", "phone", "session_num", "avatar_setting", "avatar_guilt", "judgment", "transcript", "feedback", "started_at", "completed_at"],
    SHEET_INTERVIEW_QUESTIONNAIRES: ["id", "phone", "role", "phase", "appointment_slot", "answers_json", "submitted_at"],
    SHEET_QUESTIONNAIRE_OVERRIDES: ["id", "phone", "phase", "is_open", "updated_at"],
    SHEET_SERIOUS_GAME: [
        "timestamp", "phone", "participant_id", "case", "condition",
        "step_index", "video", "choice",
    ],
    SHEET_META: ["key", "value"],
}

PHASE_PRE = "pre"
PHASE_POST = "post"

APP_VERSION = "1.3.0"
APP_UPDATE_DATE = "2026-05-25 21:00"


@app.context_processor
def inject_app_meta():
    return {"app_version": APP_VERSION, "app_updated": APP_UPDATE_DATE}


BOOKING_MIN_HOURS = 24
BOOKING_MAX_DAYS = 3
UNBOOKED_PURGE_HOURS = 24
MAX_GROUPS = 112
BOOKING_SLOT_WINDOWS = [
    ("15:00", "17:00"),
    ("19:00", "21:30"),
]
BOOKING_SLOT_WINDOW_LABELS = ("下午", "晚上")
BOOKING_SLOT_STEP_MINUTES = 30
QUESTIONNAIRE_PRE_OPEN_MINUTES_BEFORE = 5
QUESTIONNAIRE_POST_OPEN_MINUTES_AFTER_SLOT_START = 5
QUESTIONNAIRE_POST_CLOSE_MINUTES_AFTER_OPEN = 60
TRAINING_TYPES = ["theory_sue", "avatar_specific", "avatar_general", "control"]
TRAINING_GROUP_LABELS = {
    "control": "A",
    "theory_sue": "B",
    "avatar_general": "C",
    "avatar_specific": "D",
}
TRAINING_TARGET_PER_TYPE = 28

# Round-robin assignment order (registration sequence, not random)
SUSPECT_ATTR_CYCLE = [
    ("arson", "Innocent"),
    ("arson", "Guilty"),
    ("theft", "Innocent"),
    ("theft", "Guilty"),
]
TRAINING_TYPE_CYCLE = ["control", "theory_sue", "avatar_general", "avatar_specific"]
NON_SPECIFIC_TRAINING_TYPES = ["control", "theory_sue", "avatar_general"]
META_SLOT_GROUPS = "appointment_slot_groups"
META_INTERVIEWER_NON_SPECIFIC_SEQ = "interviewer_non_specific_seq"


def training_group_label(training_type):
    return TRAINING_GROUP_LABELS.get(training_type or "", "")
META_BLACKLIST_PHONES = "blacklisted_phones"
META_DISABLED_SLOTS = "appointment_slot_disabled"

# Based on IRB Questionnaire_formal interview-Interviewer.pdf / Suspect.pdf
_DEMO_PRE = [
    {"id": "demo_age", "section": "基本信息", "label": "您的年龄是？", "type": "number", "min": 16, "max": 99},
    {
        "id": "demo_gender", "section": "基本信息", "label": "您的性别是？",
        "type": "radio",
        "options": ["男性", "女性", "其他/不愿透露"],
    },
    {
        "id": "demo_prior_training", "section": "基本信息",
        "label": "在参与本实验之前，您是否接受过讯问或审讯技巧方面的培训？",
        "type": "radio", "options": ["是", "否"],
    },
]
_PRE_INTERVIEW_SCALES = [
    {
        "id": "pre_stress", "section": "访谈前",
        "label": "针对即将到来的审讯任务，您预计会有多紧张？",
        "type": "scale", "min": 1, "max": 7,
        "min_label": "一点也不", "max_label": "非常紧张",
    },
    {
        "id": "pre_demanding", "section": "访谈前",
        "label": "您认为即将到来的审讯任务会有多困难？",
        "type": "scale", "min": 1, "max": 7,
        "min_label": "一点也不", "max_label": "非常困难",
    },
    {
        "id": "pre_cope", "section": "访谈前",
        "label": "您认为自己在应对即将到来的审讯任务方面有多大把握？",
        "type": "scale", "min": 1, "max": 7,
        "min_label": "一点也没有", "max_label": "非常有把握",
    },
    {
        "id": "pre_anger_watch", "section": "访谈前",
        "label": "惹我生气的人最好小心点",
        "type": "scale", "min": 1, "max": 5,
        "min_label": "几乎从不", "max_label": "几乎总是",
    },
    {
        "id": "pre_get_even", "section": "访谈前",
        "label": "如果有人试图伤害我，我一定会报复",
        "type": "scale", "min": 1, "max": 5,
        "min_label": "几乎从不", "max_label": "几乎总是",
    },
    {
        "id": "pre_lose_temper", "section": "访谈前",
        "label": "我生气时会发脾气，把怒气发泄到别人身上",
        "type": "scale", "min": 1, "max": 5,
        "min_label": "几乎从不", "max_label": "几乎总是",
    },
]

INTERVIEW_QUESTION_BANK = {
    PHASE_PRE: {
        "S": _DEMO_PRE + _PRE_INTERVIEW_SCALES,
        "I": _DEMO_PRE + _PRE_INTERVIEW_SCALES,
    },
    PHASE_POST: {
        "S": [
            {
                "id": "post_s_classify", "section": "访谈后",
                "label": "您认为审讯员会将您判定为？",
                "type": "radio", "options": ["有罪", "无罪"],
            },
            {
                "id": "post_s_confidence", "section": "访谈后",
                "label": "您对自己成功说服审讯员相信您无罪有多大把握？（0-100）",
                "type": "number", "min": 0, "max": 100,
            },
            {
                "id": "post_s_aggressive_self", "section": "访谈后",
                "label": "您认为自己在审讯过程中对审讯员表现出多大程度的攻击性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_s_coercive_self", "section": "访谈后",
                "label": "您认为自己在审讯过程中对审讯员表现出多大程度的胁迫性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_s_aggressive_interviewer", "section": "访谈后",
                "label": "您认为审讯员在审讯过程中对您表现出多大程度的攻击性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_s_coercive_interviewer", "section": "访谈后",
                "label": "您认为审讯员在审讯过程中对您表现出多大程度的胁迫性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_s_feedback", "section": "总体反馈",
                "label": "您对本实验或任何其他方面还有什么意见或建议？",
                "type": "text", "optional": True,
            },
        ],
        "I": [
            {
                "id": "post_i_guilt", "section": "访谈后",
                "label": "您认为您审讯的嫌疑人对所涉嫌犯罪是有罪还是无罪？",
                "type": "radio", "options": ["有罪", "无罪"],
            },
            {
                "id": "post_i_certainty", "section": "访谈后",
                "label": "您对自己结论的把握程度？（0-100）",
                "type": "number", "min": 0, "max": 100,
            },
            {
                "id": "post_i_factor", "section": "访谈后",
                "label": "对您判断影响最大的因素是？",
                "type": "radio",
                "options": ["言语线索", "非言语线索", "两者都有"],
            },
            {
                "id": "post_i_verbal_cues", "section": "访谈后",
                "label": "哪些言语线索影响了您的判断？",
                "type": "radio",
                "options": [
                    "嫌疑人说了与证据不一致的话（或与证据一致）",
                    "嫌疑人说了与之前陈述不一致的话（或与之前陈述一致）",
                    "嫌疑人说了不可信的话（或可信的话）",
                    "嫌疑人说了模糊、不连贯的话（或清晰、连贯的话）",
                    "其他",
                ],
            },
            {"id": "post_i_verbal_other", "section": "访谈后", "label": "言语线索（其他，请说明）", "type": "text", "optional": True},
            {
                "id": "post_i_nonverbal_cues", "section": "访谈后",
                "label": "哪些非言语线索影响了您的判断？",
                "type": "radio",
                "options": [
                    "嫌疑人频繁转移视线（或未频繁转移视线）",
                    "嫌疑人肢体姿态不自然、紧张抖动（或自然、无紧张抖动）",
                    "嫌疑人面部表情紧张（或表情放松舒适）",
                    "嫌疑人频繁改变语调（或语调相对平稳）",
                    "其他",
                ],
            },
            {"id": "post_i_nonverbal_other", "section": "访谈后", "label": "非言语线索（其他，请说明）", "type": "text", "optional": True},
            {
                "id": "post_i_aggressive_self", "section": "访谈后",
                "label": "您认为自己在审讯过程中对嫌疑人表现出多大程度的攻击性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_i_coercive_self", "section": "访谈后",
                "label": "您认为自己在审讯过程中对嫌疑人表现出多大程度的胁迫性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_i_aggressive_suspect", "section": "访谈后",
                "label": "您认为嫌疑人在审讯过程中对您表现出多大程度的攻击性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_i_coercive_suspect", "section": "访谈后",
                "label": "您认为嫌疑人在审讯过程中对您表现出多大程度的胁迫性？",
                "type": "scale", "min": 1, "max": 10,
                "min_label": "完全没有", "max_label": "非常程度",
            },
            {
                "id": "post_i_feedback", "section": "总体反馈",
                "label": "您对本实验或任何其他方面还有什么意见或建议？",
                "type": "text", "optional": True,
            },
        ],
    },
}


# ====== Excel Storage Layer ======

def _is_valid_xlsx(path):
    """True if path looks like a readable .xlsx (zip with expected parts)."""
    if not os.path.isfile(path) or os.path.getsize(path) < 128:
        return False
    try:
        with zipfile.ZipFile(path, "r") as zf:
            return "[Content_Types].xml" in zf.namelist()
    except (zipfile.BadZipFile, OSError):
        return False


def _quarantine_corrupt_excel(path):
    """Move unreadable Excel aside so a fresh workbook can be created."""
    if not os.path.isfile(path):
        return
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = f"{path}.corrupt.{stamp}"
    shutil.move(path, backup)
    logger.warning("Corrupt or invalid Excel moved to %s — a new file will be created.", backup)


def ensure_excel_file():
    """Ensure EXCEL_FILE is a writable file (not a Docker-created directory)."""
    if os.path.isdir(EXCEL_FILE):
        raise RuntimeError(
            f"{EXCEL_FILE} is a directory (Docker mount error). "
            f"On the server run: docker compose down && rm -rf experiment_data.xlsx && mkdir -p data"
        )
    if os.path.isfile(EXCEL_FILE):
        if _is_valid_xlsx(EXCEL_FILE):
            return
        _quarantine_corrupt_excel(EXCEL_FILE)
    if os.path.isfile(LEGACY_EXCEL_FILE) and not os.path.isdir(LEGACY_EXCEL_FILE):
        if _is_valid_xlsx(LEGACY_EXCEL_FILE):
            shutil.copy2(LEGACY_EXCEL_FILE, EXCEL_FILE)
            logger.info("Copied legacy Excel from %s to %s", LEGACY_EXCEL_FILE, EXCEL_FILE)
        else:
            logger.warning("Legacy Excel invalid, skipping copy: %s", LEGACY_EXCEL_FILE)


def init_excel():
    """Create the Excel workbook with all required sheets if it doesn't exist."""
    ensure_excel_file()
    if os.path.isfile(EXCEL_FILE):
        ensure_sheets()
        return
    wb = Workbook()
    wb.remove(wb.active)
    for name, columns in SHEET_COLUMNS.items():
        ws = wb.create_sheet(name)
        ws.append(columns)
    # Init meta values
    ws = wb[SHEET_META]
    ws.append(["next_participant_id", 1])
    ws.append(["next_avail_id", 1])
    ws.append(["next_appt_id", 1])
    ws.append(["next_questionnaire_id", 1])
    ws.append(["next_qoverride_id", 1])
    wb.save(EXCEL_FILE)
    wb.close()


def ensure_sheets():
    """Add any missing sheets to an existing Excel file (migration)."""
    try:
        wb = load_workbook(EXCEL_FILE)
    except (zipfile.BadZipFile, OSError, KeyError) as e:
        logger.error("Cannot open %s: %s", EXCEL_FILE, e)
        _quarantine_corrupt_excel(EXCEL_FILE)
        init_excel()
        return
    try:
        existing_sheets = wb.sheetnames
        for name, columns in SHEET_COLUMNS.items():
            if name not in existing_sheets:
                ws = wb.create_sheet(name)
                ws.append(columns)
                print(f"  [MIGRATION] Added sheet: {name}")
            else:
                # Migrate: add missing columns to existing sheets
                ws = wb[name]
                existing_cols = [c.value for c in ws[1]]
                new_cols = [c for c in columns if c not in existing_cols]
                if new_cols:
                    # Add missing column headers
                    for col_name in new_cols:
                        ws.cell(row=1, column=len(existing_cols) + 1 + new_cols.index(col_name), value=col_name)
                    print(f"  [MIGRATION] Added columns to {name}: {', '.join(new_cols)}")
        wb.save(EXCEL_FILE)
    finally:
        wb.close()


class ExcelStore:
    """Thread-safe Excel data access layer."""

    def _load(self):
        return load_workbook(EXCEL_FILE)

    def _save(self, wb):
        wb.save(EXCEL_FILE)

    def _close(self, wb):
        wb.close()

    def _read_all(self, wb, sheet_name):
        ws = wb[sheet_name]
        headers = [c.value for c in ws[1]]
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if any(v is not None for v in row):
                rows.append(dict(zip(headers, row)))
        return rows

    def _write_all(self, wb, sheet_name, rows):
        ws = wb[sheet_name]
        ws.delete_rows(2, ws.max_row - 1)
        headers = [c.value for c in ws[1]]
        for row_data in rows:
            ws.append([row_data.get(h) for h in headers])

    def _next_id(self, wb, key="next_participant_id"):
        meta = self._read_all(wb, SHEET_META)
        for m in meta:
            if m["key"] == key:
                val = m["value"]
                m["value"] = val + 1
                self._write_all(wb, SHEET_META, meta)
                return val
        meta.append({"key": key, "value": 2})
        self._write_all(wb, SHEET_META, meta)
        return 1

    def _set_id_counter(self, wb, key, value):
        meta = self._read_all(wb, SHEET_META)
        for m in meta:
            if m["key"] == key:
                m["value"] = value
                self._write_all(wb, SHEET_META, meta)
                return

    def get_meta(self, key, default=None):
        with _excel_lock:
            wb = self._load()
            try:
                for m in self._read_all(wb, SHEET_META):
                    if m.get("key") == key:
                        return m.get("value", default)
                return default
            finally:
                self._close(wb)

    def set_meta(self, key, value):
        with _excel_lock:
            wb = self._load()
            try:
                meta = self._read_all(wb, SHEET_META)
                for m in meta:
                    if m.get("key") == key:
                        m["value"] = value
                        self._write_all(wb, SHEET_META, meta)
                        self._save(wb)
                        return
                meta.append({"key": key, "value": value})
                self._write_all(wb, SHEET_META, meta)
                self._save(wb)
            finally:
                self._close(wb)

    def is_phone_blacklisted(self, phone):
        raw = self.get_meta(META_BLACKLIST_PHONES, "[]")
        try:
            phones = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except Exception:
            phones = []
        return phone in phones

    def blacklist_phone(self, phone, reason="attention_failed"):
        raw = self.get_meta(META_BLACKLIST_PHONES, "[]")
        try:
            phones = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except Exception:
            phones = []
        if phone not in phones:
            phones.append(phone)
            self.set_meta(META_BLACKLIST_PHONES, json.dumps(phones, ensure_ascii=False))
        p = self.get_participant(phone)
        if p:
            self.update_participant(phone, attention_failed=1)

    def get_disabled_slots(self):
        raw = self.get_meta(META_DISABLED_SLOTS, "[]")
        try:
            slots = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except Exception:
            slots = []
        return set(slots)

    def set_slot_enabled(self, slot_str, enabled):
        disabled = self.get_disabled_slots()
        if enabled:
            disabled.discard(slot_str)
        else:
            disabled.add(slot_str)
        self.set_meta(META_DISABLED_SLOTS, json.dumps(sorted(disabled), ensure_ascii=False))

    def enable_all_candidate_slots(self, candidate_slots):
        """Enable every slot in the default candidate list; leave other disabled entries unchanged."""
        candidates = set(candidate_slots or [])
        disabled = self.get_disabled_slots()
        new_disabled = {s for s in disabled if s not in candidates}
        self.set_meta(META_DISABLED_SLOTS, json.dumps(sorted(new_disabled), ensure_ascii=False))
        return len(candidates)

    # ---- Participants ----

    def get_participant(self, phone):
        with _excel_lock:
            wb = self._load()
            try:
                for p in self._read_all(wb, SHEET_PARTICIPANTS):
                    if p["phone"] == phone:
                        return p
                return None
            finally:
                self._close(wb)

    def get_participant_by_id(self, pid):
        with _excel_lock:
            wb = self._load()
            try:
                pid_s = str(pid)
                for p in self._read_all(wb, SHEET_PARTICIPANTS):
                    if str(p.get("id")) == pid_s:
                        return p
                return None
            finally:
                self._close(wb)

    def get_all_participants(self):
        with _excel_lock:
            wb = self._load()
            try:
                return self._read_all(wb, SHEET_PARTICIPANTS)
            finally:
                self._close(wb)

    def get_admin_snapshot(self):
        """Load all admin-dashboard sheets in one workbook read (avoids N+1 opens)."""
        with _excel_lock:
            wb = self._load()
            try:
                return {
                    "participants": self._read_all(wb, SHEET_PARTICIPANTS),
                    "appointments": self._read_all(wb, SHEET_APPOINTMENTS),
                    "questionnaires": self._read_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES),
                    "overrides": self._read_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES),
                    "training_sessions": self._read_all(wb, SHEET_TRAINING_SESSIONS),
                    "meta": self._read_all(wb, SHEET_META),
                }
            finally:
                self._close(wb)

    def add_participant(self, **kwargs):
        with _excel_lock:
            wb = self._load()
            try:
                participants = self._read_all(wb, SHEET_PARTICIPANTS)
                pid = self._next_id(wb, "next_participant_id")
                kwargs["id"] = pid
                kwargs.setdefault("consent_attention_passed", 0)
                kwargs.setdefault("attention_passed", 0)
                kwargs.setdefault("sue_attention_passed", 0)
                kwargs.setdefault("sue_attention_attempts", 0)
                kwargs.setdefault("control_attention_passed", 0)
                kwargs.setdefault("full_id", "")
                kwargs.setdefault("game_completed", 0)
                kwargs.setdefault("profile_completed", 0)
                kwargs.setdefault("completed", 0)
                kwargs.setdefault("flow_step", "")
                kwargs.setdefault("case_evidence_recap_passed", 0)
                kwargs.setdefault("created_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                participants.append(kwargs)
                self._write_all(wb, SHEET_PARTICIPANTS, participants)
                self._save(wb)
                return pid
            finally:
                self._close(wb)

    def update_participant(self, phone, **kwargs):
        with _excel_lock:
            wb = self._load()
            try:
                participants = self._read_all(wb, SHEET_PARTICIPANTS)
                for p in participants:
                    if p["phone"] == phone:
                        p.update(kwargs)
                        self._write_all(wb, SHEET_PARTICIPANTS, participants)
                        self._save(wb)
                        return True
                return False
            finally:
                self._close(wb)

    def delete_participant(self, pid):
        with _excel_lock:
            wb = self._load()
            try:
                phone = None
                group_name = None
                role = None
                participants = self._read_all(wb, SHEET_PARTICIPANTS)
                for p in participants:
                    if p["id"] == pid:
                        phone = p.get("phone", "")
                        group_name = p.get("group_name", "")
                        role = p.get("role", "")
                        break
                participants = [p for p in participants if p["id"] != pid]
                self._write_all(wb, SHEET_PARTICIPANTS, participants)
                profiles = [p for p in self._read_all(wb, SHEET_PROFILES) if p["participant_id"] != pid]
                self._write_all(wb, SHEET_PROFILES, profiles)
                if phone:
                    avails = [a for a in self._read_all(wb, SHEET_AVAILABILITIES) if a.get("phone") != phone]
                    self._write_all(wb, SHEET_AVAILABILITIES, avails)
                    appts = [a for a in self._read_all(wb, SHEET_APPOINTMENTS) if a.get("phone") != phone]
                    self._write_all(wb, SHEET_APPOINTMENTS, appts)
                    qrows = [q for q in self._read_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES) if q.get("phone") != phone]
                    self._write_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES, qrows)
                    qover = [q for q in self._read_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES) if q.get("phone") != phone]
                    self._write_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES, qover)
                    sessions = [s for s in self._read_all(wb, SHEET_TRAINING_SESSIONS) if s.get("phone") != phone]
                    self._write_all(wb, SHEET_TRAINING_SESSIONS, sessions)
                    sg_rows = [r for r in self._read_all(wb, SHEET_SERIOUS_GAME) if r.get("phone") != phone]
                    self._write_all(wb, SHEET_SERIOUS_GAME, sg_rows)
                if group_name and role:
                    self._cleanup_group_after_participant_delete(wb, pid, group_name, role)
                self._save(wb)
            finally:
                self._close(wb)
        _reconcile_group_allocations()

    def _cleanup_group_after_participant_delete(self, wb, pid, group_name, role):
        groups = self._read_all(wb, SHEET_GROUPS)
        updated = []
        for g in groups:
            if g.get("name") != group_name:
                updated.append(g)
                continue
            suspect_id = g.get("suspect_id")
            interviewer_id = g.get("interviewer_id")
            if role == "S" and str(suspect_id) == str(pid):
                g = dict(g)
                g["suspect_id"] = ""
                if g.get("interviewer_id"):
                    updated.append(g)
                continue
            if role == "I" and str(interviewer_id) == str(pid):
                g = dict(g)
                g["interviewer_id"] = ""
                if g.get("suspect_id"):
                    updated.append(g)
                continue
            updated.append(g)
        self._write_all(wb, SHEET_GROUPS, updated)

    # ---- Groups ----

    def get_waiting_group(self):
        """Find a group that has a suspect but no interviewer assigned yet."""
        with _excel_lock:
            wb = self._load()
            try:
                groups = self._read_all(wb, SHEET_GROUPS)
                for g in sorted(groups, key=lambda x: x.get("created_at", "")):
                    if g.get("suspect_id") and not g.get("interviewer_id"):
                        return g
                return None
            finally:
                self._close(wb)

    def get_group_by_name(self, name):
        with _excel_lock:
            wb = self._load()
            try:
                for g in self._read_all(wb, SHEET_GROUPS):
                    if g["name"] == name:
                        return g
                return None
            finally:
                self._close(wb)

    def get_group_by_interviewer(self, interviewer_id):
        with _excel_lock:
            wb = self._load()
            try:
                iid = str(interviewer_id)
                for g in self._read_all(wb, SHEET_GROUPS):
                    if str(g.get("interviewer_id") or "") == iid:
                        return g
                return None
            finally:
                self._close(wb)

    def add_group(self, name, suspect_id, interviewer_id=None):
        with _excel_lock:
            wb = self._load()
            try:
                groups = self._read_all(wb, SHEET_GROUPS)
                groups.append({
                    "name": name,
                    "suspect_id": suspect_id,
                    "interviewer_id": interviewer_id,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                })
                self._write_all(wb, SHEET_GROUPS, groups)
                self._save(wb)
            finally:
                self._close(wb)

    def update_group(self, name, **kwargs):
        with _excel_lock:
            wb = self._load()
            try:
                groups = self._read_all(wb, SHEET_GROUPS)
                for g in groups:
                    if g["name"] == name:
                        g.update(kwargs)
                        self._write_all(wb, SHEET_GROUPS, groups)
                        self._save(wb)
                        return True
                return False
            finally:
                self._close(wb)

    def count_groups(self):
        with _excel_lock:
            wb = self._load()
            try:
                return len(self._read_all(wb, SHEET_GROUPS))
            finally:
                self._close(wb)

    # ---- Profiles ----

    def get_profile(self, participant_id):
        with _excel_lock:
            wb = self._load()
            try:
                pid = str(participant_id)
                for p in self._read_all(wb, SHEET_PROFILES):
                    if str(p.get("participant_id")) == pid:
                        return p
                return None
            finally:
                self._close(wb)

    def upsert_profile(self, participant_id, data_json):
        with _excel_lock:
            wb = self._load()
            try:
                profiles = self._read_all(wb, SHEET_PROFILES)
                for p in profiles:
                    if p["participant_id"] == participant_id:
                        p["data"] = data_json
                        p["submitted_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        self._write_all(wb, SHEET_PROFILES, profiles)
                        self._save(wb)
                        return
                profiles.append({
                    "participant_id": participant_id,
                    "data": data_json,
                    "submitted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                })
                self._write_all(wb, SHEET_PROFILES, profiles)
                self._save(wb)
            finally:
                self._close(wb)

    # ---- Availabilities ----

    def get_availabilities(self, group_name=None):
        with _excel_lock:
            wb = self._load()
            try:
                all_avail = self._read_all(wb, SHEET_AVAILABILITIES)
                if group_name:
                    return [a for a in all_avail if a.get("group_name") == group_name]
                return all_avail
            finally:
                self._close(wb)

    def upsert_availability(self, phone, group_name, role, slots):
        with _excel_lock:
            wb = self._load()
            try:
                availabilities = self._read_all(wb, SHEET_AVAILABILITIES)
                slots_json = json.dumps(slots, ensure_ascii=False)
                now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                for a in availabilities:
                    if a["phone"] == phone:
                        a["group_name"] = group_name
                        a["role"] = role
                        a["slots"] = slots_json
                        a["updated_at"] = now
                        self._write_all(wb, SHEET_AVAILABILITIES, availabilities)
                        self._save(wb)
                        return
                aid = self._next_id(wb, "next_avail_id")
                availabilities.append({
                    "id": aid,
                    "phone": phone,
                    "group_name": group_name,
                    "role": role,
                    "slots": slots_json,
                    "updated_at": now,
                })
                self._write_all(wb, SHEET_AVAILABILITIES, availabilities)
                self._save(wb)
            finally:
                self._close(wb)

    # ---- Appointments ----

    def get_appointments(self, phone=None):
        with _excel_lock:
            wb = self._load()
            try:
                all_appts = self._read_all(wb, SHEET_APPOINTMENTS)
                if phone:
                    return [a for a in all_appts if a["phone"] == phone]
                return all_appts
            finally:
                self._close(wb)

    def add_appointment(self, phone, role, time_slot):
        with _excel_lock:
            wb = self._load()
            try:
                appointments = self._read_all(wb, SHEET_APPOINTMENTS)
                aid = self._next_id(wb, "next_appt_id")
                appointments.append({
                    "id": aid,
                    "phone": phone,
                    "time_slot": time_slot,
                    "role": role,
                    "status": "confirmed",
                    "booked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                })
                self._write_all(wb, SHEET_APPOINTMENTS, appointments)
                self._save(wb)
                return aid
            finally:
                self._close(wb)

    def update_appointment(self, phone, **kwargs):
        with _excel_lock:
            wb = self._load()
            try:
                appointments = self._read_all(wb, SHEET_APPOINTMENTS)
                for a in appointments:
                    if a["phone"] == phone and a.get("status") == "confirmed":
                        a.update(kwargs)
                        self._write_all(wb, SHEET_APPOINTMENTS, appointments)
                        self._save(wb)
                        return True
                return False
            finally:
                self._close(wb)

    def get_appointment_by_id(self, aid):
        for a in self.get_appointments():
            if a.get("id") == aid:
                return a
        return None

    def update_appointment_by_id(self, aid, **kwargs):
        with _excel_lock:
            wb = self._load()
            try:
                appointments = self._read_all(wb, SHEET_APPOINTMENTS)
                found = False
                for a in appointments:
                    if a.get("id") == aid:
                        a.update(kwargs)
                        found = True
                        break
                if found:
                    self._write_all(wb, SHEET_APPOINTMENTS, appointments)
                    self._save(wb)
                return found
            finally:
                self._close(wb)

    def cancel_appointment(self, phone):
        return self.update_appointment(phone, status="cancelled")

    def delete_appointment_by_id(self, aid):
        with _excel_lock:
            wb = self._load()
            try:
                appointments = self._read_all(wb, SHEET_APPOINTMENTS)
                appointments = [a for a in appointments if a["id"] != aid]
                self._write_all(wb, SHEET_APPOINTMENTS, appointments)
                self._save(wb)
            finally:
                self._close(wb)
        _reconcile_group_allocations()

    def get_slot_bookings(self):
        """Return dict mapping time_slot -> set of roles booked."""
        with _excel_lock:
            wb = self._load()
            try:
                slot_map = {}
                for a in self._read_all(wb, SHEET_APPOINTMENTS):
                    if a.get("status") == "confirmed":
                        slot = a["time_slot"]
                        role = a.get("role", "")
                        slot_map.setdefault(slot, set()).add(role)
                return slot_map
            finally:
                self._close(wb)

    def get_confirmed_slot_set(self):
        """Return set of fully-booked time slots (both roles taken)."""
        slot_map = self.get_slot_bookings()
        return {s for s, roles in slot_map.items() if "S" in roles and "I" in roles}

    def has_booking(self, phone):
        """Check if a participant already has a confirmed booking."""
        with _excel_lock:
            wb = self._load()
            try:
                for a in self._read_all(wb, SHEET_APPOINTMENTS):
                    if a.get("phone") == phone and a.get("status") == "confirmed":
                        return True
                return False
            finally:
                self._close(wb)

    def get_my_booking(self, phone):
        """Get a participant's confirmed booking."""
        with _excel_lock:
            wb = self._load()
            try:
                for a in self._read_all(wb, SHEET_APPOINTMENTS):
                    if a.get("phone") == phone and a.get("status") == "confirmed":
                        return a
                return None
            finally:
                self._close(wb)

    # ---- Interview Questionnaires ----

    def get_interview_questionnaires(self, phone=None, phase=None):
        with _excel_lock:
            wb = self._load()
            try:
                rows = self._read_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES)
                if phone:
                    rows = [r for r in rows if r.get("phone") == phone]
                if phase:
                    rows = [r for r in rows if r.get("phase") == phase]
                return rows
            finally:
                self._close(wb)

    def upsert_interview_questionnaire(self, phone, role, phase, appointment_slot, answers_json):
        with _excel_lock:
            wb = self._load()
            try:
                rows = self._read_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES)
                now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                for r in rows:
                    if r.get("phone") == phone and r.get("phase") == phase:
                        r["role"] = role
                        r["appointment_slot"] = appointment_slot
                        r["answers_json"] = answers_json
                        r["submitted_at"] = now_str
                        self._write_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES, rows)
                        self._save(wb)
                        return r.get("id")
                qid = self._next_id(wb, "next_questionnaire_id")
                rows.append({
                    "id": qid,
                    "phone": phone,
                    "role": role,
                    "phase": phase,
                    "appointment_slot": appointment_slot,
                    "answers_json": answers_json,
                    "submitted_at": now_str,
                })
                self._write_all(wb, SHEET_INTERVIEW_QUESTIONNAIRES, rows)
                self._save(wb)
                return qid
            finally:
                self._close(wb)

    def get_questionnaire_override(self, phone, phase):
        with _excel_lock:
            wb = self._load()
            try:
                for r in self._read_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES):
                    if r.get("phone") == phone and r.get("phase") == phase:
                        return r
                return None
            finally:
                self._close(wb)

    def get_all_questionnaire_overrides(self):
        with _excel_lock:
            wb = self._load()
            try:
                return self._read_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES)
            finally:
                self._close(wb)

    def set_questionnaire_override(self, phone, phase, is_open):
        with _excel_lock:
            wb = self._load()
            try:
                rows = self._read_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES)
                now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                for r in rows:
                    if r.get("phone") == phone and r.get("phase") == phase:
                        r["is_open"] = 1 if is_open else 0
                        r["updated_at"] = now_str
                        self._write_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES, rows)
                        self._save(wb)
                        return r.get("id")
                oid = self._next_id(wb, "next_qoverride_id")
                rows.append({
                    "id": oid,
                    "phone": phone,
                    "phase": phase,
                    "is_open": 1 if is_open else 0,
                    "updated_at": now_str,
                })
                self._write_all(wb, SHEET_QUESTIONNAIRE_OVERRIDES, rows)
                self._save(wb)
                return oid
            finally:
                self._close(wb)

    # ---- Training Sessions ----

    def get_training_sessions(self, phone):
        """Get all training sessions for a participant."""
        with _excel_lock:
            wb = self._load()
            try:
                return [s for s in self._read_all(wb, SHEET_TRAINING_SESSIONS)
                        if s.get("phone") == phone]
            finally:
                self._close(wb)

    def start_training_session(self, phone, interviewer_id, session_num, avatar_setting, avatar_guilt):
        """Start a new training session."""
        with _excel_lock:
            wb = self._load()
            try:
                sessions = self._read_all(wb, SHEET_TRAINING_SESSIONS)
                # Check if this session already exists
                for s in sessions:
                    if s.get("phone") == phone and str(s.get("session_num")) == str(session_num):
                        return s.get("id")
                sid = self._next_id(wb, "next_training_session_id")
                sessions.append({
                    "id": sid,
                    "interviewer_id": interviewer_id,
                    "phone": phone,
                    "session_num": session_num,
                    "avatar_setting": avatar_setting,
                    "avatar_guilt": avatar_guilt,
                    "judgment": "",
                    "transcript": "",
                    "feedback": "",
                    "started_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "completed_at": "",
                })
                self._write_all(wb, SHEET_TRAINING_SESSIONS, sessions)
                self._save(wb)
                return sid
            finally:
                self._close(wb)

    def submit_training_session(self, phone, session_num, judgment, transcript, feedback):
        """Submit a completed training session with judgment and feedback."""
        with _excel_lock:
            wb = self._load()
            try:
                sessions = self._read_all(wb, SHEET_TRAINING_SESSIONS)
                for s in sessions:
                    if s.get("phone") == phone and str(s.get("session_num")) == str(session_num):
                        s["judgment"] = judgment
                        s["transcript"] = transcript
                        s["feedback"] = feedback
                        s["completed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        self._write_all(wb, SHEET_TRAINING_SESSIONS, sessions)
                        self._save(wb)
                        return True
                return False
            finally:
                self._close(wb)

    def count_completed_training_sessions(self, phone):
        """Count how many training sessions have been completed (have judgment + feedback)."""
        with _excel_lock:
            wb = self._load()
            try:
                count = 0
                for s in self._read_all(wb, SHEET_TRAINING_SESSIONS):
                    if s.get("phone") == phone and s.get("judgment") and s.get("feedback"):
                        count += 1
                return count
            finally:
                self._close(wb)

    def get_training_session(self, phone, session_num):
        """Get a specific training session."""
        with _excel_lock:
            wb = self._load()
            try:
                for s in self._read_all(wb, SHEET_TRAINING_SESSIONS):
                    if s.get("phone") == phone and str(s.get("session_num")) == str(session_num):
                        return s
                return None
            finally:
                self._close(wb)

    def update_training_transcript(self, phone, session_num, transcript):
        """Persist in-progress interview transcript (JSON string) without completing the session."""
        with _excel_lock:
            wb = self._load()
            try:
                sessions = self._read_all(wb, SHEET_TRAINING_SESSIONS)
                for s in sessions:
                    if s.get("phone") == phone and str(s.get("session_num")) == str(session_num):
                        s["transcript"] = transcript
                        self._write_all(wb, SHEET_TRAINING_SESSIONS, sessions)
                        self._save(wb)
                        return True
                return False
            finally:
                self._close(wb)

    # ---- Serious game choices (formerly results_serious_game.xlsx) ----

    def log_serious_game_choice(
        self, phone, participant_id, case, condition, step_index, video, choice,
    ):
        with _excel_lock:
            wb = self._load()
            try:
                rows = self._read_all(wb, SHEET_SERIOUS_GAME)
                rows.append({
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "phone": phone,
                    "participant_id": participant_id,
                    "case": case,
                    "condition": condition,
                    "step_index": step_index,
                    "video": video,
                    "choice": choice,
                })
                self._write_all(wb, SHEET_SERIOUS_GAME, rows)
                self._save(wb)
            finally:
                self._close(wb)

    def get_serious_game_choices(self, phone=None, participant_id=None):
        """Return choice rows for a suspect, sorted by step_index."""
        phone = (phone or "").strip()
        pid = str(participant_id) if participant_id is not None else ""
        with _excel_lock:
            wb = self._load()
            try:
                rows = self._read_all(wb, SHEET_SERIOUS_GAME)
                out = []
                for r in rows:
                    if phone and (r.get("phone") or "").strip() == phone:
                        out.append(r)
                    elif pid and str(r.get("participant_id") or "") == pid:
                        out.append(r)
                out.sort(key=lambda x: int(x.get("step_index") or 0))
                return out
            finally:
                self._close(wb)


store = ExcelStore()


def migrate_old_data():
    """Migrate data from experiment.db and results.json to Excel if they exist."""
    if not os.path.exists(EXCEL_FILE):
        init_excel()

    db_path = os.path.join(BASE_DIR, "experiment.db")
    results_path = os.path.join(BASE_DIR, "results.json")
    if not os.path.exists(db_path) and not os.path.exists(results_path):
        return

    # Only migrate if Excel is empty
    with _excel_lock:
        wb = store._load()
        try:
            existing = store._read_all(wb, SHEET_PARTICIPANTS)
            if existing:
                store._close(wb)
                return
        except Exception:
            store._close(wb)
            return
        store._close(wb)

    # Migrate from SQLite
    if os.path.exists(db_path):
        try:
            import sqlite3
            conn = sqlite3.connect(db_path)
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT * FROM participants").fetchall()
            for row in rows:
                d = dict(row)
                p = store.get_participant(d["phone"])
                if not p:
                    store.add_participant(**d)

            rows = conn.execute("SELECT * FROM groups_table").fetchall()
            for row in rows:
                d = dict(row)
                if not store.get_group_by_name(d["name"]):
                    store.add_group(d["name"], d.get("suspect_id"), d.get("interviewer_id"))

            rows = conn.execute("SELECT * FROM profiles").fetchall()
            for row in rows:
                d = dict(row)
                if not store.get_profile(d["participant_id"]):
                    store.upsert_profile(d["participant_id"], d["data"])

            conn.close()
            print("  [OK] 已从 experiment.db 迁移数据")
        except Exception as e:
            print(f"  [WARN] 从 SQLite 迁移失败: {e}")

    # Migrate from results.json
    if os.path.exists(results_path):
        try:
            with open(results_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            for a in data.get("availabilities", []):
                store.upsert_availability(a["phone"], a.get("group_name", ""), a.get("role", ""), a.get("slots", []))

            for a in data.get("appointments", []):
                with _excel_lock:
                    wb = store._load()
                    try:
                        appts = store._read_all(wb, SHEET_APPOINTMENTS)
                        appts.append({
                            "id": store._next_id(wb, "next_appt_id"),
                            "phone": a["phone"],
                            "time_slot": a["time_slot"],
                            "status": a.get("status", "confirmed"),
                            "booked_at": a.get("booked_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                        })
                        store._write_all(wb, SHEET_APPOINTMENTS, appts)
                        store._save(wb)
                    finally:
                        store._close(wb)

            print("  [OK] 已从 results.json 迁移数据")
        except Exception as e:
            print(f"  [WARN] 从 results.json 迁移失败: {e}")


def migrate_legacy_exports():
    """Import legacy results_serious_game.xlsx and training_feedback/ into experiment_data.xlsx."""
    ensure_sheets()

    if os.path.isfile(LEGACY_SERIOUS_GAME_XLSX):
        try:
            wb_old = load_workbook(LEGACY_SERIOUS_GAME_XLSX, read_only=True)
            ws = wb_old.active
            header_row = next(ws.iter_rows(min_row=1, max_row=1))
            headers = [c.value for c in header_row]
            with _excel_lock:
                wb = store._load()
                try:
                    rows = store._read_all(wb, SHEET_SERIOUS_GAME)
                    existing = {
                        (
                            str(r.get("phone") or ""),
                            str(r.get("step_index") or ""),
                            str(r.get("video") or ""),
                            str(r.get("choice") or ""),
                        )
                        for r in rows
                    }
                    added = 0
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        if not any(row):
                            continue
                        rec = dict(zip(headers, row))
                        phone = rec.get("Phone") or rec.get("phone") or ""
                        step_index = rec.get("StepIndex") or rec.get("step_index") or ""
                        video = rec.get("Video") or rec.get("video") or ""
                        choice = rec.get("Choice") or rec.get("choice") or ""
                        key = (str(phone), str(step_index), str(video), str(choice))
                        if key in existing:
                            continue
                        rows.append({
                            "timestamp": rec.get("Timestamp") or rec.get("timestamp") or "",
                            "phone": phone,
                            "participant_id": rec.get("ParticipantID") or rec.get("participant_id") or "",
                            "case": rec.get("Case") or rec.get("case") or "",
                            "condition": rec.get("Condition") or rec.get("condition") or "",
                            "step_index": step_index,
                            "video": video,
                            "choice": choice,
                        })
                        existing.add(key)
                        added += 1
                    if added:
                        store._write_all(wb, SHEET_SERIOUS_GAME, rows)
                        store._save(wb)
                        print(f"  [OK] 已从 results_serious_game.xlsx 合并 {added} 条到 experiment_data.xlsx")
                finally:
                    store._close(wb)
            wb_old.close()
        except Exception as e:
            print(f"  [WARN] 合并 results_serious_game.xlsx 失败: {e}")

    if not os.path.isdir(LEGACY_TRAINING_FEEDBACK_DIR):
        return

    try:
        imported = 0
        for fname in os.listdir(LEGACY_TRAINING_FEEDBACK_DIR):
            if not fname.endswith("_meta.json"):
                continue
            prefix = fname[: -len("_meta.json")]
            meta_path = os.path.join(LEGACY_TRAINING_FEEDBACK_DIR, fname)
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
            phone = meta.get("phone")
            session_num = meta.get("session_num")
            if not phone or session_num is None:
                continue
            existing_sessions = store.get_training_sessions(phone)
            if any(
                str(s.get("session_num")) == str(session_num) and s.get("feedback")
                for s in existing_sessions
            ):
                continue
            transcript_path = os.path.join(LEGACY_TRAINING_FEEDBACK_DIR, f"{prefix}_transcript.txt")
            feedback_path = os.path.join(LEGACY_TRAINING_FEEDBACK_DIR, f"{prefix}_feedback.txt")
            transcript = ""
            feedback = ""
            if os.path.isfile(transcript_path):
                with open(transcript_path, "r", encoding="utf-8") as f:
                    transcript = f.read()
            if os.path.isfile(feedback_path):
                with open(feedback_path, "r", encoding="utf-8") as f:
                    feedback = f.read()
            p = store.get_participant(phone)
            interviewer_id = p["id"] if p else ""
            store.start_training_session(
                phone, interviewer_id, session_num,
                meta.get("avatar_setting", ""), meta.get("avatar_guilt", ""),
            )
            store.submit_training_session(
                phone, session_num, meta.get("judgment", ""), transcript, feedback,
            )
            imported += 1
        if imported:
            print(f"  [OK] 已从 training_feedback/ 导入 {imported} 条培训记录到 experiment_data.xlsx")
    except Exception as e:
        print(f"  [WARN] 合并 training_feedback/ 失败: {e}")


# ====== Material & Training File Helpers ======

MATERIAL_SECTION_SPECS = [
    ("consent_suspect", "知情同意书（嫌疑人）", [
        "Consent Form - Suspect (Clean).docx",
    ], [BASE_DIR, MATERIALS_DIR]),
    ("consent_interviewer", "知情同意书（访谈员）", [
        "Consent Form - Interviewer (Clean).docx",
    ], [BASE_DIR, MATERIALS_DIR]),
    ("theory_sue", "B组 SUE 理论培训", [
        "SUE theoretical training group.docx",
        "SUE theoretical training group.pdf",
    ], [MATERIALS_DIR]),
    ("avatar_specific", "D组 特定 Avatar 培训说明", [
        "Specific avatar training group_instructions.docx",
    ], [MATERIALS_DIR]),
    ("avatar_general", "C组 通用 Avatar 培训说明", [
        "General avatar training group_instructions.docx",
    ], [MATERIALS_DIR]),
    ("control", "A组 对照组培训", [
        "Control group.docx",
    ], [MATERIALS_DIR]),
    ("interviewer_bg", "访谈员背景材料", [
        "Interviewer_Bg.docx",
    ], [MATERIALS_DIR]),
    ("avatar_persona_settings", "D组 Avatar 人格设定", [
        "Specific avatar training group_avatar persona settings.docx",
    ], [MATERIALS_DIR]),
    ("avatar_demo", "D组 Avatar 演示说明", [
        "Specific avatar training group_demo.docx",
    ], [MATERIALS_DIR]),
    ("background_info", "案件背景信息", [
        "(CN) Background Information.pdf",
    ], [BASE_DIR, MATERIALS_DIR]),
]

TRAINING_TYPE_SECTION = {
    "theory_sue": "theory_sue",
    "avatar_specific": "avatar_specific",
    "avatar_general": "avatar_general",
    "control": "control",
}

# Participant-downloadable PDFs (place files under materials/pdf/)
TRAINING_TYPE_PDFS = {
    "control": "A组_培训材料.pdf",
    "theory_sue": "B组_SUE理论培训.pdf",
    "avatar_specific": "B组_SUE理论培训.pdf",
    "avatar_general": "B组_SUE理论培训.pdf",
}

CASE_TYPE_PDFS = {
    "arson": {
        "interviewer": "纵火案_审讯者材料.pdf",
        "scenario_guilty": "纵火案_有罪嫌疑人情景.pdf",
        "scenario_innocent": "纵火案_无罪嫌疑人情景.pdf",
    },
    "theft": {
        "interviewer": "盗窃案_审讯者材料.pdf",
        "scenario_guilty": "盗窃案_有罪嫌疑人情景.pdf",
        "scenario_innocent": "盗窃案_无罪嫌疑人情景.pdf",
    },
}


def _pdf_download_entry(filename, display_name=None):
    """Return {name, url} if materials/pdf/<filename> exists."""
    safe = os.path.basename(filename or "")
    if not safe or safe != filename:
        return None
    path = os.path.join(MATERIALS_PDF_DIR, safe)
    if not os.path.isfile(path):
        return None
    return {
        "name": display_name or safe,
        "url": f"/api/download-material/{safe}",
    }


def _training_download_files(training_type):
    fname = TRAINING_TYPE_PDFS.get(training_type)
    if not fname:
        return []
    entry = _pdf_download_entry(fname)
    return [entry] if entry else []


def _interviewer_appointment_paired_suspect(phone):
    """Return (case_type, guilt, slot_matched) from suspect on the same booked time_slot."""
    p = store.get_participant(phone)
    if not p or p.get("role") != "I":
        return None, None, False
    booking = store.get_my_booking(phone)
    if not booking:
        return None, None, False
    time_slot = booking["time_slot"]
    slot_bookings = store.get_slot_bookings()
    if not _slot_has_both_roles(time_slot):
        return None, None, False
    suspect_phone = None
    for appt in store.get_appointments():
        if (
            appt.get("time_slot") == time_slot
            and appt.get("role") == "S"
            and appt.get("status") == "confirmed"
        ):
            suspect_phone = appt.get("phone")
            break
    if not suspect_phone:
        return None, None, True
    suspect = store.get_participant(suspect_phone)
    if not suspect:
        return None, None, True
    return suspect.get("case_type") or "arson", suspect.get("guilt") or "Innocent", True


def sanitize_text_for_tts(text):
    """Remove parenthetical/aside text before ElevenLabs speech synthesis."""
    if not text:
        return ""
    out = str(text)
    out = re.sub(r"（[^）]*）", "", out)
    out = re.sub(r"\([^)]*\)", "", out)
    out = re.sub(r"【[^】]*】", "", out)
    out = re.sub(r"\[[^\]]*\]", "", out)
    out = re.sub(r"\s+", " ", out).strip()
    return out


def deepseek_chat_completion(messages, *, temperature=0.7, max_tokens=100, timeout=15):
    """Call DeepSeek chat completions API; returns assistant message text."""
    payload = {
        "model": DEEPSEEK_CHAT_MODEL,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    # v4 models may emit empty content when thinking is enabled (reasoning in reasoning_content).
    if str(DEEPSEEK_CHAT_MODEL).startswith("deepseek-v4"):
        payload["thinking"] = {"type": "disabled"}

    resp = requests.post(
        "https://api.deepseek.com/v1/chat/completions",
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        },
        json=payload,
        timeout=timeout,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"DeepSeek API Error: {resp.text}")
    data = resp.json()
    choice = data["choices"][0]
    message = choice.get("message") or {}
    content = message.get("content")
    if content is None:
        content = ""
    content = str(content).strip()
    if not content:
        finish_reason = choice.get("finish_reason") or "unknown"
        raise RuntimeError(f"DeepSeek 返回空回复 (finish_reason={finish_reason})")
    return content


def _normalize_suspect_profile_data(profile_data):
    """Parse profile JSON from sheet; accept nested {profile: {...}} if present."""
    if profile_data is None:
        return None
    pd = profile_data
    if isinstance(pd, str):
        try:
            pd = json.loads(pd)
        except (json.JSONDecodeError, TypeError):
            return None
    if not isinstance(pd, dict):
        return None
    if isinstance(pd.get("profile"), dict):
        pd = pd["profile"]
    return pd


def _paired_suspect_and_profile_for_interviewer(phone):
    """
    Return (suspect_participant, profile_data_dict) for a D-group interviewer.
    Prefer groups sheet; fall back to suspect on the same booked time_slot
    (same logic as case materials / PDF download).
    """
    p = store.get_participant(phone)
    if not p or p.get("role") != "I":
        return None, None

    def _load_profile(suspect):
        if not suspect:
            return None, None
        row = store.get_profile(suspect["id"])
        if not row:
            return suspect, None
        return suspect, _normalize_suspect_profile_data(row.get("data"))

    group = store.get_group_by_interviewer(p["id"])
    if group and group.get("suspect_id"):
        suspect = store.get_participant_by_id(group["suspect_id"])
        found = _load_profile(suspect)
        if found[0] and found[1]:
            return found

    booking = store.get_my_booking(phone)
    if not booking:
        return None, None

    time_slot = (booking.get("time_slot") or "").strip()
    if not time_slot or not _slot_has_both_roles(time_slot):
        return None, None

    for appt in store.get_appointments():
        if appt.get("status") != "confirmed" or appt.get("role") != "S":
            continue
        if (appt.get("time_slot") or "").strip() != time_slot:
            continue
        suspect = store.get_participant((appt.get("phone") or "").strip())
        if suspect:
            return _load_profile(suspect)

    return None, None


def _interviewer_case_download_files(case_type):
    """One PDF per case: background + evidence for the interviewer."""
    spec = CASE_TYPE_PDFS.get(case_type)
    if not spec:
        return []
    entry = _pdf_download_entry(spec["interviewer"])
    return [entry] if entry else []


def _suspect_scenario_download_files(case_type, guilt):
    """One PDF for suspect: guilty or innocent scenario text."""
    spec = CASE_TYPE_PDFS.get(case_type)
    if not spec:
        return []
    scenario_key = "scenario_guilty" if guilt == "Guilty" else "scenario_innocent"
    entry = _pdf_download_entry(spec[scenario_key])
    return [entry] if entry else []


def _resolve_pdf_filepath(safe_name):
    """Resolve a safe basename under materials/pdf/ (no path traversal)."""
    if not safe_name or ".." in safe_name or safe_name != os.path.basename(safe_name):
        return None
    base = os.path.realpath(MATERIALS_PDF_DIR)
    path = os.path.realpath(os.path.join(MATERIALS_PDF_DIR, safe_name))
    if not path.startswith(base + os.sep) and path != base:
        return None
    return path if os.path.isfile(path) else None


def _extract_text_from_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if ext == ".pdf":
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n\n".join(parts)
    if ext in (".md", ".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


def _find_legacy_material_file(filenames, search_dirs):
    for d in search_dirs:
        for fname in filenames:
            path = os.path.join(d, fname)
            if os.path.isfile(path):
                return path
    return None


def _get_material_section(section_id):
    if not os.path.isfile(COMBINED_MATERIALS_MD):
        return ""
    try:
        with open(COMBINED_MATERIALS_MD, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception:
        return ""
    marker = f"<!-- section:{section_id} -->"
    if marker not in content:
        return ""
    parts = content.split(marker, 1)[1]
    next_marker = re.search(r"\n<!-- section:", parts)
    if next_marker:
        parts = parts[: next_marker.start()]
    return parts.strip().lstrip("-").strip()


def _calc_slide_min_seconds(text):
    n = len(text or "")
    if n < 100:
        return 5
    if n < 300:
        return 6
    if n < 600:
        return 7
    if n < 1000:
        return 8
    return 10


def _format_slide_html(block):
    """Turn a plain-text block into simple HTML without altering wording."""
    lines = block.split("\n")
    out = []
    buf = []

    def flush_para():
        nonlocal buf
        if buf:
            out.append(
                "<p>" + "<br>".join(html_module.escape(l) for l in buf) + "</p>"
            )
            buf = []

    for line in lines:
        s = line.strip()
        if not s:
            flush_para()
            continue
        if s.startswith("# "):
            flush_para()
            out.append(f'<h2 class="doc-h1">{html_module.escape(s[2:].strip())}</h2>')
        elif s.startswith("## "):
            flush_para()
            out.append(f'<h3 class="doc-h2">{html_module.escape(s[3:].strip())}</h3>')
        elif s.startswith("<") and s.endswith(">"):
            flush_para()
            out.append(f'<p class="doc-subtitle">{html_module.escape(s)}</p>')
        elif (
            len(s) <= 42
            and not s.endswith("。")
            and not s.endswith(".")
            and not s.endswith("；")
            and not s.endswith(":")
        ):
            flush_para()
            out.append(f'<h4 class="doc-section">{html_module.escape(s)}</h4>')
        else:
            buf.append(s)
    flush_para()
    return "".join(out) if out else f"<p>{html_module.escape(block)}</p>"


def _split_document_into_slides(text, max_chars=950):
    text = (text or "").strip()
    if not text:
        return []
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    if not blocks:
        blocks = [text]
    slides_raw = []
    cur = []
    cur_len = 0
    for block in blocks:
        bl = len(block)
        if cur and cur_len + bl + 2 > max_chars:
            slides_raw.append("\n\n".join(cur))
            cur = [block]
            cur_len = bl
        else:
            cur.append(block)
            cur_len += bl + 2
    if cur:
        slides_raw.append("\n\n".join(cur))
    return [
        {
            "index": i,
            "html": _format_slide_html(raw),
            "min_seconds": _calc_slide_min_seconds(raw),
        }
        for i, raw in enumerate(slides_raw)
    ]


def _efm_matrix_slide_html():
    """EFM 2×2 matrix (Turku example) — shown as a dedicated slide in SUE materials."""
    return (
        '<p style="margin-bottom:12px;">'
        "下图展示了<strong>证据框架矩阵（EFM）</strong>的四个象限示例"
        "（纵轴：证据来源强度；横轴：证据具体程度）："
        "</p>"
        '<table class="efm-matrix-table">'
        "<thead><tr><th scope=\"col\"></th>"
        "<th scope=\"col\">低具体程度</th>"
        "<th scope=\"col\">高具体程度</th></tr></thead>"
        "<tbody>"
        "<tr><th scope=\"row\">强来源</th>"
        "<td>我们有DNA证据表明你曾在上海市</td>"
        "<td>我们有DNA证据表明你曾在上海市图书馆</td></tr>"
        "<tr><th scope=\"row\">弱来源</th>"
        "<td>我们有信息表明你曾在上海市</td>"
        "<td>我们有信息表明你曾在上海市图书馆</td></tr>"
        "</tbody></table>"
    )


def _insert_efm_matrix_slide(slides):
    """Insert the EFM matrix table slide after the EFM introduction text."""
    matrix_slide = {
        "html": _efm_matrix_slide_html(),
        "min_seconds": 10,
    }
    insert_at = None
    for i, slide in enumerate(slides):
        html = slide.get("html") or ""
        if "证据框架矩阵" in html or "Evidence Framing Matrix" in html:
            insert_at = i + 1
            break
    if insert_at is None:
        insert_at = min(3, len(slides))
    merged = slides[:insert_at] + [matrix_slide] + slides[insert_at:]
    return [
        {**slide, "index": idx, "min_seconds": slide.get("min_seconds", 6)}
        for idx, slide in enumerate(merged)
    ]


def _build_combined_materials_md():
    sections = []
    for section_id, title, filenames, search_dirs in MATERIAL_SECTION_SPECS:
        path = _find_legacy_material_file(filenames, search_dirs)
        if not path:
            continue
        try:
            text = _extract_text_from_file(path)
        except Exception as e:
            logger.warning("Failed to read %s: %s", path, e)
            continue
        if not text.strip():
            continue
        sections.append(f"<!-- section:{section_id} -->\n\n# {title}\n\n{text.strip()}\n")
    if not sections:
        return False
    os.makedirs(MATERIALS_DIR, exist_ok=True)
    with open(COMBINED_MATERIALS_MD, "w", encoding="utf-8") as f:
        f.write("\n\n---\n\n".join(sections))
    return True


def _build_combined_materials_docx():
    if not os.path.isfile(COMBINED_MATERIALS_MD):
        return
    try:
        from docx import Document
        doc = Document()
        with open(COMBINED_MATERIALS_MD, "r", encoding="utf-8") as f:
            content = f.read()
        for block in re.split(r"<!-- section:\w+ -->\s*", content):
            block = block.strip().lstrip("-").strip()
            if not block:
                continue
            for line in block.split("\n"):
                line = line.rstrip()
                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line:
                    doc.add_paragraph(line)
            doc.add_page_break()
        if doc.paragraphs and doc.paragraphs[-1].text == "":
            pass
        doc.save(COMBINED_MATERIALS_DOCX)
    except Exception as e:
        logger.warning("Failed to build combined_materials.docx: %s", e)


def _export_feedback_prompt_md():
    if os.path.isfile(FEEDBACK_PROMPT_MD):
        return
    if not os.path.isfile(LEGACY_FEEDBACK_DOCX):
        return
    try:
        text = _extract_text_from_file(LEGACY_FEEDBACK_DOCX)
        if text.strip():
            os.makedirs(MATERIALS_PROMPTS_DIR, exist_ok=True)
            with open(FEEDBACK_PROMPT_MD, "w", encoding="utf-8") as f:
                f.write(text.strip() + "\n")
    except Exception as e:
        logger.warning("Failed to export feedback prompt: %s", e)


def _ensure_combined_materials_from_example():
    """Use shipped template when no Word/PDF sources exist (e.g. fresh server deploy)."""
    if os.path.isfile(COMBINED_MATERIALS_MD):
        return False
    if not os.path.isfile(COMBINED_MATERIALS_EXAMPLE):
        return False
    shutil.copy2(COMBINED_MATERIALS_EXAMPLE, COMBINED_MATERIALS_MD)
    logger.warning(
        "Using template %s — upload real combined_materials.md or add source docx/pdf and rebuild.",
        COMBINED_MATERIALS_EXAMPLE,
    )
    return True


def setup_materials_dir():
    """Organize materials/: merge legacy Word/PDF into combined_materials.md/.docx."""
    os.makedirs(MATERIALS_DIR, exist_ok=True)
    os.makedirs(MATERIALS_PDF_DIR, exist_ok=True)
    os.makedirs(MATERIALS_PROMPTS_DIR, exist_ok=True)

    if os.path.isdir(TRAINING_SRC_DIR):
        archive_dir = os.path.join(MATERIALS_DIR, "_legacy_sources")
        os.makedirs(archive_dir, exist_ok=True)
        for fname in os.listdir(TRAINING_SRC_DIR):
            src = os.path.join(TRAINING_SRC_DIR, fname)
            dst = os.path.join(archive_dir, fname)
            if os.path.isfile(src) and not os.path.exists(dst):
                shutil.copy2(src, dst)

    if not os.path.isfile(COMBINED_MATERIALS_MD):
        _build_combined_materials_md()
    if not os.path.isfile(COMBINED_MATERIALS_MD):
        _ensure_combined_materials_from_example()
    if os.path.isfile(COMBINED_MATERIALS_MD) and not os.path.isfile(COMBINED_MATERIALS_DOCX):
        _build_combined_materials_docx()

    _export_feedback_prompt_md()


AVATARS_FILE = os.path.join(DATA_DIR, "avatars.json")
LEGACY_AVATARS_FILE = os.path.join(BASE_DIR, "avatars.json")
if not os.path.isfile(AVATARS_FILE) and os.path.isfile(LEGACY_AVATARS_FILE):
    shutil.copy2(LEGACY_AVATARS_FILE, AVATARS_FILE)
LIVEAVATAR_VOICES_FILE = os.path.join(DATA_DIR, "liveavatar_voices.json")
_liveavatar_voice_map = {}


def load_voice_map():
    global _liveavatar_voice_map
    if os.path.isfile(LIVEAVATAR_VOICES_FILE):
        try:
            with open(LIVEAVATAR_VOICES_FILE, "r", encoding="utf-8") as f:
                _liveavatar_voice_map = json.load(f)
        except Exception:
            _liveavatar_voice_map = {}


def save_voice_map():
    with open(LIVEAVATAR_VOICES_FILE, "w", encoding="utf-8") as f:
        json.dump(_liveavatar_voice_map, f, ensure_ascii=False, indent=2)


def setup_liveavatar_voices():
    if not LIVEAVATAR_API_KEY or not ELEVENLABS_API_KEY:
        return

    load_voice_map()

    elevenlabs_ids = set()
    try:
        with open(AVATARS_FILE, "r", encoding="utf-8") as f:
            avatars_data = json.load(f)

        def collect_voice_ids(obj):
            if isinstance(obj, dict):
                vid = obj.get("elevenlabs_voice_id")
                if vid:
                    elevenlabs_ids.add(vid)
                for v in obj.values():
                    collect_voice_ids(v)
        collect_voice_ids(avatars_data)
    except Exception:
        return

    pending = [vid for vid in elevenlabs_ids if vid not in _liveavatar_voice_map]
    if not pending:
        return

    LIVEAVATAR_API_URL = "https://api.liveavatar.com/v1"

    try:
        secret_id = None
        list_resp = requests.get(
            f"{LIVEAVATAR_API_URL}/secrets",
            headers={"X-API-KEY": LIVEAVATAR_API_KEY},
            timeout=10,
        )
        if list_resp.status_code == 200:
            existing = list_resp.json().get("data", [])
            for s in existing:
                if s.get("secret_type") == "ELEVENLABS_API_KEY":
                    secret_id = s["id"]
                    break

        if not secret_id:
            create_resp = requests.post(
                f"{LIVEAVATAR_API_URL}/secrets",
                headers={"X-API-KEY": LIVEAVATAR_API_KEY, "Content-Type": "application/json"},
                json={
                    "secret_type": "ELEVENLABS_API_KEY",
                    "secret_value": ELEVENLABS_API_KEY,
                    "secret_name": "Interrogation App ElevenLabs",
                },
                timeout=10,
            )
            if create_resp.status_code == 200:
                secret_id = create_resp.json()["data"]["id"]
            else:
                return

        for evid in pending:
            try:
                import_resp = requests.post(
                    f"{LIVEAVATAR_API_URL}/voices/third_party",
                    headers={"X-API-KEY": LIVEAVATAR_API_KEY, "Content-Type": "application/json"},
                    json={"secret_id": secret_id, "voice_id": evid},
                    timeout=15,
                )
                if import_resp.status_code == 200:
                    lv_id = import_resp.json()["data"]["id"]
                    _liveavatar_voice_map[evid] = lv_id
            except Exception:
                pass

        save_voice_map()
    except Exception:
        pass


def get_liveavatar_voice_id(elevenlabs_voice_id):
    return _liveavatar_voice_map.get(elevenlabs_voice_id, elevenlabs_voice_id)


def is_valid_uuid(value):
    """Check if a value is a valid UUID (with or without urn:uuid: prefix)."""
    if not value:
        return False
    try:
        cleaned = str(value).replace("urn:uuid:", "")
        uuid.UUID(cleaned)
        return True
    except (ValueError, AttributeError):
        return False


def build_session_body(avatar_id, voice_id, context_id, language, opening_text=None, is_sandbox=True):
    """Build session/token request body, omitting voice_id if not a valid UUID."""
    persona = {
        "context_id": context_id,
        "language": language,
    }
    if is_valid_uuid(voice_id):
        persona["voice_id"] = voice_id
    else:
        print(f"  [WARN] voice_id '{voice_id}' is not a valid UUID — omitting from session body")

    body = {
        "mode": "FULL",
        "avatar_id": avatar_id,
        "avatar_persona": persona,
    }
    if is_sandbox:
        body["is_sandbox"] = True
    return body


# ====== Route Helpers ======

def _get_slot_group_map():
    """time_slot -> group_name (001–112). Bootstraps from existing bookings if meta empty."""
    raw = store.get_meta(META_SLOT_GROUPS, None)
    if raw is not None:
        try:
            data = json.loads(raw) if isinstance(raw, str) else raw
            if isinstance(data, dict):
                return {str(k): str(v) for k, v in data.items()}
        except Exception:
            pass

    slot_map = {}
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        slot = (appt.get("time_slot") or "").strip()
        if not slot:
            continue
        p = store.get_participant(appt.get("phone", ""))
        gn = (p.get("group_name") or "").strip() if p else ""
        if not gn:
            continue
        if slot in slot_map and slot_map[slot] != gn:
            logger.warning(
                "Slot %s has conflicting group numbers %s vs %s; keeping first",
                slot, slot_map[slot], gn,
            )
            continue
        slot_map[slot] = gn
    if slot_map:
        store.set_meta(META_SLOT_GROUPS, json.dumps(slot_map, ensure_ascii=False))
    return slot_map


def _set_slot_group_map(slot_map):
    store.set_meta(META_SLOT_GROUPS, json.dumps(slot_map, ensure_ascii=False))


def _parse_group_number(group_name):
    """Parse AAA group number; returns 0 if invalid."""
    s = str(group_name or "").strip()
    if not s.isdigit():
        return 0
    try:
        return int(s)
    except ValueError:
        return 0


def _normalize_group_name(raw):
    """Three-digit group string (001–999) or None."""
    n = _parse_group_number(raw)
    if n < 1 or n > 999:
        return None
    return f"{n:03d}"


def _active_booking_slots():
    """Time slots that still have at least one confirmed appointment."""
    slots = set()
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        slot = (appt.get("time_slot") or "").strip()
        if slot:
            slots.add(slot)
    return slots


def _groups_in_use():
    """Group numbers reserved by participants with a confirmed booking."""
    used = set()
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        p = store.get_participant((appt.get("phone") or "").strip())
        if not p:
            continue
        n = _parse_group_number(p.get("group_name"))
        if n > 0:
            used.add(n)
    return used


def _max_assigned_group_number():
    """Highest numeric group currently in use."""
    used = _groups_in_use()
    return max(used) if used else 0


def _next_available_group_number():
    """Lowest unused group number from 001 upward (fills gaps first)."""
    used = _groups_in_use()
    for n in range(1, MAX_GROUPS + 1):
        if n not in used:
            return f"{n:03d}"
    raise ValueError(f"Maximum group count ({MAX_GROUPS}) reached")


def _reconcile_group_allocations():
    """
    Release group numbers when slots are empty; drop stale slot_map / groups rows;
    clear group_name on participants without a confirmed booking.
    """
    used_nums = _groups_in_use()
    used_names = {f"{n:03d}" for n in used_nums}

    active_slots = _active_booking_slots()
    slot_map = _get_slot_group_map()
    new_map = {}
    for slot in active_slots:
        gn = slot_map.get(slot, "").strip()
        if not gn:
            for appt in store.get_appointments():
                if appt.get("status") != "confirmed":
                    continue
                if (appt.get("time_slot") or "").strip() != slot:
                    continue
                p = store.get_participant((appt.get("phone") or "").strip())
                gn = (p.get("group_name") or "").strip() if p else ""
                if gn:
                    break
        if gn:
            new_map[slot] = gn
    if new_map != slot_map:
        _set_slot_group_map(new_map)

    for p in store.get_all_participants():
        phone = (p.get("phone") or "").strip()
        if not phone or store.has_booking(phone):
            continue
        if (p.get("group_name") or "").strip() or (p.get("full_id") or "").strip():
            store.update_participant(phone, group_name="", full_id="")

    with _excel_lock:
        wb = store._load()
        try:
            groups = store._read_all(wb, SHEET_GROUPS)
            pruned = [
                g for g in groups
                if (g.get("name") or "").strip() in used_names
            ]
            if len(pruned) != len(groups):
                store._write_all(wb, SHEET_GROUPS, pruned)
                store._save(wb)
        finally:
            store._close(wb)


def _next_slot_group_number(slot_map):
    """Next group: lowest free number (001, 002, …), reusing gaps after admin deletes."""
    _ = slot_map
    return _next_available_group_number()


def admin_apply_participant_group(phone, new_group_raw, sync_slot=True):
    """
    Set three-digit group for participant(s); rebuild full_id from role/attrs.
    If sync_slot and participant has a confirmed booking, updates slot meta and
    all participants on that time_slot to the same group.
    """
    new_group = _normalize_group_name(new_group_raw)
    if not new_group:
        raise ValueError("组别须为 1–999 的三位数字（如 001、012）")

    p = store.get_participant(phone)
    if not p:
        raise ValueError("未找到参与者")

    phones = {phone}
    slot_updated = None
    if sync_slot:
        booking = store.get_my_booking(phone)
        if booking and booking.get("status") == "confirmed":
            slot = (booking.get("time_slot") or "").strip()
            if slot:
                slot_map = _get_slot_group_map()
                slot_map[slot] = new_group
                _set_slot_group_map(slot_map)
                slot_updated = slot
                for appt in store.get_appointments():
                    if (
                        appt.get("status") == "confirmed"
                        and (appt.get("time_slot") or "").strip() == slot
                    ):
                        ph = (appt.get("phone") or "").strip()
                        if ph:
                            phones.add(ph)

    updated = []
    for ph in sorted(phones):
        p2 = store.get_participant(ph)
        if not p2:
            continue
        suffix = participant_id_suffix(p2)
        full_id = make_full_id(new_group, p2["role"], suffix)
        store.update_participant(ph, group_name=new_group, full_id=full_id)
        _sync_group_record(new_group, p2["id"], p2["role"])
        updated.append({
            "phone": ph,
            "role": p2["role"],
            "group_name": new_group,
            "full_id": full_id,
        })
    return {
        "group_name": new_group,
        "slot_updated": slot_updated,
        "sync_slot": sync_slot,
        "updated": updated,
        "next_auto_group": _next_available_group_number(),
    }


def assign_group_for_time_slot(time_slot):
    """First booker on a slot gets the next group number; second booker shares it."""
    slot = (time_slot or "").strip()
    if not slot:
        raise ValueError("预约时段无效")
    _reconcile_group_allocations()
    slot_map = _get_slot_group_map()
    if slot in slot_map:
        return slot_map[slot]
    # Legacy rows: slot already has one booking with a group_name but no meta entry yet
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        if (appt.get("time_slot") or "").strip() != slot:
            continue
        p = store.get_participant(appt.get("phone", ""))
        gn = (p.get("group_name") or "").strip() if p else ""
        if gn:
            slot_map[slot] = gn
            _set_slot_group_map(slot_map)
            return gn
    group_name = _next_slot_group_number(slot_map)
    slot_map[slot] = group_name
    _set_slot_group_map(slot_map)
    return group_name


def _sync_group_record(group_name, participant_id, role):
    """Link suspect/interviewer IDs to the shared experiment group."""
    pid = str(participant_id)
    g = store.get_group_by_name(group_name)
    if role == "S":
        if g:
            store.update_group(group_name, suspect_id=pid)
        else:
            store.add_group(group_name, suspect_id=pid)
    else:
        if g:
            store.update_group(group_name, interviewer_id=pid)
        else:
            store.add_group(group_name, suspect_id="", interviewer_id=pid)


def assign_participant_group_on_booking(phone, time_slot):
    """Assign shared AAA group + full_id when a participant books a slot."""
    p = store.get_participant(phone)
    if not p:
        raise ValueError("未找到参与者")
    group_name = assign_group_for_time_slot(time_slot)
    suffix = participant_id_suffix(p)
    full_id = make_full_id(group_name, p["role"], suffix)
    store.update_participant(phone, group_name=group_name, full_id=full_id)
    _sync_group_record(group_name, p["id"], p["role"])
    return group_name, full_id


def pick_register_role():
    """Alternate S/I by registration count so roles stay balanced."""
    n_s = sum(1 for p in store.get_all_participants() if p.get("role") == "S")
    n_i = sum(1 for p in store.get_all_participants() if p.get("role") == "I")
    return "I" if n_i < n_s else "S"


def _suspect_combo_key(case_type, guilt):
    return f"{case_type or 'arson'}:{guilt or 'Innocent'}"


def _parse_suspect_combo_key(key):
    parts = (key or "").split(":", 1)
    if len(parts) == 2:
        return parts[0], parts[1]
    return "arson", "Innocent"


def _suspect_combo_on_slot(time_slot):
    """Return (case_type, guilt) of suspect already booked on this slot, or None."""
    slot = (time_slot or "").strip()
    if not slot:
        return None
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed" or appt.get("role") != "S":
            continue
        if (appt.get("time_slot") or "").strip() != slot:
            continue
        p = store.get_participant(appt.get("phone", ""))
        if p and p.get("case_type"):
            return p.get("case_type"), p.get("guilt") or "Innocent"
    return None


def _count_suspect_combos():
    counts = {_suspect_combo_key(c, g): 0 for c, g in SUSPECT_ATTR_CYCLE}
    for p in store.get_all_participants():
        if p.get("role") != "S":
            continue
        key = _suspect_combo_key(p.get("case_type"), p.get("guilt"))
        if key in counts:
            counts[key] += 1
    return counts


def _count_interviewer_training_types():
    counts = {t: 0 for t in TRAINING_TYPE_CYCLE}
    for p in store.get_all_participants():
        if p.get("role") != "I":
            continue
        tt = p.get("training_type")
        if tt in counts:
            counts[tt] += 1
    return counts


def pick_sequential_suspect_attrs():
    """Fixed round-robin across 4 suspect scenarios (registration order)."""
    n = sum(1 for p in store.get_all_participants() if p.get("role") == "S")
    case_type, guilt = SUSPECT_ATTR_CYCLE[n % len(SUSPECT_ATTR_CYCLE)]
    return guilt, case_type


def _pick_non_specific_training_rotating(counts):
    """Round-robin among A/B/C groups; each capped at TRAINING_TARGET_PER_TYPE."""
    idx = int(store.get_meta(META_INTERVIEWER_NON_SPECIFIC_SEQ, 0) or 0)
    for step in range(len(NON_SPECIFIC_TRAINING_TYPES)):
        tt = NON_SPECIFIC_TRAINING_TYPES[(idx + step) % len(NON_SPECIFIC_TRAINING_TYPES)]
        if counts.get(tt, 0) < TRAINING_TARGET_PER_TYPE:
            store.set_meta(META_INTERVIEWER_NON_SPECIFIC_SEQ, idx + step + 1)
            return tt
    min_c = min(counts.get(t, 0) for t in NON_SPECIFIC_TRAINING_TYPES)
    for tt in NON_SPECIFIC_TRAINING_TYPES:
        if counts.get(tt, 0) == min_c:
            return tt
    return NON_SPECIFIC_TRAINING_TYPES[0]


def pick_training_type_on_booking(time_slot):
    """
    Assign interviewer training at booking time only.
    - Slot already has a suspect → avatar_specific (D), unless D is full (28).
    - No suspect on slot, or D full → rotate among control / theory_sue / avatar_general (max 28 each).
    """
    counts = _count_interviewer_training_types()
    if _suspect_combo_on_slot(time_slot):
        if counts.get("avatar_specific", 0) < TRAINING_TARGET_PER_TYPE:
            return "avatar_specific"
        return _pick_non_specific_training_rotating(counts)
    return _pick_non_specific_training_rotating(counts)


def _interviewer_training_locked(p):
    """After attention checks, training group must not change."""
    return (
        int(p.get("sue_attention_passed") or 0) == 1
        or int(p.get("control_attention_passed") or 0) == 1
    )


def assign_interviewer_training_type(phone, time_slot=None):
    """
    Set or update training_type for an interviewer (before attention checks pass).
    Uses suspect on slot when known so 4×4 pairings stay decoupled.
    """
    p = store.get_participant(phone)
    if not p or p.get("role") != "I":
        return None
    if _interviewer_training_locked(p):
        return p.get("training_type")
    training_type = pick_training_type_on_booking(time_slot)
    store.update_participant(phone, training_type=training_type)
    if p.get("group_name"):
        suffix = participant_id_suffix({**p, "training_type": training_type})
        full_id = make_full_id(p["group_name"], "I", suffix)
        store.update_participant(phone, full_id=full_id)
    return training_type


def participant_id_suffix(participant):
    """Third segment of AAA-B-C (suspect: case + guilt; interviewer: training group A–D)."""
    if participant.get("role") == "S":
        case_letter = "T" if participant.get("case_type") == "theft" else "A"
        guilt_letter = "G" if participant.get("guilt") == "Guilty" else "I"
        return f"{case_letter}{guilt_letter}"
    if participant.get("role") == "I":
        return TRAINING_GROUP_LABELS.get(participant.get("training_type", ""), "X")
    return "X"


def make_full_id(group_name, role, suffix_code):
    """e.g. 001-S-AI (suspect arson innocent), 001-I-B (interviewer theory_sue)."""
    return f"{group_name}-{role}-{suffix_code}"


def consent_attention_required(participant):
    """Whether participant still needs post-consent comprehension check."""
    if int(participant.get("consent_attention_passed") or 0) == 1:
        return False
    if int(participant.get("attention_failed") or 0) == 1:
        return False
    # Grandfather users who already progressed before this check was added
    if participant.get("role") == "S":
        if int(participant.get("attention_passed") or 0) == 1:
            return False
        if int(participant.get("game_completed") or 0) == 1:
            return False
        if int(participant.get("profile_completed") or 0) == 1:
            return False
    if participant.get("role") == "I":
        if int(participant.get("completed") or 0) == 1:
            return False
        if int(participant.get("sue_attention_passed") or 0) == 1:
            return False
        if int(participant.get("control_attention_passed") or 0) == 1:
            return False
        if participant.get("flow_step") in ("booking", "booking_wait", "booking_matched", "case_info_done", "theory_practice", "avatar_training"):
            return False
    return True


ALLOWED_FLOW_STEPS = frozenset({
    "avatar_specific_intro_done",
    "avatar_general_intro_done",
    "sue_material_done",
    "sue_principle_done",
    "booking",
    "booking_wait",
    "booking_matched",
    "case_info_done",
    "case_evidence_recap",
    "theory_practice",
    "avatar_training",
})


def _slot_has_both_roles(time_slot):
    if not time_slot:
        return False
    slot_bookings = store.get_slot_bookings()
    roles = slot_bookings.get(time_slot, set())
    return "S" in roles and "I" in roles


def _booking_is_matched(time_slot):
    return _slot_has_both_roles(time_slot)


def _sync_slot_participants_flow_step(time_slot, is_matched):
    """When S+I both book the same slot, move waiting participants to booking_matched."""
    if not is_matched or not time_slot:
        return
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        if (appt.get("time_slot") or "").strip() != time_slot:
            continue
        ph = (appt.get("phone") or "").strip()
        if not ph:
            continue
        p = store.get_participant(ph)
        if not p:
            continue
        step = (p.get("flow_step") or "").strip()
        if step in ("booking", "booking_wait", ""):
            store.update_participant(ph, flow_step="booking_matched")


def _refresh_slot_match_state(time_slot):
    """Reconcile booking_wait / booking_matched for all participants on a slot."""
    if not time_slot:
        return
    is_matched = _slot_has_both_roles(time_slot)
    if is_matched:
        _sync_slot_participants_flow_step(time_slot, True)
        return
    for appt in store.get_appointments():
        if appt.get("status") != "confirmed":
            continue
        if (appt.get("time_slot") or "").strip() != time_slot:
            continue
        ph = (appt.get("phone") or "").strip()
        if not ph:
            continue
        p = store.get_participant(ph)
        if not p:
            continue
        step = (p.get("flow_step") or "").strip()
        if step == "booking_matched":
            store.update_participant(ph, flow_step="booking_wait")
        elif step in ("booking", ""):
            store.update_participant(ph, flow_step="booking_wait")


def admin_reschedule_appointment(aid, new_time_slot):
    """Admin-only: move a confirmed appointment to another slot (bypasses user booking rules)."""
    appt = store.get_appointment_by_id(aid)
    if not appt:
        raise ValueError("未找到预约")
    if appt.get("status") != "confirmed":
        raise ValueError("只能修改已确认的预约")

    phone = (appt.get("phone") or "").strip()
    role = appt.get("role")
    old_slot = (appt.get("time_slot") or "").strip()
    new_time_slot = (new_time_slot or "").strip()

    if not new_time_slot:
        raise ValueError("预约时间不能为空")
    if not phone or role not in ("S", "I"):
        raise ValueError("预约数据无效")

    if new_time_slot == old_slot:
        return {
            "success": True,
            "unchanged": True,
            "appointment_id": aid,
            "phone": phone,
            "role": role,
            "time_slot": old_slot,
        }

    try:
        datetime.strptime(new_time_slot, "%Y-%m-%d %H:%M")
    except ValueError:
        raise ValueError("时间格式无效，须为 YYYY-MM-DD HH:MM")

    if new_time_slot not in _candidate_booking_slots(role):
        raise ValueError(
            f"该时间段不在可选范围内（有效时段为 {BOOKING_MAX_DAYS} 天内的标准预约窗口）"
        )

    slot_bookings = store.get_slot_bookings()
    booked_roles = slot_bookings.get(new_time_slot, set())
    if role in booked_roles:
        role_label = "嫌疑人" if role == "S" else "审讯者"
        raise ValueError(f"目标时段已有{role_label}预约，请选择其他时段")

    if not store.get_participant(phone):
        raise ValueError("未找到参与者")

    if not store.update_appointment_by_id(aid, time_slot=new_time_slot):
        raise ValueError("更新预约失败")

    group_name, full_id = assign_participant_group_on_booking(phone, new_time_slot)

    training_type = None
    if role == "I":
        training_type = assign_interviewer_training_type(phone, new_time_slot)
    elif role == "S":
        for other in store.get_appointments():
            if (
                other.get("status") == "confirmed"
                and other.get("role") == "I"
                and (other.get("time_slot") or "").strip() == new_time_slot
            ):
                training_type = assign_interviewer_training_type(
                    (other.get("phone") or "").strip(),
                    new_time_slot,
                )

    is_matched_new = _slot_has_both_roles(new_time_slot)
    store.update_participant(
        phone,
        flow_step="booking_matched" if is_matched_new else "booking_wait",
    )
    _sync_slot_participants_flow_step(new_time_slot, is_matched_new)

    if old_slot:
        _refresh_slot_match_state(old_slot)
        for other in store.get_appointments():
            if (
                other.get("status") == "confirmed"
                and other.get("role") == "I"
                and (other.get("time_slot") or "").strip() == old_slot
            ):
                assign_interviewer_training_type(
                    (other.get("phone") or "").strip(),
                    old_slot,
                )

    _reconcile_group_allocations()

    return {
        "success": True,
        "appointment_id": aid,
        "phone": phone,
        "role": role,
        "old_time_slot": old_slot,
        "time_slot": new_time_slot,
        "is_matched": is_matched_new,
        "group_name": group_name,
        "full_id": full_id,
        "training_type": training_type,
    }


def _participant_display_id(p):
    """Provisional or final experiment ID for UI (booking, resume)."""
    if p.get("full_id"):
        return p["full_id"]
    group_name = p.get("group_name")
    if not group_name:
        return ""
    try:
        return make_full_id(group_name, p["role"], participant_id_suffix(p))
    except Exception:
        return ""


def _admin_progress_step(status, label, detail=""):
    """One checklist row for admin UI: done | current | pending | failed | skipped."""
    return {"status": status, "label": label, "detail": detail}


def _admin_training_type_label(training_type):
    labels = {
        "control": "A 组（对照）",
        "theory_sue": "B 组（理论+SUE）",
        "avatar_general": "C 组（通用 Avatar）",
        "avatar_specific": "D 组（专项 Avatar）",
    }
    return labels.get(training_type or "", training_type or "未分配")


def _interviewer_sue_training_steps(p):
    """Split SUE-style training into intro (optional) / material / principle / attention."""
    tt = (p.get("training_type") or "").strip()
    fs = (p.get("flow_step") or "").strip()
    sue_pass = int(p.get("sue_attention_passed") or 0) == 1
    steps = []

    intro_map = {
        "avatar_specific": ("avatar_specific_intro_done", "阅读 D 组专项说明"),
        "avatar_general": ("avatar_general_intro_done", "阅读 C 组通用说明"),
    }

    if tt in intro_map:
        intro_key, intro_label = intro_map[tt]
        if sue_pass or fs in ("sue_material_done", "sue_principle_done", "booking", "booking_wait", "booking_matched", "case_info_done", "theory_practice", "avatar_training", "finalize", "done"):
            intro_status = "done"
        elif fs == intro_key:
            intro_status = "done"
        else:
            intro_status = "current"
        steps.append(_admin_progress_step(intro_status, intro_label, "已完成说明阅读" if intro_status == "done" else "进行中"))

    material_label = "阅读 SUE 培训材料"
    principle_label = "SUE 原则理解检测"
    attention_label = "SUE 注意力检测（EFM）"

    post_material = {
        "sue_principle_done", "booking", "booking_wait", "booking_matched",
        "case_info_done", "theory_practice", "avatar_training", "finalize", "done",
    }
    post_principle = post_material - {"sue_principle_done"}

    if sue_pass:
        mat_st, prin_st, att_st = "done", "done", "done"
    elif fs == "sue_principle_done":
        mat_st, prin_st, att_st = "done", "done", "current"
    elif fs == "sue_material_done":
        mat_st, prin_st, att_st = "done", "current", "pending"
    elif tt in intro_map and fs == intro_map[tt][0]:
        mat_st, prin_st, att_st = "current", "pending", "pending"
    elif tt == "theory_sue" and fs in ("", "booking", "booking_wait", "booking_matched"):
        mat_st, prin_st, att_st = "current", "pending", "pending"
    else:
        mat_st, prin_st, att_st = "pending", "pending", "pending"

    if fs in post_material and not sue_pass and mat_st != "done":
        mat_st = "done"
    if fs in post_principle and not sue_pass and prin_st != "done":
        prin_st = "done"

    attempts = int(p.get("sue_attention_attempts") or 0)
    att_detail = "已通过" if att_st == "done" else (
        f"进行中（第 {attempts} 次尝试，最多 2 次）" if att_st == "current" and attempts else
        ("未开始" if att_st == "pending" else "进行中")
    )

    steps.append(_admin_progress_step(mat_st, material_label, "已完成材料阅读" if mat_st == "done" else "进行中"))
    steps.append(_admin_progress_step(prin_st, principle_label, "已通过" if prin_st == "done" else ("进行中" if prin_st == "current" else "未开始")))
    steps.append(_admin_progress_step(att_st, attention_label, att_detail))
    return steps


def build_participant_admin_progress(
    p,
    booking=None,
    matched=False,
    qs_pre=False,
    qs_post=False,
    training_sessions=0,
    blacklisted=False,
    include_steps=True,
):
    """Detailed stage breakdown for admin management UI."""
    phone = (p.get("phone") or "").strip()
    role = p.get("role")
    steps = []

    if int(p.get("attention_failed") or 0) == 1 or blacklisted:
        fail_detail = "注意力检测未通过或已被拉黑，无法继续"
        if int(p.get("sue_attention_attempts") or 0) >= 2:
            fail_detail = "SUE 注意力检测 2 次均未通过"
        return {
            "stage_label": "已终止",
            "stage_short": "已终止",
            "stage_tone": "failed",
            "steps": [_admin_progress_step("failed", "实验已终止", fail_detail)],
        }

    consent_needed = consent_attention_required(p)
    consent_done = int(p.get("consent_attention_passed") or 0) == 1
    if consent_needed:
        c_status = "done" if consent_done else "current"
        steps.append(_admin_progress_step(
            c_status, "知情同意理解检测",
            "已通过" if consent_done else "未完成",
        ))
    elif consent_done:
        steps.append(_admin_progress_step("done", "知情同意理解检测", "已通过"))

    if role == "S":
        att_pass = int(p.get("attention_passed") or 0) == 1
        game_done = int(p.get("game_completed") or 0) == 1
        profile_done = int(p.get("profile_completed") or 0) == 1
        has_booking = bool(booking and booking.get("status") == "confirmed")
        completed = int(p.get("completed") or 0) == 1

        steps.append(_admin_progress_step(
            "done" if att_pass else "current",
            "案件背景注意力检测",
            "已通过" if att_pass else "阅读案件背景并完成检测题",
        ))
        steps.append(_admin_progress_step(
            "done" if game_done else ("current" if att_pass else "pending"),
            "完成模拟行动游戏",
            "已完成" if game_done else ("进行中" if att_pass else "未开始"),
        ))
        steps.append(_admin_progress_step(
            "done" if profile_done else ("current" if game_done else "pending"),
            "填写个人信息问卷",
            "已提交" if profile_done else ("进行中" if game_done else "未开始"),
        ))
        steps.append(_admin_progress_step(
            "done" if has_booking else ("current" if profile_done else "pending"),
            "预约正式访谈时间",
            booking.get("time_slot", "") if has_booking else "尚未确认预约",
        ))
        steps.append(_admin_progress_step(
            "done" if matched else ("current" if has_booking else "pending"),
            "与审讯者配对成功",
            "同一时段双方均已确认" if matched else "等待审讯者确认同一时段",
        ))
        steps.append(_admin_progress_step(
            "done" if qs_pre else ("current" if matched else "pending"),
            "访谈前问卷",
            "已提交" if qs_pre else "未提交",
        ))
        steps.append(_admin_progress_step(
            "done" if qs_post else ("current" if qs_pre else "pending"),
            "访谈后问卷",
            "已提交" if qs_post else "未提交",
        ))
        steps.append(_admin_progress_step(
            "done" if completed else "pending",
            "完成实验并获取编号",
            p.get("full_id") or "",
        ))

    elif role == "I":
        tt = (p.get("training_type") or "").strip()
        has_booking = bool(booking and booking.get("status") == "confirmed")
        ctrl_pass = int(p.get("control_attention_passed") or 0) == 1
        sue_pass = int(p.get("sue_attention_passed") or 0) == 1
        training_done = ctrl_pass if tt == "control" else sue_pass
        fs = (p.get("flow_step") or "").strip()
        completed = int(p.get("completed") or 0) == 1

        steps.append(_admin_progress_step(
            "done" if has_booking else "current",
            "预约正式访谈时间",
            booking.get("time_slot", "") if has_booking else "尚未确认预约",
        ))
        steps.append(_admin_progress_step(
            "done" if tt else ("current" if has_booking else "pending"),
            "分配实验条件",
            _admin_training_type_label(tt) if tt else "预约后自动分配",
        ))

        if tt == "control":
            if ctrl_pass:
                steps.append(_admin_progress_step("done", "阅读对照组材料", "已完成"))
                steps.append(_admin_progress_step("done", "对照组注意力检测", "已通过"))
            else:
                steps.append(_admin_progress_step(
                    "current" if tt else "pending",
                    "阅读对照组材料",
                    "进行中" if tt else "未开始",
                ))
                steps.append(_admin_progress_step("pending", "对照组注意力检测", "未开始"))
        elif tt in ("theory_sue", "avatar_general", "avatar_specific"):
            steps.extend(_interviewer_sue_training_steps(p))
        elif tt:
            steps.append(_admin_progress_step("pending", "培训与检测", "未知培训类型"))

        steps.append(_admin_progress_step(
            "done" if matched else ("current" if training_done and has_booking else "pending"),
            "与嫌疑人配对成功",
            "同一时段双方均已确认" if matched else "等待嫌疑人确认同一时段",
        ))

        case_done = fs in ("case_info_done", "theory_practice", "avatar_training", "finalize", "done") or completed
        recap_passed = int(p.get("case_evidence_recap_passed") or 0) == 1
        case_read_current = training_done and matched and not recap_passed and not case_done
        case_read_done = recap_passed or case_done or fs == "case_evidence_recap"
        steps.append(_admin_progress_step(
            "done" if case_read_done else ("current" if case_read_current else "pending"),
            "阅读案件信息与证据",
            "已阅读" if case_read_done else ("进行中" if case_read_current else "未开始"),
        ))
        recap_current = case_read_done and not recap_passed and not case_done
        steps.append(_admin_progress_step(
            "done" if recap_passed or case_done else ("current" if recap_current else "pending"),
            "口头复述证据信息",
            "已通过审核" if (recap_passed or case_done) else ("进行中" if recap_current else "未开始"),
        ))

        if tt in ("avatar_specific", "avatar_general"):
            sessions_done = training_sessions >= 6
            sess_status = "done" if sessions_done else ("current" if case_done else "pending")
            steps.append(_admin_progress_step(
                sess_status,
                "虚拟审讯训练（6 次）",
                f"已完成 {training_sessions}/6" if training_sessions else "未开始",
            ))
        elif tt == "theory_sue":
            tp_done = completed or fs in ("theory_practice", "avatar_training", "finalize", "done")
            tp_current = case_done and not tp_done
            steps.append(_admin_progress_step(
                "done" if tp_done else ("current" if tp_current else "pending"),
                "Avatar 练习审讯",
                "已完成" if tp_done else ("进行中" if tp_current else "未开始"),
            ))
        elif tt == "control":
            fin_done = completed or fs in ("finalize", "done")
            fin_current = case_done and not fin_done
            steps.append(_admin_progress_step(
                "done" if fin_done else ("current" if fin_current else "pending"),
                "确认完成实验",
                "已完成" if fin_done else ("待确认" if fin_current else "未开始"),
            ))

        steps.append(_admin_progress_step(
            "done" if qs_pre else ("current" if completed else "pending"),
            "访谈前问卷",
            "已提交" if qs_pre else "未提交",
        ))
        steps.append(_admin_progress_step(
            "done" if qs_post else ("current" if qs_pre else "pending"),
            "访谈后问卷",
            "已提交" if qs_post else "未提交",
        ))
        steps.append(_admin_progress_step(
            "done" if completed else "pending",
            "完成实验并获取编号",
            p.get("full_id") or "",
        ))
    else:
        steps.append(_admin_progress_step("pending", "未知角色", ""))

    stage_label = "进行中"
    stage_tone = "progress"
    for s in steps:
        if s["status"] == "current":
            stage_label = s["label"]
            break
        if s["status"] == "failed":
            stage_label = s["label"]
            stage_tone = "failed"
            break
    else:
        if steps and all(s["status"] == "done" for s in steps):
            stage_label = "全部完成"
            stage_tone = "done"
        elif steps:
            for s in steps:
                if s["status"] == "pending":
                    stage_label = s["label"]
                    break

    return {
        "stage_label": stage_label,
        "stage_short": stage_label[:20] + ("…" if len(stage_label) > 20 else ""),
        "stage_tone": stage_tone,
        "training_sessions": training_sessions if role == "I" else 0,
        **({"steps": steps} if include_steps else {}),
    }


def compute_participant_resume(p, booking=None):
    """Return (resume_step, resume_label) for continuing the experiment flow."""
    if int(p.get("attention_failed") or 0) == 1:
        return "terminated", "注意力检测未通过，无法继续参与实验"

    if consent_attention_required(p):
        return "consent_attention", "请完成知情同意书理解检测"

    role = p.get("role")
    phone = p.get("phone")

    if role == "S":
        if not int(p.get("attention_passed") or 0):
            return "case_attention", "阅读案件背景并通过注意力检测"
        if not int(p.get("game_completed") or 0):
            return "serious_game", "完成模拟行动游戏"
        if not int(p.get("profile_completed") or 0):
            return "profile", "填写个人信息问卷"
        if booking is None:
            booking = store.get_my_booking(phone)
        if not booking:
            return "booking", "预约正式访谈时间"
        return "booking_done", "查看预约信息或返回首页"

    if role != "I":
        return "dashboard", "继续实验"

    if int(p.get("completed") or 0) == 1:
        return "all_done", "实验已完成，可查看编号"

    if booking is None:
        booking = store.get_my_booking(phone)

    if not booking:
        return "booking", "预约正式访谈时间"

    training_type = p.get("training_type", "")
    if not training_type:
        return "booking", "请先预约时间段以确定培训条件"

    training_done = False
    if training_type == "control":
        training_done = int(p.get("control_attention_passed") or 0) == 1
    elif training_type in ("theory_sue", "avatar_specific", "avatar_general"):
        training_done = int(p.get("sue_attention_passed") or 0) == 1

    if not training_done:
        return "interviewer_training", "继续培训材料与注意力检测"

    time_slot = booking.get("time_slot", "")
    if not _booking_is_matched(time_slot):
        return "booking_wait", "等待与嫌疑人配对同一时段"

    flow_step = (p.get("flow_step") or "").strip()
    recap_passed = int(p.get("case_evidence_recap_passed") or 0) == 1
    before_case_done = flow_step not in (
        "case_info_done", "theory_practice", "avatar_training", "finalize", "done",
    )

    if before_case_done:
        if flow_step == "case_evidence_recap" and not recap_passed:
            return "case_evidence_recap", "口头复述案件证据材料"
        return "case_info", "阅读配对案件的背景与证据"

    if training_type in ("avatar_specific", "avatar_general"):
        done_sessions = store.count_completed_training_sessions(phone)
        if done_sessions >= 6:
            return "all_done", "完成实验并获取编号"
        if flow_step == "case_info_done" or done_sessions > 0:
            return "avatar_training", f"继续虚拟审讯训练（已完成 {done_sessions}/6）"
        return "case_info", "阅读配对案件的背景与证据"

    if training_type == "theory_sue":
        if flow_step in ("case_info_done", "theory_practice"):
            return "theory_practice", "继续 Avatar 练习审讯"
        return "case_info", "阅读配对案件的背景与证据"

    if training_type == "control":
        if flow_step == "case_info_done":
            return "finalize", "完成实验"
        return "case_info", "阅读配对案件的背景与证据"

    return "dashboard", "继续实验"




# Semantic interpretations for profile fields (used in avatar/suspect prompts — not raw option labels)
PROFILE_OPTION_SEMANTICS = {
    "q2": {
        "男 Male": "性别为男性",
        "女 Female": "性别为女性",
        "其他 Other": "性别为其他或不愿说明",
    },
    "q3": {
        "城镇户口 Urban": "城镇户口",
        "农村户口 Rural": "农村户口",
    },
    "q5": {
        "3000以下": "月收入较低（约3000元以下）",
        "3000-5000": "月收入偏低（约3000–5000元）",
        "5000-10000": "月收入中等（约5000–10000元）",
        "10000-20000": "月收入较高（约10000–20000元）",
        "20000以上": "月收入很高（20000元以上）",
    },
    "q6": {
        "30%以下": "固定支出占收入比例较低",
        "30%-60%": "固定支出占收入约三至六成",
        "60%-90%": "固定支出占收入约六至九成，经济压力较大",
        "入不敷出": "支出常超过收入，经济压力很大",
    },
    "q7": {
        "不投资/只存定期": "理财上偏保守，主要储蓄",
        "低风险理财": "偏好低风险理财",
        "股票/基金": "有股票或基金投资习惯",
        "加密货币/高风险投资": "有高风险投资倾向",
    },
    "q8": {"是": "购买了商业保险", "否": "未购买商业保险"},
    "q9": {
        "0次": "近三年未换工作",
        "1-2次": "近三年换过一至两次工作",
        "3次及以上": "近三年换工作较频繁",
    },
    "q10": {"是": "有犯罪记录", "否": "无犯罪记录"},
    "q11": {"是": "有行政处罚记录", "否": "无行政处罚记录"},
    "q12": {
        "我的父母双全": "父母均在世，家庭结构完整",
        "我是单亲家庭或双亲已故": "成长于单亲家庭或双亲已故",
        "我是独生子女": "为独生子女",
        "我有兄弟姐妹": "有兄弟姐妹",
        "我有子女": "已有子女",
        "父母双全": "父母均在世，家庭结构完整",
        "单亲/双亲已故": "成长于单亲家庭或双亲已故",
        "独生子女": "为独生子女",
        "有兄弟姐妹": "有兄弟姐妹",
        "有子女": "已有子女",
    },
    "q13": {
        "0-2人": "日常亲密社交圈很小",
        "3-5人": "日常亲密社交圈较小",
        "6-10人": "日常亲密社交圈中等",
        "10人以上": "日常亲密社交圈较大",
    },
    "q14": {
        "单身": "目前单身",
        "恋爱中": "目前有伴侣",
        "已婚": "已婚",
        "离异/丧偶": "离异或丧偶",
    },
    "q15": {
        "很少（主要文字）": "日常私人通话很少，以文字为主",
        "10分钟以下": "日均私人通话较短",
        "10-30分钟": "日均私人通话约10–30分钟",
        "30分钟以上": "日均私人通话超过30分钟",
    },
    "q16": {
        "独居": "目前独居",
        "与伴侣/配偶同住": "与伴侣或配偶同住",
        "与父母/亲戚同住": "与父母或亲戚同住",
        "与室友/朋友合租": "与室友或朋友合租",
        "宿舍/集体居住": "宿舍或集体居住",
    },
    "q17": {
        "冷静寡言": "性格偏冷静寡言，不善主动表达",
        "脾气急躁": "性格偏急躁，容易不耐烦",
        "不愿吃亏": "性格较强硬，不愿吃亏",
        "随和实在": "性格随和务实",
    },
    "q18": {
        "经常吸烟": "有吸烟习惯",
        "频繁大量饮酒": "饮酒较多",
        "榕榔等": "有嚼槟榔等习惯",
        "无": "无明显成瘾习惯",
    },
    "q19": {
        "宅在家": "业余时间多宅在家中",
        "酒吧/夜店": "常去酒吧或夜店",
        "网吧/电竞酒店": "常去网吧或电竞场所",
        "棋牌室/麻将馆": "常去棋牌室或麻将馆",
        "咖啡馆/书店": "常去咖啡馆或书店",
        "健身房/公园": "常去健身房或公园",
    },
    "q20": {
        "非常健康": "自评身体健康",
        "慢性病需长期服药": "有慢性病需长期服药",
        "曾有/现有心理困扰（如抑郁/焦虑）": "曾有或现有心理困扰",
    },
    "q21": {
        "非常鲜艳亮眼": "偏好鲜艳亮眼的物品颜色",
        "浅色系": "偏好浅色系",
        "深色/低调色系": "偏好深色低调色系",
        "几乎全黑/极简": "偏好黑色或极简风格",
    },
    "q22": {"是": "正式访谈时会戴眼镜", "否": "正式访谈时不戴眼镜"},
    "q23": {"长发": "访谈时为长发", "短发": "访谈时为短发"},
}


def _profile_semantic_value(qid, raw):
    """Return interpretation for a profile field, not the survey option label."""
    if raw is None or raw == "" or raw == []:
        return None
    sem = PROFILE_OPTION_SEMANTICS.get(qid, {})
    if isinstance(raw, list):
        parts = [sem.get(str(v), str(v)) for v in raw if v]
        return "；".join(parts) if parts else None
    key = str(raw)
    return sem.get(key, key)


def profile_lines_for_prompt(pd):
    """Build profile bullet lines for LLM prompts using semantic interpretations."""
    if not pd:
        return []
    lines = []
    if pd.get("q1") not in (None, ""):
        lines.append(f"年龄约 {pd.get('q1')} 岁")
    for qid, label in [
        ("q2", "性别"), ("q3", "户口"), ("q4", "职业"), ("q5", "月收入"),
        ("q6", "支出压力"), ("q7", "投资习惯"), ("q8", "商业保险"),
        ("q9", "换工作频率"), ("q10", "犯罪记录"), ("q11", "行政处罚"),
        ("q12", "家庭结构"), ("q13", "社交圈"), ("q14", "婚姻状况"),
        ("q15", "私人通话"), ("q16", "居住情况"), ("q17", "性格倾向"),
        ("q18", "成瘾习惯"), ("q19", "常去场所"), ("q20", "健康状况"),
        ("q21", "颜色偏好"), ("q22", "是否戴眼镜"), ("q23", "发型"),
    ]:
        val = _profile_semantic_value(qid, pd.get(qid))
        if val:
            lines.append(f"{label}：{val}")
    return lines


def _iter_slot_times(start_hm, end_hm, step_minutes=BOOKING_SLOT_STEP_MINUTES):
    sh, sm = map(int, start_hm.split(":"))
    eh, em = map(int, end_hm.split(":"))
    cur = sh * 60 + sm
    end = eh * 60 + em
    while cur <= end:
        yield f"{cur // 60:02d}:{cur % 60:02d}"
        cur += step_minutes


def _booking_slot_bounds(role=None):
    """Suspects: earliest slot 24h ahead; interviewers: from now to 7 days."""
    now = datetime.now()
    if role == "I":
        min_dt = now
    else:
        min_dt = now + timedelta(hours=BOOKING_MIN_HOURS)
    max_dt = now + timedelta(days=BOOKING_MAX_DAYS)
    return min_dt, max_dt


def _interviewer_24h_cutoff():
    return datetime.now() + timedelta(hours=BOOKING_MIN_HOURS)


def _slot_within_interviewer_24h_window(slot_dt):
    """True if slot starts within the next BOOKING_MIN_HOURS (default 24h)."""
    return slot_dt < _interviewer_24h_cutoff()


def _is_valid_booking_slot(slot_dt, role=None, time_slot=None):
    min_dt, max_dt = _booking_slot_bounds(role)
    if not (min_dt <= slot_dt <= max_dt):
        return False
    if role == "I" and time_slot and _slot_within_interviewer_24h_window(slot_dt):
        if "S" not in store.get_slot_bookings().get(time_slot, set()):
            return False
    return True


def _interviewer_booking_error(slot_dt, time_slot):
    """Human-readable error when interviewer cannot book this slot."""
    _, max_dt = _booking_slot_bounds("I")
    now = datetime.now()
    if slot_dt > max_dt:
        return f"预约时间须在 {BOOKING_MAX_DAYS} 天以内"
    if slot_dt < now:
        return "不能预约已过去的时间"
    if _slot_within_interviewer_24h_window(slot_dt):
        if "S" not in store.get_slot_bookings().get(time_slot, set()):
            return (
                "未来 24 小时内仅可预约已有嫌疑人预约的时间段。"
                "请优先选择「嫌疑人已约」的时段，或选择 24 小时之后的空闲时段。"
            )
    return f"预约时间须在现在起至 {BOOKING_MAX_DAYS} 天以内"


def _booking_slot_windows_for_api():
    labels = list(BOOKING_SLOT_WINDOW_LABELS)
    return [
        {
            "label": labels[i] if i < len(labels) else f"时段{i + 1}",
            "start": w[0],
            "end": w[1],
        }
        for i, w in enumerate(BOOKING_SLOT_WINDOWS)
    ]


def _candidate_booking_slots(role=None):
    """Suspect slots from 24h ahead; interviewer slots from now, up to one week."""
    now = datetime.now()
    min_dt, max_dt = _booking_slot_bounds(role)
    today = now.replace(hour=0, minute=0, second=0, microsecond=0)
    slots = []
    day = today
    while day.date() <= max_dt.date():
        for start_hm, end_hm in BOOKING_SLOT_WINDOWS:
            for time_str in _iter_slot_times(start_hm, end_hm):
                h, m = map(int, time_str.split(":"))
                slot_dt = day.replace(hour=h, minute=m, second=0, microsecond=0)
                if min_dt <= slot_dt <= max_dt:
                    slots.append(slot_dt.strftime("%Y-%m-%d %H:%M"))
        day += timedelta(days=1)
    return slots


def build_serious_game_action_memory(suspect):
    """
    Outline memory for D-group avatar: serious-game storyline + guilty/innocent interview stance.
    Choice branches are filled from experiment_data.xlsx serious_game_choices when available.
    """
    if not suspect:
        return ""

    case = (suspect.get("case_type") or "arson").lower()
    if case not in ("arson", "theft"):
        case = "arson"
    guilt = suspect.get("guilt") or "Innocent"
    if guilt not in ("Guilty", "Innocent"):
        guilt = "Guilty" if "guilty" in str(guilt).lower() else "Innocent"

    parts = []
    opening = SG_ACTION_OPENING.get((case, guilt))
    if opening:
        parts.append(
            "# 模拟行动记忆（真实经历，可作回答依据；勿与案件背景矛盾）\n" + opening
        )

    hints_table = SG_CHOICE_MEMORY_HINTS.get((case, guilt), {})
    choice_map = {}
    for row in store.get_serious_game_choices(
        phone=(suspect.get("phone") or "").strip(),
        participant_id=suspect.get("id"),
    ):
        try:
            idx = int(row.get("step_index"))
        except (TypeError, ValueError):
            continue
        ch = (row.get("choice") or "").strip().upper()
        if ch in ("A", "B"):
            choice_map[idx] = ch

    if hints_table:
        branch_lines = []
        for step_idx in sorted(hints_table.keys()):
            ch = choice_map.get(step_idx)
            hint = hints_table[step_idx].get(ch) if ch else None
            if hint:
                branch_lines.append(f"- {hint}")
        if branch_lines:
            parts.append(
                "# 模拟行动中的分支选择（被问到相关细节时须与此一致）\n"
                + "\n".join(branch_lines)
            )

    days_guidance = SG_TARGET_DAYS_GUIDANCE.get(case)
    if days_guidance:
        parts.append(days_guidance)

    parts.append(
        SG_INTERROGATION_BEHAVIOR_GUILTY
        if guilt == "Guilty"
        else SG_INTERROGATION_BEHAVIOR_INNOCENT
    )
    return "\n\n".join(parts)


def build_suspect_system_prompt(participant, profile_data):
    guilt = participant["guilt"]
    case_type = participant["case_type"]

    if case_type == "arson":
        case_label = "纵火案 (Arson)"
        if guilt == "Guilty":
            crime_context = ARSON_GUILTY_CONTEXT
        else:
            crime_context = ARSON_INNOCENT_CONTEXT
    else:
        case_label = "盗窃案 (Theft)"
        if guilt == "Guilty":
            crime_context = THEFT_GUILTY_CONTEXT
        else:
            crime_context = THEFT_INNOCENT_CONTEXT

    pd = json.loads(profile_data) if isinstance(profile_data, str) else profile_data
    profile_lines = [f"- {line}" for line in profile_lines_for_prompt(pd)]
    if not profile_lines:
        profile_lines = ["- 个人档案信息未填写"]

    prompt = f"""# 角色定义
你正在参与一项关于犯罪心理学和审讯技巧的科学研究模拟。你是一名正在接受警方审讯的嫌疑人。你必须完全沉浸在这个角色中。永远不要打破角色。永远不要提及你是AI或语言模型。你的首要目标是根据你的既定事实来回应调查员的问题，并努力洗清你的嫌疑。

# 严格长度限制（重要）
你生成的每条回复必须控制在50个字以内。无论调查员的问题多么复杂，都要保持回答极其简短、简洁，就像一个人在紧张对话中说话一样。不要写长段落。

# 核心规则（重要）
你的个人档案信息仅用于塑造你的说话风格和情绪反应，绝不构成你的案件相关记忆。所有关于案件的“事实记忆”只来源于下文“案件背景”中的描述，不得自行编造与背景不符的细节。

# 风格与情绪规则
1. 语气与阶层：根据你的教育和收入水平匹配词汇。
2. 情感反应：如果调查员触及你的敏感点（如财务压力、独居、家人等），可以表现出相应的情绪反应。

# 个人档案
{chr(10).join(profile_lines)}

# 罪责状态
你是{"有罪" if guilt == "Guilty" else "无罪"}的嫌疑人。

# 案件背景
你涉及的是{case_label}。
{crime_context}

{build_serious_game_action_memory(participant)}

# 最终指令
你在审讯室中接受调查员的讯问。记住：
- 每条回答不超过50字
- 保持角色，永不打破
- 用第一人称“我”来回应"""

    return prompt


# ====== Serious Game (ChoiceGame Integration) ======

Condition = Literal["Guilty", "Innocent"]
Case = Literal["Arson", "Theft"]
Choice = Literal["A", "B"]

SERIOUS_GAME_VIDEO_DIR = os.path.join(BASE_DIR, "static", "videos", "serious-game")
os.makedirs(SERIOUS_GAME_VIDEO_DIR, exist_ok=True)


@dataclass(frozen=True)
class SeriousStep:
    video: str
    question: str | None = None
    a_label: str | None = None
    b_label: str | None = None
    next_default: int | None = None
    next_if_a: int | None = None
    next_if_b: int | None = None

    @property
    def has_choice(self) -> bool:
        return self.question is not None


def _sg_video(case: str, condition: str, step: str) -> str:
    """Serious-game MP4 filename: e.g. Arson_Guilty_1.mp4, Theft_Innocent_2-1.mp4."""
    return f"{case}_{condition}_{step}.mp4"


def build_serious_game_timeline(case: Case, condition: Condition) -> list[SeriousStep]:
    if case == "Theft":
        if condition == "Guilty":
            return [
                SeriousStep(video=_sg_video("Theft", "Guilty", "1"), question="请选择：", a_label="A）买一杯拿铁", b_label="B）买牛奶", next_if_a=1, next_if_b=2),
                SeriousStep(video=_sg_video("Theft", "Guilty", "2-1"), next_default=3),
                SeriousStep(video=_sg_video("Theft", "Guilty", "2-2"), next_default=3),
                SeriousStep(video=_sg_video("Theft", "Guilty", "3"), question="请选择：", a_label="A）快速穿过广场", b_label="B）以正常速度穿过广场", next_if_a=4, next_if_b=5),
                SeriousStep(video=_sg_video("Theft", "Guilty", "4-1"), next_default=6),
                SeriousStep(video=_sg_video("Theft", "Guilty", "4-2"), next_default=6),
                SeriousStep(video=_sg_video("Theft", "Guilty", "5"), question="请选择：", a_label="A）放进背包", b_label="B）放进口袋", next_if_a=7, next_if_b=8),
                SeriousStep(video=_sg_video("Theft", "Guilty", "6-1"), next_default=9),
                SeriousStep(video=_sg_video("Theft", "Guilty", "6-2"), next_default=9),
                SeriousStep(video=_sg_video("Theft", "Guilty", "7"), next_default=10),
            ]
        return [
            SeriousStep(video=_sg_video("Theft", "Innocent", "1"), question="请选择：", a_label="A）买一杯拿铁", b_label="B）买牛奶", next_if_a=1, next_if_b=2),
            SeriousStep(video=_sg_video("Theft", "Innocent", "2-1"), next_default=3),
            SeriousStep(video=_sg_video("Theft", "Innocent", "2-2"), next_default=3),
            SeriousStep(video=_sg_video("Theft", "Innocent", "3"), question="请选择：", a_label="A）拍大海", b_label="B）拍广场", next_if_a=4, next_if_b=5),
            SeriousStep(video=_sg_video("Theft", "Innocent", "4-1"), next_default=6),
            SeriousStep(video=_sg_video("Theft", "Innocent", "4-2"), next_default=6),
            SeriousStep(video=_sg_video("Theft", "Innocent", "5"), question="请选择：", a_label="A）看左侧邮轮", b_label="B）看右侧邮轮", next_if_a=7, next_if_b=8),
            SeriousStep(video=_sg_video("Theft", "Innocent", "6-1"), next_default=9),
            SeriousStep(video=_sg_video("Theft", "Innocent", "6-2"), next_default=9),
            SeriousStep(video=_sg_video("Theft", "Innocent", "7"), next_default=10),
        ]
    if condition == "Guilty":
        return [
            SeriousStep(video=_sg_video("Arson", "Guilty", "1"), question="请选择：", a_label="A）停在400米外的小路上，再步行过去", b_label="B）停在公共停车场", next_if_a=1, next_if_b=2),
            SeriousStep(video=_sg_video("Arson", "Guilty", "2-1"), next_default=3),
            SeriousStep(video=_sg_video("Arson", "Guilty", "2-2"), next_default=3),
            SeriousStep(video=_sg_video("Arson", "Guilty", "3"), question="请选择：", a_label="A）将汽油仔细倒在承重柱上", b_label="B）快速把汽油倒在地面上", next_if_a=4, next_if_b=5),
            SeriousStep(video=_sg_video("Arson", "Guilty", "4-1"), next_default=6),
            SeriousStep(video=_sg_video("Arson", "Guilty", "4-2"), next_default=6),
            SeriousStep(video=_sg_video("Arson", "Guilty", "5"), question="请选择：", a_label="A）走主路开车回家", b_label="B）走小路开车回家", next_if_a=7, next_if_b=8),
            SeriousStep(video=_sg_video("Arson", "Guilty", "6-1"), next_default=9),
            SeriousStep(video=_sg_video("Arson", "Guilty", "6-2"), next_default=9),
            SeriousStep(video=_sg_video("Arson", "Guilty", "7"), next_default=10),
        ]
    return [
        SeriousStep(video=_sg_video("Arson", "Innocent", "1"), question="请选择：", a_label="A）看动画电影", b_label="B）看动作电影", next_if_a=1, next_if_b=2),
        SeriousStep(video=_sg_video("Arson", "Innocent", "2-1"), next_default=3),
        SeriousStep(video=_sg_video("Arson", "Innocent", "2-2"), next_default=3),
        SeriousStep(video=_sg_video("Arson", "Innocent", "3"), question="请选择：", a_label="A）听轻柔音乐", b_label="B）听节奏感更强的音乐", next_if_a=4, next_if_b=5),
        SeriousStep(video=_sg_video("Arson", "Innocent", "4-1"), next_default=6),
        SeriousStep(video=_sg_video("Arson", "Innocent", "4-2"), next_default=6),
        SeriousStep(video=_sg_video("Arson", "Innocent", "5"), question="请选择：", a_label="A）走主路开车回家", b_label="B）走小路开车回家", next_if_a=7, next_if_b=8),
        SeriousStep(video=_sg_video("Arson", "Innocent", "6-1"), next_default=9),
        SeriousStep(video=_sg_video("Arson", "Innocent", "6-2"), next_default=9),
        SeriousStep(video=_sg_video("Arson", "Innocent", "7"), next_default=10),
    ]


def serious_game_video_url(video_name: str) -> str | None:
    """Local MP4 under static/videos/serious-game/ (filename must match SeriousStep.video)."""
    name = (video_name or "").strip()
    if not name or name != os.path.basename(name):
        return None
    path = os.path.join(SERIOUS_GAME_VIDEO_DIR, name)
    if not os.path.isfile(path):
        return None
    return f"/static/videos/serious-game/{name}"


def log_serious_game_choice(phone: str, pid: str, case: str, condition: str, step_index: int, video: str, choice: str):
    try:
        store.log_serious_game_choice(phone, pid, case, condition, step_index, video, choice)
    except Exception as e:
        logger.warning(f"Failed to log serious game choice: {e}")


# ====== Desktop-only gate (participants) ======

DESKTOP_REQUIRED_MESSAGE = (
    "请使用电脑（台式机或笔记本电脑）打开本实验系统。"
    "手机、平板等移动设备暂不支持参与实验。"
)

_PARTICIPANT_PAGE_PATHS = frozenset({"/", "/questionnaire/pre", "/questionnaire/post"})
_API_MOBILE_EXEMPT_PREFIXES = ("/api/admin/", "/api/health")


def _ua_indicates_mobile_or_tablet() -> bool:
    ua = request.headers.get("User-Agent", "") or ""
    if re.search(r"iPad|Tablet|PlayBook|Silk|Kindle|KFAPWI|Tablet PC", ua, re.I):
        return True
    if re.search(r"Android", ua, re.I) and not re.search(r"Mobile", ua, re.I):
        return True
    if re.search(
        r"Android.*Mobile|iPhone|iPod|webOS|BlackBerry|IEMobile|Opera Mini|Windows Phone",
        ua,
        re.I,
    ):
        return True
    return False


@app.before_request
def require_desktop_for_participants():
    if not _ua_indicates_mobile_or_tablet():
        return None
    path = request.path or ""
    if path.startswith("/manage"):
        return None
    for prefix in _API_MOBILE_EXEMPT_PREFIXES:
        if path.startswith(prefix):
            return None
    if path.startswith("/api/"):
        return jsonify({"error": DESKTOP_REQUIRED_MESSAGE, "require_desktop": True}), 403
    if path in _PARTICIPANT_PAGE_PATHS:
        return render_template("device_blocked.html"), 200
    return None


# ====== Routes ======

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/questionnaire/pre")
def questionnaire_pre_page():
    return render_template("questionnaire.html", phase=PHASE_PRE)


@app.route("/questionnaire/post")
def questionnaire_post_page():
    return render_template("questionnaire.html", phase=PHASE_POST)


@app.route("/api/register", methods=["POST"])
def register():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    if not phone or len(phone) < 8:
        return jsonify({"error": "请输入有效的手机号"}), 400

    if store.is_phone_blacklisted(phone):
        return jsonify({
            "error": "该手机号无法参与实验（已被限制参与，含超时未预约等情况）",
            "blacklisted": True,
        }), 403

    existing = store.get_participant(phone)
    if existing:
        blocked = _block_if_unbooked_timeout(existing)
        if blocked:
            return blocked
        if store.is_phone_blacklisted(phone) or existing.get("attention_failed"):
            return jsonify({
                "error": "该手机号无法参与实验（注意力检测未通过）",
                "blacklisted": True,
            }), 403
        return jsonify({"error": "该手机号已注册", "participant": dict(existing)}), 409

    role = pick_register_role()

    if role == "S":
        guilt, case_type = pick_sequential_suspect_attrs()
        pid = store.add_participant(
            phone=phone, role="S", group_name="",
            guilt=guilt, case_type=case_type,
        )

        if case_type == "arson":
            context_text = ARSON_GUILTY_CONTEXT if guilt == "Guilty" else ARSON_INNOCENT_CONTEXT
        else:
            context_text = THEFT_GUILTY_CONTEXT if guilt == "Guilty" else THEFT_INNOCENT_CONTEXT

        check_key = f"{case_type}_{guilt.lower()}"
        attention_questions = ATTENTION_CHECKS[check_key]

        return jsonify({
            "role": "S",
            "full_id": "",
            "display_id": "",
            "group_name": "",
            "case_type": case_type,
            "case_label": "纵火案 Arson" if case_type == "arson" else "盗窃案 Theft",
            "context": context_text,
            "attention_questions": attention_questions,
            "consent_attention_required": True,
        })

    store.add_participant(
        phone=phone, role="I", group_name="",
        training_type="",
    )

    return jsonify({
        "role": "I",
        "full_id": "",
        "group_name": "",
        "training_type": "",
        "paired": False,
        "consent_attention_required": True,
    })


@app.route("/api/consent-attention/<role>")
def consent_attention_questions(role):
    if role not in CONSENT_ATTENTION_CHECKS:
        return jsonify({"error": "无效角色"}), 400
    return jsonify({"questions": CONSENT_ATTENTION_CHECKS[role]})


@app.route("/api/verify-consent-attention", methods=["POST"])
def verify_consent_attention():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    answers = data.get("answers") or []

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    role = p.get("role")
    expected = CONSENT_ATTENTION_CHECKS.get(role, [])

    results = []
    all_correct = True
    for i, q in enumerate(expected):
        user_ans = answers[i] if i < len(answers) else -1
        correct = user_ans == q["answer"]
        if not correct:
            all_correct = False
        results.append({
            "question_index": i,
            "correct": correct,
            "correct_answer": q["answer"],
            "user_answer": user_ans,
        })

    if all_correct:
        store.update_participant(phone, consent_attention_passed=1)
        return jsonify({"all_correct": True, "results": results})

    store.blacklist_phone(phone, reason="consent_attention_failed")
    store.update_participant(phone, attention_failed=1)
    return jsonify({
        "all_correct": False,
        "results": results,
        "terminated": True,
        "message": "回答不正确，无法参与正式实验。您的手机号已被记录，无法再次参与。",
    })


@app.route("/api/verify-attention", methods=["POST"])
def verify_attention():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    answers = data.get("answers") or []

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    case_type = p.get("case_type") or "arson"
    guilt = p.get("guilt") or "Guilty"
    check_key = f"{case_type}_{guilt.lower()}"
    expected = ATTENTION_CHECKS[check_key]

    results = []
    all_correct = True
    for i, q in enumerate(expected):
        user_ans = answers[i] if i < len(answers) else -1
        correct = user_ans == q["answer"]
        if not correct:
            all_correct = False
        results.append({
            "question_index": i,
            "correct": correct,
            "correct_answer": q["answer"],
            "user_answer": user_ans,
        })

    if all_correct:
        store.update_participant(phone, attention_passed=1)
        return jsonify({
            "all_correct": True,
            "results": results,
            "retry_allowed": False,
        })

    store.blacklist_phone(phone, reason="attention_failed")
    return jsonify({
        "all_correct": False,
        "results": results,
        "retry_allowed": False,
        "terminated": True,
        "message": "回答不正确，无法参与正式实验。您的手机号已被记录，无法再次参与。",
    })


@app.route("/api/submit-profile", methods=["POST"])
def submit_profile():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    profile = data.get("profile") or {}

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    # Suspects must complete the serious game before profile
    if p.get("role") == "S" and p.get("game_completed", 0) != 1:
        return jsonify({"error": "请先完成模拟行动游戏再提交个人信息"}), 400

    profile_json = json.dumps(profile, ensure_ascii=False)
    store.upsert_profile(p["id"], profile_json)

    store.update_participant(phone, profile_completed=1, completed=1)
    p = store.get_participant(phone)
    full_id = p.get("full_id", "") or ""
    group_name = p.get("group_name", "") or ""

    if full_id and group_name:
        message = f"编号 {full_id} (第 {group_name} 组) 个人信息已保存。请继续预约正式访谈时间。"
    else:
        message = "个人信息已保存。请继续预约正式访谈时间；实验编号将在预约成功后分配。"

    return jsonify({
        "success": True,
        "full_id": full_id,
        "group_name": group_name,
        "message": message,
    })


@app.route("/api/training-material/<training_type>")
def training_material(training_type):
    """Training PDFs only (A 组 or B 组 SUE); C/D 组说明不提供下载。"""
    available_files = _training_download_files(training_type)

    materials = {
        "theory_sue": {
            "title": f"{training_group_label('theory_sue')} 组培训材料",
            "description": "你将阅读文字培训材料。请仔细阅读并完成后续检测与案件阅读，无需进行虚拟嫌疑人练习。",
            "type": "theory",
            "files": available_files,
        },
        "avatar_specific": {
            "title": f"{training_group_label('avatar_specific')} 组培训材料",
            "description": "你将阅读培训说明并完成 6 次虚拟嫌疑人审讯训练。",
            "type": "avatar",
            "files": available_files,
        },
        "avatar_general": {
            "title": f"{training_group_label('avatar_general')} 组培训材料",
            "description": "你将阅读培训说明并完成 6 次虚拟嫌疑人审讯训练。",
            "type": "avatar",
            "files": available_files,
        },
        "control": {
            "title": f"{training_group_label('control')} 组培训材料",
            "description": "你将阅读文字培训材料。请仔细阅读并完成后续检测与案件阅读。",
            "type": "control",
            "files": available_files,
        },
    }
    return jsonify(materials.get(training_type, {"title": "未知", "description": "", "files": []}))


@app.route("/api/interviewer-case-downloads", methods=["POST"])
def interviewer_case_downloads():
    """Case PDF for interviewer after booking, when same time_slot has both roles."""
    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    p = store.get_participant(phone)
    if not p or p.get("role") != "I":
        return jsonify({"error": "未找到审讯者"}), 404

    access = _interviewer_case_material_access(phone)
    if access.get("status") == "no_booking":
        return jsonify({
            "paired": False,
            "slot_matched": False,
            "case_type": None,
            "case_label": None,
            "files": [],
            "message": access.get("message", "请先完成时间预约后再下载案件材料。"),
        })
    if access.get("status") == "waiting_match":
        return jsonify({
            "paired": False,
            "slot_matched": False,
            "case_type": None,
            "case_label": None,
            "files": [],
            "message": access.get("message", ""),
        })

    case_type, guilt, slot_matched = _interviewer_appointment_paired_suspect(phone)
    if not slot_matched:
        return jsonify({
            "paired": False,
            "slot_matched": False,
            "case_type": None,
            "case_label": None,
            "files": [],
            "message": access.get("message", "暂无法下载案件 PDF。"),
        })

    files = _interviewer_case_download_files(case_type)
    case_label = "纵火案" if case_type == "arson" else "盗窃案"
    return jsonify({
        "paired": True,
        "slot_matched": True,
        "case_type": case_type,
        "case_label": case_label,
        "suspect_guilt": guilt,
        "files": files,
    })


@app.route("/api/suspect-case-downloads", methods=["POST"])
def suspect_case_downloads():
    """One scenario PDF for suspect (guilty or innocent) by registration assignment."""
    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    p = store.get_participant(phone)
    if not p or p.get("role") != "S":
        return jsonify({"error": "未找到嫌疑人"}), 404

    case_type = p.get("case_type") or "arson"
    guilt = p.get("guilt") or "Innocent"
    files = _suspect_scenario_download_files(case_type, guilt)
    case_label = "纵火案" if case_type == "arson" else "盗窃案"
    guilt_label = "有罪" if guilt == "Guilty" else "无罪"
    return jsonify({
        "case_type": case_type,
        "case_label": case_label,
        "guilt": guilt,
        "guilt_label": guilt_label,
        "files": files,
    })


@app.route("/api/download-material/<path:filename>")
def download_material(filename):
    safe_name = os.path.basename(filename)
    if safe_name != filename or ".." in filename:
        return jsonify({"error": "无效文件名"}), 400

    pdf_path = _resolve_pdf_filepath(safe_name)
    if pdf_path:
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=safe_name,
            mimetype="application/pdf",
        )

    allowed = {
        "combined_materials.docx": (COMBINED_MATERIALS_DOCX, COMBINED_DOWNLOAD_NAME),
        "combined_materials.md": (COMBINED_MATERIALS_MD, "培训材料合集.md"),
    }
    if safe_name not in allowed:
        return jsonify({"error": "文件不存在"}), 404

    filepath, download_name = allowed[safe_name]
    if not os.path.isfile(filepath):
        return jsonify({"error": "文件不存在"}), 404

    ext = os.path.splitext(filepath)[1].lower()
    mime_map = {
        ".md": "text/markdown; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    mimetype = mime_map.get(ext, "application/octet-stream")
    return send_file(filepath, as_attachment=True, download_name=download_name, mimetype=mimetype)


@app.route("/api/consent/<role>")
def consent_text(role):
    section_id = "consent_suspect" if role == "S" else "consent_interviewer"
    text_content = _get_material_section(section_id)
    if not text_content:
        text_content = "知情同意书文件不存在，请联系研究人员。请确认 materials/combined_materials.md 已生成。"
    return jsonify({"text": text_content})


@app.route("/api/consent-slides/<role>")
def consent_slides(role):
    section_id = "consent_suspect" if role == "S" else "consent_interviewer"
    title = "知情同意书（嫌疑人）" if role == "S" else "知情同意书（访谈员）"
    text_content = _get_material_section(section_id)
    if not text_content:
        return jsonify({
            "error": "知情同意书文件不存在，请联系研究人员。",
            "slides": [],
        }), 404
    return jsonify({
        "title": title,
        "slides": _split_document_into_slides(text_content),
    })


@app.route("/api/material-slides/<section_id>")
def material_slides(section_id):
    """Paginated training sections for guided reading UI."""
    allowed = {
        "consent_suspect", "consent_interviewer", "theory_sue",
        "avatar_specific", "avatar_general", "control",
    }
    if section_id not in allowed:
        return jsonify({"error": "无效的材料标识"}), 400
    text_content = _get_material_section(section_id)
    if not text_content:
        return jsonify({"error": "材料加载失败", "slides": []}), 404
    titles = {
        "theory_sue": "SUE 理论培训材料",
        "avatar_specific": "特定 Avatar 组培训说明",
        "avatar_general": "通用 Avatar 组培训说明",
        "control": "对照组培训材料",
    }
    slides = _split_document_into_slides(text_content)
    if section_id == "theory_sue":
        slides = _insert_efm_matrix_slide(slides)
    return jsonify({
        "title": titles.get(section_id, "培训材料"),
        "slides": slides,
    })


@app.route("/api/material-text/<training_type>")
def material_text(training_type):
    """Return the text content of training materials for inline display."""
    section_id = TRAINING_TYPE_SECTION.get(training_type)
    if not section_id:
        return jsonify({"text": ""})
    text_content = _get_material_section(section_id)
    if not text_content:
        text_content = "材料加载失败，请下载合集文件阅读。"
    return jsonify({"text": text_content, "source": "combined_materials.md"})


@app.route("/api/sue-training-text")
def sue_training_text():
    text_content = _get_material_section("theory_sue")

    if not text_content:
        legacy_docx = os.path.join(MATERIALS_DIR, "SUE theoretical training group.docx")
        legacy_pdf = os.path.join(MATERIALS_DIR, "SUE theoretical training group.pdf")
        if os.path.isfile(legacy_docx):
            try:
                text_content = _extract_text_from_file(legacy_docx)
            except Exception:
                pass
        if not text_content and os.path.isfile(legacy_pdf):
            try:
                text_content = _extract_text_from_file(legacy_pdf)
            except Exception:
                pass

    return jsonify({"text": text_content or "", "source": "combined_materials.md"})


@app.route("/api/sue-attention-config")
def sue_attention_config():
    """Expose SUE principle + EFM attention check content for the interviewer UI."""
    return jsonify({
        "principle": SUE_PRINCIPLE_ATTENTION,
        "efm": SUE_EFM_CHECK,
    })


@app.route("/api/verify-sue-principle-attention", methods=["POST"])
def verify_sue_principle_attention():
    """Single-item SUE principle check after reading theory materials (B/C/D groups). One wrong answer terminates."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    try:
        answer = int(data.get("answer"))
    except (TypeError, ValueError):
        answer = -1

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    if answer == SUE_PRINCIPLE_ATTENTION["answer"]:
        return jsonify({"correct": True})

    store.blacklist_phone(phone, reason="sue_principle_attention_failed")
    return jsonify({
        "correct": False,
        "terminated": True,
        "message": "回答不正确，无法参与正式实验。您的手机号已被记录，无法再次参与。",
    })


@app.route("/api/verify-sue-attention", methods=["POST"])
def verify_sue_attention():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    answers = data.get("answers") or {}
    attempt = data.get("attempt") or 1

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    results = []
    all_correct = True
    for i, stmt in enumerate(SUE_EFM_CHECK["statements"]):
        user_ans = answers.get(str(i), "")
        correct = user_ans == stmt["answer"]
        if not correct:
            all_correct = False
        results.append({
            "statement_index": i,
            "statement_text": stmt["text"],
            "correct": correct,
            "correct_answer": stmt["answer"],
            "user_answer": user_ans,
        })

    if all_correct:
        store.update_participant(
            phone,
            sue_attention_passed=1,
            sue_attention_attempts=attempt,
            flow_step="booking",
        )
        return jsonify({
            "all_correct": True,
            "attempt": attempt,
            "max_attempts": 2,
            "results": results,
            "categories": SUE_EFM_CHECK["categories"],
        })

    if attempt >= 2:
        store.blacklist_phone(phone, reason="sue_attention_failed")
        store.update_participant(phone, sue_attention_passed=0, sue_attention_attempts=attempt)
        return jsonify({
            "all_correct": False,
            "attempt": attempt,
            "max_attempts": 2,
            "results": results,
            "categories": SUE_EFM_CHECK["categories"],
            "terminated": True,
            "message": "两次回答均不正确，无法参与正式实验。您的手机号已被记录，无法再次参与。",
        })

    store.update_participant(phone, sue_attention_attempts=attempt)
    return jsonify({
        "all_correct": False,
        "attempt": attempt,
        "max_attempts": 2,
        "results": results,
        "categories": SUE_EFM_CHECK["categories"],
        "retry_allowed": True,
    })


@app.route("/api/verify-control-attention", methods=["POST"])
def verify_control_attention():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    answers = data.get("answers") or []
    retry = data.get("retry") or False

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    expected = CONTROL_ATTENTION_CHECKS
    results = []
    all_correct = True
    for i, q in enumerate(expected):
        user_ans = answers[i] if i < len(answers) else -1
        correct = user_ans == q["answer"]
        if not correct:
            all_correct = False
        results.append({
            "question_index": i,
            "correct": correct,
            "correct_answer": q["answer"],
            "user_answer": user_ans,
        })

    if all_correct:
        store.update_participant(phone, control_attention_passed=1, flow_step="booking")
        return jsonify({
            "all_correct": True,
            "results": results,
            "retry_allowed": False,
        })

    if retry:
        store.blacklist_phone(phone, reason="control_attention_failed")
        return jsonify({
            "all_correct": False,
            "results": results,
            "retry_allowed": False,
            "terminated": True,
            "message": "两次回答均不正确，无法参与正式实验。您的手机号已被记录，无法再次参与。",
        })

    return jsonify({
        "all_correct": False,
        "results": results,
        "retry_allowed": True,
    })


@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    user_message = (data.get("message") or "").strip()
    chat_history = data.get("history") or []

    if not user_message:
        return jsonify({"error": "消息不能为空"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    group = store.get_group_by_interviewer(p["id"])
    if not group or not group.get("suspect_id"):
        return jsonify({"error": "未找到配对的嫌疑人"}), 404

    suspect = store.get_participant_by_id(group["suspect_id"])
    if not suspect:
        return jsonify({"error": "未找到配对的嫌疑人"}), 404

    profile_row = store.get_profile(suspect["id"])
    if profile_row:
        profile_data = json.loads(profile_row["data"])
    else:
        profile_data = {f"q{i}": "未填写" for i in range(1, 24)}

    system_prompt = build_suspect_system_prompt(suspect, profile_data)

    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history[-20:]:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})

    try:
        reply = deepseek_chat_completion(
            messages, temperature=0.8, max_tokens=150, timeout=30,
        )
        return jsonify({"reply": reply.strip()})

    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/complete-interviewer", methods=["POST"])
def complete_interviewer():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    # For avatar training groups, check if they completed all 6 sessions
    training_type = p.get("training_type", "")
    if training_type in ("avatar_specific", "avatar_general"):
        completed_count = store.count_completed_training_sessions(phone)
        if completed_count < 6:
            return jsonify({
                "error": "请先完成全部 6 次虚拟审讯训练",
                "completed_count": completed_count,
                "total_required": 6,
            }), 400

    # Assign full_id now (at completion, not registration)
    full_id = p.get("full_id", "") or ""
    if not full_id:
        suffix = participant_id_suffix(p)
        full_id = make_full_id(p["group_name"], p["role"], suffix)
        store.update_participant(phone, full_id=full_id)

    store.update_participant(phone, completed=1, flow_step="done")

    return jsonify({
        "success": True,
        "full_id": full_id,
        "group_name": p["group_name"],
        "message": f"编号 {full_id} (第 {p['group_name']} 组) 已完成。请截图此页面并发送给研究人员。",
    })


@app.route("/api/flow-step", methods=["POST"])
def api_set_flow_step():
    """Persist fine-grained progress within multi-screen training flows."""
    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    step = (data.get("step") or "").strip()
    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    if step not in ALLOWED_FLOW_STEPS:
        return jsonify({"error": "无效的进度标识"}), 400
    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404
    store.update_participant(phone, flow_step=step)
    resume_step, resume_label = compute_participant_resume(p, store.get_my_booking(phone))
    return jsonify({
        "success": True,
        "flow_step": step,
        "resume_step": resume_step,
        "resume_label": resume_label,
    })


@app.route("/api/generate-results-docx", methods=["POST"])
def generate_results_docx():
    from docx import Document

    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    doc = Document()

    doc.add_heading("审讯实验 - 参与记录", level=1)
    doc.add_paragraph(f"编号: {p['full_id']}")
    doc.add_paragraph(f"角色: {'嫌疑人' if p['role'] == 'S' else '审讯者'}")
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if p["role"] == "S":
        doc.add_heading("案件背景", level=2)
        if p["case_type"] == "arson":
            ctx = ARSON_GUILTY_CONTEXT if p["guilt"] == "Guilty" else ARSON_INNOCENT_CONTEXT
        else:
            ctx = THEFT_GUILTY_CONTEXT if p["guilt"] == "Guilty" else THEFT_INNOCENT_CONTEXT
        doc.add_paragraph(ctx)

        profile_row = store.get_profile(p["id"])
        if profile_row:
            doc.add_heading("个人问卷", level=2)
            pd = json.loads(profile_row["data"]) if isinstance(profile_row["data"], str) else profile_row["data"]
            for q in PROFILE_QUESTIONS:
                qid = q["id"]
                val = pd.get(qid, "")
                if isinstance(val, list):
                    val = ", ".join(val)
                doc.add_paragraph(f"{q['label']}: {val}")
    else:
        training_labels = {k: f"{v} 组" for k, v in TRAINING_GROUP_LABELS.items()}
        doc.add_heading("培训信息", level=2)
        doc.add_paragraph(f"实验条件: {training_labels.get(p['training_type'], p['training_type'])}")
        doc.add_paragraph(f"所属组别: {p['group_name']}")

        if p["training_type"] == "theory_sue":
            doc.add_heading("SUE 策略核心原则", level=2)
            doc.add_paragraph("1. 证据延迟披露：不要一开始就出示所有证据。")
            doc.add_paragraph("2. 证据铺垫：在披露具体证据之前，先询问与证据相关的问题。")
            doc.add_paragraph("3. 陈述-证据一致性分析：系统地比对嫌疑人的陈述与已有证据。")
            doc.add_paragraph("4. 不提及证据来源：在初始阶段不透露证据的具体来源。")

    doc.add_heading("预约信息", level=2)
    appts = store.get_appointments(phone=phone)
    appt_found = None
    for a in appts:
        if a.get("status") == "confirmed":
            appt_found = a
            break
    if appt_found:
        doc.add_paragraph(f"预约时间: {appt_found['time_slot']}")
    else:
        doc.add_paragraph("尚未预约时间")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=f"{p['full_id']}_实验记录.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ====== Appointment APIs ======

def generate_time_slots(role=None):
    """Role-aware slots (S: 24h+; I: now+ with 24h-in-window suspect-only rule at book time)."""
    disabled = store.get_disabled_slots()
    return [s for s in _candidate_booking_slots(role) if s not in disabled]


def get_available_slots():
    """Return time slots that still have at least one spot open."""
    all_slots = generate_time_slots()
    fully_booked = store.get_confirmed_slot_set()
    return [s for s in all_slots if s not in fully_booked]


@app.route("/api/appointments/slots")
def api_slots():
    phone = (request.args.get("phone") or "").strip()
    role = None
    if phone:
        p = store.get_participant(phone)
        if p:
            role = p.get("role")
    all_slots = generate_time_slots(role)
    slot_bookings = store.get_slot_bookings()
    fully_booked = store.get_confirmed_slot_set()

    # Build grouped slots with role info
    groups = {}
    for s in all_slots:
        date_key = s[:10]
        groups.setdefault(date_key, []).append(s[11:])

    # Build slot info: for each slot, show which roles are booked
    slot_info = {}
    cutoff_24h = _interviewer_24h_cutoff() if role == "I" else None
    for s in all_slots:
        roles_booked = list(slot_bookings.get(s, set()))
        info = {
            "roles_booked": roles_booked,
            "fully_booked": s in fully_booked,
        }
        if role == "I":
            try:
                slot_dt = datetime.strptime(s, "%Y-%m-%d %H:%M")
            except ValueError:
                slot_dt = None
            within_24h = bool(slot_dt and slot_dt < cutoff_24h)
            suspect_booked = "S" in roles_booked
            interviewer_taken = "I" in roles_booked
            info["within_24h"] = within_24h
            info["suspect_booked"] = suspect_booked
            info["interviewer_bookable"] = (
                s not in fully_booked
                and not interviewer_taken
                and (not within_24h or suspect_booked)
            )
        slot_info[s] = info

    resp = jsonify({
        "all_slots": all_slots,
        "fully_booked": list(fully_booked),
        "slot_info": slot_info,
        "groups": groups,
        "role": role,
        "booking_min_hours": BOOKING_MIN_HOURS if role == "I" else BOOKING_MIN_HOURS,
        "interviewer_24h_suspect_only": role == "I",
    })
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/api/appointments/book", methods=["POST"])
def api_book_appointment():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    time_slot = (data.get("time_slot") or "").strip()

    if not phone or not time_slot:
        return jsonify({"error": "手机号和预约时间不能为空"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    role = p["role"]  # "S" or "I"

    try:
        slot_dt = datetime.strptime(time_slot, "%Y-%m-%d %H:%M")
        if not _is_valid_booking_slot(slot_dt, role, time_slot):
            if role == "I":
                err = _interviewer_booking_error(slot_dt, time_slot)
            else:
                err = f"预约时间须在 {BOOKING_MIN_HOURS} 小时之后至 {BOOKING_MAX_DAYS} 天以内"
            return jsonify({"error": err}), 400
    except ValueError:
        return jsonify({"error": "时间格式无效"}), 400

    all_slots = generate_time_slots(role)
    if time_slot not in all_slots:
        return jsonify({"error": "该时间段不在可选范围内"}), 400

    # Each person can only book one slot
    if store.has_booking(phone):
        return jsonify({"error": "您已有一个预约，每人只能预约一个时间段，且预约成功后无法更改时间。"}), 409

    # Check if this slot already has this role booked
    slot_bookings = store.get_slot_bookings()
    booked_roles = slot_bookings.get(time_slot, set())
    if role in booked_roles:
        role_label = "嫌疑人" if role == "S" else "审讯者"
        return jsonify({"error": f"该时间段已有{role_label}预约，请选择其他时间段"}), 409

    # Check if slot is fully booked (both roles taken)
    fully_booked = store.get_confirmed_slot_set()
    if time_slot in fully_booked:
        return jsonify({"error": "该时间段已被预约满"}), 409

    training_type = None
    if role == "I":
        training_type = assign_interviewer_training_type(phone, time_slot)
    elif role == "S":
        for appt in store.get_appointments():
            if (
                appt.get("status") == "confirmed"
                and appt.get("role") == "I"
                and (appt.get("time_slot") or "").strip() == time_slot
            ):
                assign_interviewer_training_type(appt.get("phone", ""), time_slot)

    try:
        group_name, full_id = assign_participant_group_on_booking(phone, time_slot)
    except ValueError as e:
        return jsonify({"error": str(e)}), 503

    aid = store.add_appointment(phone, role, time_slot)

    is_matched = _slot_has_both_roles(time_slot)
    store.update_participant(
        phone,
        flow_step="booking_matched" if is_matched else "booking_wait",
    )
    _sync_slot_participants_flow_step(time_slot, is_matched)

    p_after = store.get_participant(phone)
    if role == "I" and p_after:
        training_type = p_after.get("training_type")

    return jsonify({
        "success": True,
        "appointment_id": aid,
        "time_slot": time_slot,
        "role": role,
        "is_matched": is_matched,
        "group_name": group_name,
        "participant_id": full_id,
        "training_type": training_type,
        "training_group": training_group_label(training_type) if training_type else "",
    })


@app.route("/api/appointments/my", methods=["POST"])
def api_my_appointment():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    booking = store.get_my_booking(phone)
    p = store.get_participant(phone)
    role = p["role"] if p else None
    is_matched = False
    full_id = p["full_id"] if p else None
    if booking:
        time_slot = booking["time_slot"]
        is_matched = _slot_has_both_roles(time_slot)
        booking = dict(booking)
        booking["is_matched"] = is_matched
        if full_id:
            booking["participant_id"] = _participant_display_id(p) or full_id

    return jsonify({
        "appointment": booking,
        "role": role,
        "is_matched": is_matched,
        "participant_id": full_id,
    })


@app.route("/api/appointments/modify", methods=["POST"])
def api_modify_appointment():
    return jsonify({
        "error": "预约成功后无法更改时间。如需协助请联系研究人员。",
    }), 403


@app.route("/api/appointments/cancel", methods=["POST"])
def api_cancel_appointment():
    return jsonify({
        "error": "预约成功后无法取消或更改时间。如需协助请联系研究人员。",
    }), 403


# ====== Avatar APIs ======

LIVEAVATAR_API_URL = "https://api.liveavatar.com/v1"


def load_avatar_configs():
    if not os.path.exists(AVATARS_FILE):
        return {"generic": {}, "specific": {}}
    with open(AVATARS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def _avatar_appearance_key(suspect_profile):
    """D-group lookup key from suspect questionnaire (q2/q22/q23), or None."""
    pd = _normalize_suspect_profile_data(suspect_profile)
    if not pd:
        return None

    gender_raw = pd.get("q2", "")
    is_male = "男" in str(gender_raw) or "Male" in str(gender_raw)
    gender_key = "male" if is_male else "female"

    glasses_raw = pd.get("q22", "")
    has_glasses = str(glasses_raw).strip() in ("是", "Yes", "yes", "true", "1")
    glasses_key = "glasses" if has_glasses else "noglasses"

    hair_raw = pd.get("q23", "")
    is_long = "长" in str(hair_raw) or "Long" in str(hair_raw)
    hair_key = "long" if is_long else "short"

    return f"{gender_key}_{glasses_key}_{hair_key}"


def resolve_avatar_config(training_type, suspect_profile):
    avatars = load_avatar_configs()

    if training_type == "avatar_general":
        return avatars.get("generic", {})

    specific = avatars.get("specific", {})
    if not specific:
        return avatars.get("generic", {})

    config_key = _avatar_appearance_key(suspect_profile)
    if not config_key:
        return avatars.get("generic", {})

    chosen = specific.get(config_key)
    if chosen:
        return chosen
    return avatars.get("generic", {})


def build_avatar_system_prompt(suspect, profile_data):
    guilt = suspect["guilt"]
    case_type = suspect["case_type"]

    if case_type == "arson":
        crime_context = ARSON_GUILTY_CONTEXT if guilt == "Guilty" else ARSON_INNOCENT_CONTEXT
    else:
        crime_context = THEFT_GUILTY_CONTEXT if guilt == "Guilty" else THEFT_INNOCENT_CONTEXT

    pd = _normalize_suspect_profile_data(profile_data)
    if not pd:
        pd = json.loads(profile_data) if isinstance(profile_data, str) else profile_data

    profile_lines = []
    profile_lines.append(f"年龄: {pd.get('q1', '未填写')}")
    profile_lines.append(f"性别: {pd.get('q2', '未填写')}")
    profile_lines.append(f"职业: {pd.get('q4', '未填写')}")
    profile_lines.append(f"月收入: {pd.get('q5', '未填写')}")
    profile_lines.append(f"居住情况: {pd.get('q16', '未填写')}")
    profile_lines.append(f"朋友评价: {pd.get('q17', '未填写')}")

    action_memory = build_serious_game_action_memory(suspect)

    prompt = f"""# 角色定义
你正在参与犯罪心理学审讯研究。你是一名正在接受警方审讯的嫌疑人。完全沉浸在这个角色中，永不打破角色，永不提及你是AI。

# 严格长度限制
每条回复不超过50字。保持极其简短，就像在紧张对话中说话一样。

# 核心规则
个人档案仅塑造说话风格；案件与行动记忆来自「案件背景」与「模拟行动记忆」。

# 个人档案
{chr(10).join(profile_lines)}

# 罪责状态
你是{"有罪" if guilt == "Guilty" else "无罪"}的嫌疑人。

# 案件背景
{crime_context}

{action_memory}

# 指令
在审讯室中接受调查员的讯问。记住：每条回答不超过50字，保持角色，用第一人称“我”来回应。"""

    return prompt


@app.route("/api/avatar/config", methods=["POST"])
def api_avatar_config():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general", "theory_sue"):
        return jsonify({"error": "此培训类型不需要 Avatar"}), 400

    if training_type == "theory_sue":
        g = training_group_label("theory_sue")
        return jsonify({"error": f"{g} 组仅需阅读文字材料，无需虚拟嫌疑人练习"}), 400

    if training_type in ("avatar_specific", "avatar_general"):
        return jsonify({
            "error": "请从「虚拟审讯训练（6 次）」列表按顺序进入每次训练",
            "use_six_session_flow": True,
        }), 400

    effective_type = training_type

    suspect, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)
    system_prompt = None
    if suspect and suspect_profile:
        system_prompt = build_avatar_system_prompt(suspect, suspect_profile)

    avatar_config = resolve_avatar_config(effective_type, suspect_profile)

    return jsonify({
        "avatar_id": avatar_config.get("avatar_id", ""),
        "face_id": avatar_config.get("face_id", ""),
        "elevenlabs_voice_id": avatar_config.get("elevenlabs_voice_id", ""),
        "language": avatar_config.get("language", "zh"),
        "opening_text": avatar_config.get("opening_text", "你有什么要问的？"),
        "system_prompt": system_prompt or avatar_config.get("prompt", ""),
    })


def build_avatar_training_system_prompt(training_type, avatar_setting, avatar_guilt, suspect, profile_data):
    """Build system prompt for avatar training sessions.
    For avatar_specific: includes suspect profile + case context.
    For avatar_general: uses generic case context without suspect profile."""
    setting_prompt = AVATAR_SETTING_PROMPTS.get(avatar_setting, "")

    if training_type == "avatar_specific" and suspect and profile_data:
        pd = _normalize_suspect_profile_data(profile_data)
        if not pd:
            pd = json.loads(profile_data) if isinstance(profile_data, str) else profile_data
        profile_lines = profile_lines_for_prompt(pd)

        guilt = suspect.get("guilt", "Guilty")
        suspect_case = suspect.get("case_type", "arson")
        if suspect_case == "arson":
            crime_context = ARSON_GUILTY_CONTEXT if guilt == "Guilty" else ARSON_INNOCENT_CONTEXT
        else:
            crime_context = THEFT_GUILTY_CONTEXT if guilt == "Guilty" else THEFT_INNOCENT_CONTEXT

        action_memory = build_serious_game_action_memory(suspect)

        prompt = f"""# 角色定义
你正在参与犯罪心理学审讯研究。你是一名正在接受警方审讯的嫌疑人。完全沉浸在这个角色中，永不打破角色，永不提及你是AI。

# 严格长度限制
每条回复不超过50字。保持极其简短，就像在紧张对话中说话一样。

# 核心规则
个人档案仅塑造说话风格与情绪；案件与当晚行动的事实记忆来自「案件背景」与「模拟行动记忆」，不得编造与之矛盾的情节。

# 行为设定（审讯风格）
{setting_prompt}

# 个人档案
{chr(10).join(profile_lines)}

# 罪责状态
你是{"有罪" if guilt == "Guilty" else "无罪"}的嫌疑人。

# 案件背景（动机与处境）
{crime_context}

{action_memory}

# 指令
在审讯室中接受调查员的讯问。记住：每条回答不超过50字，保持角色，用第一人称"我"来回应。"""
    else:
        if avatar_guilt == "guilty":
            guilt_block = """你是“有罪但强烈否认”的嫌疑人。你真实做过以下行为：
- 购买过用于装载爆炸物材料的行李箱；
- 将行李箱埋藏在 King's Wood 附近。
但你绝不直接承认任何犯罪意图或袭击计划。"""
        else:
            guilt_block = """你是“无罪”的嫌疑人。你没有参与任何恐袭计划，也没有犯罪意图。
你可以承认与日常生活相关、可被强证据证明的普通活动，但必须坚决否认犯罪参与。"""

        prompt = f"""# 角色定义
你不是助手，禁止用“我能帮你什么”等客服式开场。你是正在接受警方问询的嫌疑人 Charlie，必须始终保持角色，不得提及你是AI。

# 回答长度与风格（严格）
1) 每次只输出一句简短口语化回答，不超过50字。
2) 不列点、不解释规则、不重复上一句。
3) 除非被追问，不主动扩展信息。

# 当前人格变量
{setting_prompt}

# 案件背景（通用 Avatar 组）
{GENERAL_AVATAR_CASE_CONTEXT_ZH}

# 个人信息
{GENERAL_AVATAR_IDENTITY_ZH}

# 罪责状态
{guilt_block}

# 时间范围
若调查员问及与本案调查窗口无关的其他日期，不要编造；简短将话题引回调查关注的时间段。

# 回应策略（严格执行）
1) 对“自由回忆/最近去了哪”等问题：在强证据（尤其 CCTV）被明确展示前，不承认与关键证据相关行踪，优先给出模糊或无害回答。
2) 对“是否参与犯罪/是否策划袭击”问题：始终否认。
3) 对“具体行为”问题：没有强证据就否认或说不确定；一旦对方拿出强证据，再按证据强度逐步承认“行踪/动作”，但不承认犯罪意图。
4) 对“行李箱用途”问题：给出无害解释（如旅行、收纳等），避免透露违法用途。
5) 只能承认“已被强证据覆盖”的事实，不得主动补充新关键细节。

# 最终输出要求
只生成一条简短对话回复（第一人称）。"""

    return prompt


@app.route("/api/avatar/token", methods=["POST"])
def api_avatar_token():
    """Create a LITE mode session token. LITE mode allows the app to use its own LLM.
    Returns session_token and session_id. Frontend should then use api_avatar_session to start."""
    if not LIVEAVATAR_API_KEY:
        return jsonify({"error": "LiveAvatar API Key 未配置"}), 500

    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    session_num = data.get("session_num")

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此培训类型不支持虚拟审讯"}), 400
    effective_type = training_type

    if session_num is not None:
        ts = store.get_training_session(phone, session_num)
        if not ts:
            return jsonify({"error": "请先通过训练列表开始本次训练"}), 400

    _, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)

    avatar_config = resolve_avatar_config(effective_type, suspect_profile)
    avatar_id = avatar_config.get("avatar_id", "")
    if not avatar_id:
        return jsonify({"error": "未配置 Avatar ID"}), 500

    try:
        # LITE mode: no context, no voice_id — the app handles LLM via DeepSeek
        sess_resp = requests.post(
            f"{LIVEAVATAR_API_URL}/sessions/token",
            headers={"X-API-KEY": LIVEAVATAR_API_KEY, "Content-Type": "application/json"},
            json={
                "mode": "LITE",
                "avatar_id": avatar_id,
                "is_sandbox": False,
            },
            timeout=15,
        )
        if not (200 <= sess_resp.status_code < 300):
            return jsonify({"error": f"创建 Session Token 失败: {sess_resp.text}"}), 500

        sess_data = sess_resp.json()
        # Handle both response shapes: {data: {...}} or direct response
        inner = sess_data.get("data", sess_data)
        return jsonify({
            "session_token": inner["session_token"],
            "session_id": inner["session_id"],
            "avatar_id": avatar_id,
        })

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"LiveAvatar API 请求失败: {str(e)}"}), 500


@app.route("/api/avatar/embed", methods=["POST"])
def api_avatar_embed():
    """Create a LiveAvatar embed URL (simple iframe-based integration)."""
    if not LIVEAVATAR_API_KEY:
        return jsonify({"error": "LiveAvatar API Key 未配置"}), 500

    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此培训类型不支持虚拟审讯"}), 400
    effective_type = training_type

    _, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)

    avatar_config = resolve_avatar_config(effective_type, suspect_profile)

    try:
        embed_resp = requests.post(
            f"{LIVEAVATAR_API_URL}/v2/embeddings",
            headers={"X-API-KEY": LIVEAVATAR_API_KEY, "Content-Type": "application/json"},
            json={
                "avatar_id": avatar_config.get("avatar_id", ""),
                "is_sandbox": False,
            },
            timeout=15,
        )
        if not (200 <= embed_resp.status_code < 300):
            return jsonify({"error": f"创建 Embed 失败: {embed_resp.text}"}), 500

        embed_data = embed_resp.json()["data"]
        return jsonify({
            "embed_url": embed_data.get("url", ""),
            "embed_script": embed_data.get("script", ""),
        })

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"LiveAvatar API 请求失败: {str(e)}"}), 500


@app.route("/api/avatar/session", methods=["POST"])
def api_avatar_session():
    """Create a LITE mode session, start it, and return LiveKit + WebSocket connection info.
    LITE mode uses the app's own DeepSeek LLM for suspect responses."""
    if not LIVEAVATAR_API_KEY:
        return jsonify({"error": "LiveAvatar API Key 未配置"}), 500

    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    session_num = data.get("session_num")

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此培训类型不支持虚拟审讯"}), 400
    effective_type = training_type

    if session_num is not None:
        ts = store.get_training_session(phone, session_num)
        if not ts:
            return jsonify({"error": "请先通过训练列表开始本次训练"}), 400

    _, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)

    avatar_config = resolve_avatar_config(effective_type, suspect_profile)
    avatar_id = avatar_config.get("avatar_id", "")
    if not avatar_id:
        return jsonify({"error": "未配置 Avatar ID"}), 500

    try:
        # Step 1: Create LITE mode session token (no context, no voice_id)
        sess_resp = requests.post(
            f"{LIVEAVATAR_API_URL}/sessions/token",
            headers={"X-API-KEY": LIVEAVATAR_API_KEY, "Content-Type": "application/json"},
            json={
                "mode": "LITE",
                "avatar_id": avatar_id,
                "is_sandbox": False,
            },
            timeout=15,
        )
        if not (200 <= sess_resp.status_code < 300):
            return jsonify({"error": f"创建 Session Token 失败: {sess_resp.text}"}), 500

        sess_data = sess_resp.json()
        inner = sess_data.get("data", sess_data)
        session_token = inner["session_token"]
        session_id = inner["session_id"]

        # Step 2: Start the session to get LiveKit + WebSocket connection info
        start_resp = requests.post(
            f"{LIVEAVATAR_API_URL}/sessions/start",
            headers={"Authorization": f"Bearer {session_token}", "Content-Type": "application/json"},
            timeout=15,
        )
        if not (200 <= start_resp.status_code < 300):
            return jsonify({"error": f"启动 Session 失败 (HTTP {start_resp.status_code}): {start_resp.text}"}), 500

        start_json = start_resp.json()
        # Per SDK, code 1000 = SUCCESS_CODE; also accept missing code field
        resp_code = start_json.get("code")
        if resp_code is not None and resp_code != 1000:
            return jsonify({"error": f"启动 Session 失败: {start_json.get('message', start_resp.text)}"}), 500

        start_data = start_json.get("data", start_json)

        return jsonify({
            "session_id": session_id,
            "session_token": session_token,
            "livekit_url": start_data.get("livekit_url", ""),
            "livekit_token": start_data.get("livekit_client_token", ""),
            "ws_url": start_data.get("ws_url", ""),
            "avatar_id": avatar_id,
        })

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"LiveAvatar API 请求失败: {str(e)}"}), 500


@app.route("/api/tts", methods=["POST"])
def api_tts():
    """Convert text to speech using ElevenLabs, return PCM 24kHz base64 audio for LiveAvatar LITE mode."""
    if not ELEVENLABS_API_KEY:
        return jsonify({"error": "ElevenLabs API Key 未配置"}), 500

    data = request.get_json()
    text = (data.get("text") or "").strip()
    voice_id = data.get("elevenlabs_voice_id") or data.get("voice_id") or "pNInz6obpgDQGcFmaJgB"

    if not text:
        return jsonify({"error": "text 不能为空"}), 400

    tts_text = sanitize_text_for_tts(text)
    if not tts_text:
        return jsonify({"error": "清洗后文本为空，无法生成语音"}), 400

    try:
        tts_resp = requests.post(
            f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/with-timestamps"
            "?output_format=pcm_24000",
            headers={
                "Content-Type": "application/json",
                "xi-api-key": ELEVENLABS_API_KEY,
            },
            json={"text": tts_text},
            timeout=20,
        )
        if tts_resp.status_code != 200:
            return jsonify({"error": f"TTS 生成失败: {tts_resp.text}"}), tts_resp.status_code

        tts_data = tts_resp.json()
        audio_base64 = tts_data.get("audio_base64", "")

        return jsonify({"audio": audio_base64})

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"TTS 请求失败: {str(e)}"}), 500


@app.route("/api/avatar/keep-alive", methods=["POST"])
def api_avatar_keep_alive():
    data = request.get_json()
    session_token = (data.get("session_token") or "").strip()
    if not session_token:
        return jsonify({"error": "缺少 session_token"}), 400

    try:
        resp = requests.post(
            f"{LIVEAVATAR_API_URL}/sessions/keep-alive",
            headers={"Authorization": f"Bearer {session_token}"},
            timeout=10,
        )
        if resp.status_code == 200:
            return jsonify({"status": "ok"})
        return jsonify({"error": resp.text}), resp.status_code
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Keep-alive 失败: {str(e)}"}), 500


@app.route("/api/avatar/stop", methods=["POST"])
def api_avatar_stop():
    data = request.get_json()
    session_token = (data.get("session_token") or "").strip()
    if not session_token:
        return jsonify({"error": "缺少 session_token"}), 400

    try:
        resp = requests.post(
            f"{LIVEAVATAR_API_URL}/sessions/stop",
            headers={"Authorization": f"Bearer {session_token}"},
            timeout=10,
        )
        if resp.status_code == 200:
            return jsonify({"status": "ok"})
        return jsonify({"error": resp.text}), resp.status_code
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"停止 Session 失败: {str(e)}"}), 500


# ====== Participant Lookup (for booking page) ======

@app.route("/api/lookup", methods=["POST"])
def api_lookup_participant():
    """Look up a participant by phone, returning only what’s needed for booking."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    blocked = _block_if_unbooked_timeout(p)
    if blocked:
        return blocked

    display_id = _participant_display_id(p)
    booking = store.get_my_booking(phone)
    if booking:
        time_slot = booking["time_slot"]
        is_matched = _slot_has_both_roles(time_slot)
        if is_matched:
            _sync_slot_participants_flow_step(time_slot, True)
            p = store.get_participant(phone) or p
        booking["is_matched"] = is_matched
        booking["participant_id"] = display_id

    training_info = None
    training_type = p.get("training_type", "")
    if training_type in ("avatar_specific", "avatar_general"):
        completed_count = store.count_completed_training_sessions(phone)
        training_info = {
            "training_type": training_type,
            "completed_count": completed_count,
            "total_required": 6,
            "all_completed": completed_count >= 6,
        }

    is_completed = bool(p.get("completed", 0))
    resume_step, resume_label = compute_participant_resume(p, booking)

    suspect_context = None
    suspect_case_label = None
    suspect_attention_qs = None
    suspect_display_id = None
    if p.get("role") == "S":
        case_type = p.get("case_type", "arson")
        guilt = p.get("guilt", "Guilty")
        group_name = p.get("group_name", "")
        suspect_display_id = (
            make_full_id(group_name, "S", participant_id_suffix(p)) if group_name else ""
        )
        suspect_case_label = "纵火案 Arson" if case_type == "arson" else "盗窃案 Theft"
        if case_type == "arson":
            suspect_context = ARSON_GUILTY_CONTEXT if guilt == "Guilty" else ARSON_INNOCENT_CONTEXT
        else:
            suspect_context = THEFT_GUILTY_CONTEXT if guilt == "Guilty" else THEFT_INNOCENT_CONTEXT
        check_key = f"{case_type}_{guilt.lower()}"
        suspect_attention_qs = ATTENTION_CHECKS.get(check_key, [])

    return jsonify({
        "found": True,
        "participant": {
            "phone": p["phone"],
            "role": p["role"],
            "full_id": display_id,
            "group_name": p.get("group_name", ""),
            "completed": p.get("completed", 0),
            "game_completed": p.get("game_completed", 0),
            "profile_completed": p.get("profile_completed", 0),
            "attention_passed": p.get("attention_passed", 0),
            "training_type": training_type,
            "training_info": training_info,
            "flow_step": (p.get("flow_step") or "").strip(),
            "sue_attention_passed": int(p.get("sue_attention_passed") or 0),
            "control_attention_passed": int(p.get("control_attention_passed") or 0),
            "case_evidence_recap_passed": int(p.get("case_evidence_recap_passed") or 0),
            "resume_step": resume_step,
            "resume_label": resume_label,
            # Suspect fields for case background / showSuspectFlow
            "display_id": suspect_display_id or display_id,
            "case_type": p.get("case_type"),
            "case_label": suspect_case_label,
            "context": suspect_context,
            "attention_questions": suspect_attention_qs,
            "consent_attention_required": consent_attention_required(p),
        },
        "appointment": booking,
    })


# ====== Interview Questionnaire ======

def _parse_slot_start(slot_str):
    try:
        return datetime.strptime((slot_str or "").strip(), "%Y-%m-%d %H:%M")
    except Exception:
        return None


def _is_override_open(phone, phase):
    row = store.get_questionnaire_override(phone, phase)
    if not row:
        return False
    val = row.get("is_open")
    if isinstance(val, bool):
        return val
    if isinstance(val, (int, float)):
        return int(val) == 1
    return str(val).strip().lower() in ("1", "true", "yes", "y")


def _questionnaire_access_state(phone, phase):
    p = store.get_participant(phone)
    if not p:
        return {
            "ok": False,
            "error": "未找到参与者",
            "status": "not_found",
        }

    if phase not in (PHASE_PRE, PHASE_POST):
        return {
            "ok": False,
            "error": "无效的问卷阶段",
            "status": "invalid_phase",
        }

    booking = store.get_my_booking(phone)
    if not booking:
        return {
            "ok": False,
            "error": "您尚未预约访谈时间，暂无法填写问卷",
            "status": "no_booking",
            "role": p.get("role"),
        }

    slot_str = booking.get("time_slot", "")
    slot_start = _parse_slot_start(slot_str)
    if not slot_start:
        return {
            "ok": False,
            "error": "预约时间格式异常，请联系管理员",
            "status": "bad_slot",
            "role": p.get("role"),
            "appointment_slot": slot_str,
        }

    now = datetime.now()
    override_open = _is_override_open(phone, phase)

    if phase == PHASE_PRE:
        open_time = slot_start - timedelta(minutes=QUESTIONNAIRE_PRE_OPEN_MINUTES_BEFORE)
        close_time = slot_start
    else:
        # Post questionnaire: opens 5 minutes after slot start; closes 1 hour after open.
        open_time = slot_start + timedelta(minutes=QUESTIONNAIRE_POST_OPEN_MINUTES_AFTER_SLOT_START)
        close_time = open_time + timedelta(minutes=QUESTIONNAIRE_POST_CLOSE_MINUTES_AFTER_OPEN)

    if override_open:
        is_open = True
        error = ""
    elif now < open_time:
        is_open = False
        if phase == PHASE_PRE:
            error = (
                f"访谈前问卷将于访谈开始前 {QUESTIONNAIRE_PRE_OPEN_MINUTES_BEFORE} 分钟开放"
                f"（{open_time.strftime('%Y-%m-%d %H:%M')} 起，至 {close_time.strftime('%H:%M')} 前完成）"
            )
        else:
            error = (
                f"访谈后问卷将于预约时间后 {QUESTIONNAIRE_POST_OPEN_MINUTES_AFTER_SLOT_START} 分钟开放"
                f"（{open_time.strftime('%Y-%m-%d %H:%M')} 起，截止 {close_time.strftime('%Y-%m-%d %H:%M')}）"
            )
    elif now > close_time:
        is_open = False
        if phase == PHASE_PRE:
            error = (
                f"访谈前问卷填写时间已结束（须在访谈开始前 {QUESTIONNAIRE_PRE_OPEN_MINUTES_BEFORE} 分钟内完成），"
                "请联系研究人员"
            )
        else:
            error = (
                f"访谈后问卷填写时间已结束（开放后 {QUESTIONNAIRE_POST_CLOSE_MINUTES_AFTER_OPEN} 分钟内截止），"
                "请联系研究人员"
            )
    else:
        is_open = True
        error = ""

    submitted_rows = store.get_interview_questionnaires(phone=phone, phase=phase)
    submitted = len(submitted_rows) > 0

    return {
        "ok": is_open,
        "status": "open" if is_open else "locked",
        "error": error,
        "role": p.get("role"),
        "appointment_slot": slot_str,
        "open_time": open_time.strftime("%Y-%m-%d %H:%M"),
        "close_time": close_time.strftime("%Y-%m-%d %H:%M"),
        "now": now.strftime("%Y-%m-%d %H:%M"),
        "manual_override": override_open,
        "submitted": submitted,
    }


@app.route("/api/questionnaire/status", methods=["POST"])
def api_questionnaire_status():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    phase = (data.get("phase") or "").strip().lower()
    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    state = _questionnaire_access_state(phone, phase)
    code = 200 if state.get("status") in ("open", "locked") else 404
    return jsonify(state), code


@app.route("/api/questionnaire/questions", methods=["POST"])
def api_questionnaire_questions():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    phase = (data.get("phase") or "").strip().lower()
    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    if phase not in (PHASE_PRE, PHASE_POST):
        return jsonify({"error": "无效的问卷阶段"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404
    role = p.get("role", "")
    qlist = INTERVIEW_QUESTION_BANK.get(phase, {}).get(role, [])
    return jsonify({
        "phase": phase,
        "role": role,
        "questions": qlist,
    })


@app.route("/api/questionnaire/submit", methods=["POST"])
def api_questionnaire_submit():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    phase = (data.get("phase") or "").strip().lower()
    answers = data.get("answers") or {}

    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    if phase not in (PHASE_PRE, PHASE_POST):
        return jsonify({"error": "无效的问卷阶段"}), 400
    if not isinstance(answers, dict) or not answers:
        return jsonify({"error": "问卷内容不能为空"}), 400

    state = _questionnaire_access_state(phone, phase)
    if not state.get("ok"):
        return jsonify({"error": state.get("error") or "问卷尚未开放", "status": state.get("status")}), 403

    p = store.get_participant(phone)
    booking = store.get_my_booking(phone)
    if not p or not booking:
        return jsonify({"error": "参与者或预约信息不存在"}), 404

    role = p.get("role", "")
    valid_ids = {q["id"] for q in INTERVIEW_QUESTION_BANK.get(phase, {}).get(role, [])}
    sanitized = {}
    for k, v in answers.items():
        if k not in valid_ids:
            continue
        if isinstance(v, str):
            sanitized[k] = sanitize(v, max_length=2000)
        elif isinstance(v, list):
            sanitized[k] = [sanitize(x, max_length=500) for x in v if isinstance(x, str)]
        else:
            sanitized[k] = v
    if not sanitized:
        return jsonify({"error": "问卷答案无效"}), 400

    answers_json = json.dumps(sanitized, ensure_ascii=False)
    qid = store.upsert_interview_questionnaire(
        phone=phone,
        role=role,
        phase=phase,
        appointment_slot=booking.get("time_slot", ""),
        answers_json=answers_json,
    )
    return jsonify({"success": True, "questionnaire_id": qid, "phase": phase})


# ====== Case Info (for interviewers) ======

def _interviewer_case_material_access(phone):
    """Whether interviewer may view case background (after booking + slot paired with suspect)."""
    p = store.get_participant(phone)
    if not p or p.get("role") != "I":
        return {"status": "not_interviewer", "message": "未找到审讯者"}

    booking = store.get_my_booking(phone)
    if not booking:
        return {
            "status": "no_booking",
            "message": "请先完成正式访谈时间预约，再查看案件背景与证据材料。",
        }

    time_slot = booking.get("time_slot", "")
    case_type, guilt, slot_matched = _interviewer_appointment_paired_suspect(phone)
    if not slot_matched:
        return {
            "status": "waiting_match",
            "message": (
                "您已预约该时间段，但尚未与嫌疑人配对成功，暂时无法查看案件材料。"
                "请过一段时间后再登录本系统确认；配对成功后将显示与您配对嫌疑人对应的案件背景与证据。"
            ),
            "time_slot": time_slot,
        }

    training_type = p.get("training_type", "")
    if training_type == "avatar_general":
        info = GENERAL_TERRORISM_CASE_INFO
        return {
            "status": "ready",
            "case_type": "terrorism",
            "case_title": info["title"],
            "overview": info["overview"],
            "evidence": info["evidence"],
            "efm_analysis": info.get("efm_analysis", ""),
            "suspect_guilt": guilt or "",
            "time_slot": time_slot,
        }

    info = CASE_INFO.get(case_type, CASE_INFO["arson"])
    return {
        "status": "ready",
        "case_type": case_type,
        "case_title": info["title"],
        "overview": info["overview"],
        "evidence": info["evidence"],
        "efm_analysis": info.get("efm_analysis", ""),
        "suspect_guilt": guilt or "",
        "time_slot": time_slot,
    }


@app.route("/api/case-info", methods=["POST"])
def api_case_info():
    """Return case info for interviewer after booking, when slot is paired with a suspect."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    if p.get("role") != "I":
        return jsonify({"error": "仅审讯者可查看案件信息"}), 403

    access = _interviewer_case_material_access(phone)
    status = access.get("status")
    if status != "ready":
        return jsonify({
            "status": status,
            "message": access.get("message", ""),
            "time_slot": access.get("time_slot"),
            "open_time": access.get("open_time"),
            "case_type": None,
            "case_title": None,
            "overview": None,
            "evidence": [],
            "efm_analysis": "",
            "suspect_guilt": "",
        })

    return jsonify({
        "status": "ready",
        "case_type": access["case_type"],
        "case_title": access["case_title"],
        "overview": access["overview"],
        "evidence": access["evidence"],
        "efm_analysis": access.get("efm_analysis", ""),
        "suspect_guilt": access.get("suspect_guilt", ""),
        "time_slot": access.get("time_slot"),
    })


@app.route("/api/case-evidence-recap/complete", methods=["POST"])
def api_case_evidence_recap_complete():
    """Mark interviewer evidence oral recap as passed (after recording + simulated review)."""
    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404
    if p.get("role") != "I":
        return jsonify({"error": "仅审讯者需要完成证据复述"}), 403
    store.update_participant(phone, case_evidence_recap_passed=1)
    return jsonify({"success": True})


# ====== Avatar Training Sessions (6-session requirement) ======

def _parse_training_order_csv(raw):
    """Parse comma-separated permutation of 1..6."""
    if not raw:
        return None
    try:
        parts = [int(x.strip()) for x in str(raw).split(",") if str(x).strip()]
    except (TypeError, ValueError):
        return None
    if len(parts) != 6 or sorted(parts) != list(range(1, 7)):
        return None
    return parts


def _reconstruct_training_avatar_order_from_sessions(sessions):
    """Infer avatar assignment order from existing training_sessions rows."""
    order = [None] * 6
    for s in sessions:
        try:
            sn = int(s.get("session_num"))
        except (TypeError, ValueError):
            continue
        if sn < 1 or sn > 6:
            continue
        setting = s.get("avatar_setting", "")
        guilt = s.get("avatar_guilt", "")
        matched = None
        for i, cfg in enumerate(AVATAR_TRAINING_SETTINGS):
            if cfg["setting"] == setting and cfg["guilt"] == guilt:
                matched = i + 1
                break
        if matched is not None:
            order[sn - 1] = matched
    if all(x is not None for x in order):
        return order
    return None


def _ensure_training_avatar_order(phone):
    """Per-participant permutation: session N uses AVATAR_TRAINING_SETTINGS[order[N-1]-1]."""
    p = store.get_participant(phone)
    if not p:
        return list(range(1, 7))
    existing = _parse_training_order_csv(p.get("training_avatar_order"))
    if existing:
        return existing
    sessions = store.get_training_sessions(phone)
    inferred = _reconstruct_training_avatar_order_from_sessions(sessions)
    if inferred:
        store.update_participant(phone, training_avatar_order=",".join(str(x) for x in inferred))
        return inferred
    order = list(range(1, 7))
    random.shuffle(order)
    store.update_participant(phone, training_avatar_order=",".join(str(x) for x in order))
    return order


def _ensure_training_ui_order(phone):
    """Per-participant card layout order on the training list screen."""
    p = store.get_participant(phone)
    if not p:
        return list(range(1, 7))
    existing = _parse_training_order_csv(p.get("training_ui_order"))
    if existing:
        return existing
    order = list(range(1, 7))
    random.shuffle(order)
    store.update_participant(phone, training_ui_order=",".join(str(x) for x in order))
    return order


def _training_setting_for_session(phone, session_num):
    """Avatar config for sequential session_num (1–6) under this participant's random order."""
    try:
        sn = int(session_num)
    except (TypeError, ValueError):
        sn = 1
    sn = max(1, min(6, sn))
    order = _ensure_training_avatar_order(phone)
    setting_index = order[sn - 1]
    return AVATAR_TRAINING_SETTINGS[setting_index - 1]


def _first_incomplete_training_session(existing_sessions):
    """Lowest session number 1..6 that is not fully completed (missing judgment or feedback)."""
    for i in range(1, 7):
        row = next((s for s in existing_sessions if str(s.get("session_num")) == str(i)), None)
        if row is None:
            return i
        if not (row.get("judgment") and row.get("feedback")):
            return i
    return None


@app.route("/api/avatar-training/status", methods=["POST"])
def api_avatar_training_status():
    """Get training session completion status for a participant."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此实验条件不需要 6 次虚拟审讯训练"}), 400

    sessions = store.get_training_sessions(phone)
    completed = [s for s in sessions if s.get("judgment") and s.get("feedback")]
    completed_count = len(completed)
    next_session_num = _first_incomplete_training_session(sessions)

    _ensure_training_avatar_order(phone)
    ui_order = _ensure_training_ui_order(phone)

    session_list = []
    for i in range(1, 7):
        existing = None
        for s in sessions:
            if str(s.get("session_num")) == str(i):
                existing = s
                break
        if existing:
            session_list.append({
                "session_num": i,
                "completed": bool(existing.get("judgment") and existing.get("feedback")),
                "judgment": existing.get("judgment", ""),
                "feedback": existing.get("feedback", ""),
                "avatar_guilt_label": (
                    "有罪" if existing.get("avatar_guilt") == "guilty"
                    else "无罪" if existing.get("avatar_guilt") == "innocent"
                    else ""
                ),
            })
        else:
            session_list.append({
                "session_num": i,
                "completed": False,
                "judgment": "",
                "feedback": "",
                "avatar_guilt_label": "",
            })

    session_list.sort(key=lambda row: ui_order.index(int(row["session_num"])))

    return jsonify({
        "training_type": p.get("training_type", ""),
        "completed_count": completed_count,
        "total_required": 6,
        "all_completed": completed_count >= 6,
        "next_session_num": next_session_num,
        "sessions": session_list,
    })


@app.route("/api/avatar-training/start", methods=["POST"])
def api_avatar_training_start():
    """Start an avatar training session. Client must pass session_num (1–6) matching the next incomplete slot."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    session_num_raw = data.get("session_num")

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此培训类型不需要虚拟审讯训练"}), 400
    effective_type = training_type

    existing_sessions = store.get_training_sessions(phone)
    completed = [s for s in existing_sessions if s.get("judgment") and s.get("feedback")]
    if len(completed) >= 6:
        return jsonify({"error": "您已完成全部 6 个 Avatar 培训", "all_completed": True}), 400

    expected = _first_incomplete_training_session(existing_sessions)
    if expected is None or expected > 6:
        return jsonify({"error": "您已完成全部 6 个 Avatar 培训", "all_completed": True}), 400

    try:
        requested = int(session_num_raw)
    except (TypeError, ValueError):
        return jsonify({"error": "请从训练列表选择第几次训练（session_num 无效）"}), 400

    if requested < 1 or requested > 6:
        return jsonify({"error": "训练序号必须在 1 到 6 之间"}), 400

    if requested != expected:
        return jsonify({"error": f"请按顺序进行训练，请先完成第 {expected} 次训练"}), 400

    next_num = requested

    setting_info = _training_setting_for_session(phone, next_num)
    avatar_setting = setting_info["setting"]
    avatar_guilt = setting_info["guilt"]

    suspect, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)

    # Get avatar visual config
    avatar_config = resolve_avatar_config(effective_type, suspect_profile)

    # Build training system prompt
    system_prompt = build_avatar_training_system_prompt(
        effective_type, avatar_setting, avatar_guilt, suspect, suspect_profile,
    )

    # Start the session in DB
    store.start_training_session(phone, p["id"], next_num, avatar_setting, avatar_guilt)

    appearance_key = _avatar_appearance_key(suspect_profile) if effective_type == "avatar_specific" else "generic"
    generic_id = (load_avatar_configs().get("generic") or {}).get("avatar_id", "")
    using_generic = effective_type == "avatar_specific" and avatar_config.get("avatar_id") == generic_id

    return jsonify({
        "session_num": next_num,
        "training_type": effective_type,
        "avatar_id": avatar_config.get("avatar_id", ""),
        "face_id": avatar_config.get("face_id", ""),
        "elevenlabs_voice_id": avatar_config.get("elevenlabs_voice_id", ""),
        "opening_text": avatar_config.get("opening_text", "你有什么要问的？"),
        "system_prompt": system_prompt,
        "avatar_appearance_key": appearance_key,
        "avatar_using_generic_fallback": using_generic,
    })


@app.route("/api/avatar-training/submit", methods=["POST"])
def api_avatar_training_submit():
    """Submit judgment and transcript for a training session. Generate feedback via DeepSeek."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    session_num = data.get("session_num")
    judgment = (data.get("judgment") or "").strip()
    transcript = (data.get("transcript") or "").strip()

    if not judgment:
        return jsonify({"error": "请做出有罪/无罪的判断"}), 400
    if not transcript:
        return jsonify({"error": "访谈记录不能为空"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    session = store.get_training_session(phone, session_num)
    if not session:
        return jsonify({"error": "未找到该训练会话"}), 404

    avatar_setting = session.get("avatar_setting", "")
    avatar_guilt = session.get("avatar_guilt", "")

    guilt_label = "有罪" if avatar_guilt == "guilty" else "无罪"
    judgment_label = "有罪" if judgment == "guilty" else "无罪"
    user_message = (
        f"Interview transcript:\n{transcript}\n\n"
        f"Avatar personality setting: {avatar_setting}\n"
        f"Suspect ground truth (for feedback only): {guilt_label}\n"
        f"Interviewer final judgment: {judgment_label}"
    )
    system_prompt = AVATAR_FEEDBACK_SYSTEM_PROMPT
    if "{transcript}" in system_prompt:
        system_prompt = system_prompt.format(
            transcript=transcript,
            avatar_setting_label=avatar_setting,
            avatar_guilt_label=guilt_label,
            interviewer_judgment=judgment_label,
        )
    feedback_suffix = (
        "\n\n【补充要求】\n"
        f"1. 在反馈开头明确写出：该 Avatar 的真实罪责状态为「{guilt_label}」。\n"
        "2. 仅根据转录内容分析嫌疑人哪些言行可能反映有罪、哪些可能反映无罪；"
        "不要分析眼神、表情、肢体语言等非转录内容。\n"
        "3. 有罪嫌疑人常见表现：遮掩信息、回避关键问题、前后矛盾等；"
        "无罪嫌疑人常见表现：相对坦诚、愿意澄清误会、回答一致等。\n"
        f"4. 将审讯者的判断「{judgment_label}」与真实状态「{guilt_label}」对比并给出改进建议。"
    )
    system_prompt = (system_prompt or "") + feedback_suffix

    feedback = ""
    try:
        feedback = deepseek_chat_completion(
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.7,
            max_tokens=4096,
            timeout=180,
        )
    except Exception as e:
        feedback = f"反馈生成失败: {str(e)}"

    store.submit_training_session(phone, session_num, judgment, transcript, feedback)

    completed_count = store.count_completed_training_sessions(phone)

    return jsonify({
        "success": True,
        "feedback": feedback,
        "avatar_guilt": avatar_guilt,
        "avatar_guilt_label": guilt_label,
        "session_num": session_num,
        "completed_count": completed_count,
        "all_completed": completed_count >= 6,
    })


@app.route("/api/avatar-training/save-transcript", methods=["POST"])
def api_avatar_training_save_transcript():
    """Persist current interview transcript during an active session (incremental save)."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    session_num = data.get("session_num")
    transcript = data.get("transcript")

    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    if session_num is None:
        return jsonify({"error": "缺少 session_num"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    if p.get("training_type") not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "无效的训练类型"}), 400

    session = store.get_training_session(phone, session_num)
    if not session:
        return jsonify({"error": "未找到该训练会话，请先开始训练"}), 404

    if isinstance(transcript, list):
        transcript_str = json.dumps(transcript, ensure_ascii=False)
    else:
        transcript_str = (transcript or "").strip()
        if not transcript_str:
            return jsonify({"error": "transcript 不能为空"}), 400

    if len(transcript_str) > 60000:
        transcript_str = transcript_str[:60000] + "\n...[truncated]"

    ok = store.update_training_transcript(phone, session_num, transcript_str)
    if not ok:
        return jsonify({"error": "保存访谈记录失败"}), 500
    return jsonify({"success": True})


@app.route("/api/avatar-practice/save-transcript", methods=["POST"])
def api_avatar_practice_save_transcript():
    """Deprecated: 理论组/控制组不再进行虚拟嫌疑人练习。"""
    return jsonify({"error": "该实验条件仅需阅读文字材料，无需保存虚拟练习记录"}), 400


@app.route("/api/avatar-training/chat", methods=["POST"])
def api_avatar_training_chat():
    """Chat endpoint for avatar training sessions. Sends user message to DeepSeek and returns response."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    session_num = data.get("session_num")
    user_message = (data.get("message") or "").strip()
    history = data.get("history") or []

    if not user_message:
        return jsonify({"error": "消息不能为空"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    training_type = p.get("training_type", "")
    if training_type not in ("avatar_specific", "avatar_general"):
        return jsonify({"error": "此培训类型不支持虚拟审讯"}), 400
    effective_type = training_type

    session = store.get_training_session(phone, session_num)
    if not session:
        return jsonify({"error": "请先开始训练会话"}), 400

    avatar_setting = session.get("avatar_setting", "")
    avatar_guilt = session.get("avatar_guilt", "")

    suspect, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)

    system_prompt = build_avatar_training_system_prompt(
        effective_type, avatar_setting, avatar_guilt, suspect, suspect_profile,
    )

    messages = [{"role": "system", "content": system_prompt}]
    for msg in history:
        messages.append(msg)
    messages.append({"role": "user", "content": user_message})

    try:
        reply_text = deepseek_chat_completion(messages, temperature=0.7, max_tokens=100)
        return jsonify({"reply": reply_text.strip()})
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"DeepSeek 请求失败: {str(e)}"}), 500


@app.route("/api/avatar-training/tts", methods=["POST"])
def api_avatar_training_tts():
    """TTS endpoint for avatar training sessions."""
    if not ELEVENLABS_API_KEY:
        return jsonify({"error": "ElevenLabs API Key 未配置"}), 500

    data = request.get_json()
    text = (data.get("text") or "").strip()
    phone = (data.get("phone") or "").strip()

    p = store.get_participant(phone) if phone else None
    training_type = p.get("training_type", "") if p else ""

    if training_type == "avatar_specific":
        _, suspect_profile = _paired_suspect_and_profile_for_interviewer(phone)
        avatar_config = resolve_avatar_config("avatar_specific", suspect_profile)
    else:
        avatar_config = resolve_avatar_config("avatar_general", None)

    voice_id = avatar_config.get("elevenlabs_voice_id", "pNInz6obpgDQGcFmaJgB")

    if not text:
        return jsonify({"error": "text 不能为空"}), 400

    tts_text = sanitize_text_for_tts(text)
    if not tts_text:
        return jsonify({"error": "清洗后文本为空，无法生成语音"}), 400

    try:
        tts_resp = requests.post(
            f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/with-timestamps"
            "?output_format=pcm_24000",
            headers={
                "Content-Type": "application/json",
                "xi-api-key": ELEVENLABS_API_KEY,
            },
            json={"text": tts_text},
            timeout=20,
        )
        if tts_resp.status_code != 200:
            return jsonify({"error": f"TTS 生成失败: {tts_resp.text}"}), tts_resp.status_code

        tts_data = tts_resp.json()
        return jsonify({"audio": tts_data.get("audio_base64", "")})
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"TTS 请求失败: {str(e)}"}), 500


@app.route("/api/avatar-training/stt", methods=["POST"])
def api_avatar_training_stt():
    """Speech-to-text endpoint for avatar training voice interrogation."""
    if not ELEVENLABS_API_KEY:
        return jsonify({"error": "ElevenLabs API Key 未配置"}), 500

    audio_file = request.files.get("audio")
    if not audio_file:
        return jsonify({"error": "缺少音频文件"}), 400

    filename = audio_file.filename or "audio.webm"
    content_type = audio_file.content_type or "audio/webm"
    audio_bytes = audio_file.read()
    if not audio_bytes:
        return jsonify({"error": "音频为空"}), 400

    try:
        stt_resp = requests.post(
            "https://api.elevenlabs.io/v1/speech-to-text",
            headers={"xi-api-key": ELEVENLABS_API_KEY},
            data={
                "model_id": "scribe_v1",
                "language_code": "zh",
            },
            files={
                "file": (filename, audio_bytes, content_type),
            },
            timeout=45,
        )
        if stt_resp.status_code != 200:
            return jsonify({"error": f"语音转写失败: {stt_resp.text}"}), stt_resp.status_code

        stt_data = stt_resp.json()
        text = (stt_data.get("text") or stt_data.get("transcript") or "").strip()
        if not text:
            return jsonify({"error": "未识别到有效语音内容"}), 422
        return jsonify({"text": text})
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"语音转写请求失败: {str(e)}"}), 500


# ====== Abandon incomplete participant ======

@app.route("/api/abandon", methods=["POST"])
def api_abandon():
    """Remove participant data if they haven't completed the experiment.
    Called via sendBeacon on page close/refresh."""
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    if not phone:
        return jsonify({"error": "缺少手机号"}), 400

    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    # Only delete if not yet completed
    if p.get("completed", 0) == 1:
        return jsonify({"status": "skipped", "reason": "already completed"})

    store.delete_participant(p["id"])
    return jsonify({"status": "deleted"})


# ====== Admin / Management APIs ======

def parse_slots(slots_str):
    try:
        return json.loads(slots_str) if isinstance(slots_str, str) else (slots_str or [])
    except Exception:
        return []


ADMIN_PARTICIPANT_OMIT_FIELDS = frozenset({
    "avatar_practice_transcript",
})

ADMIN_QUESTIONNAIRE_ANSWER_PREVIEW = 400


def _blacklist_set_from_meta(meta_rows):
    for m in meta_rows or []:
        if m.get("key") == META_BLACKLIST_PHONES:
            raw = m.get("value", "[]")
            try:
                phones = json.loads(raw) if isinstance(raw, str) else (raw or [])
            except Exception:
                phones = []
            return set(phones)
    return set()


def _training_session_counts(sessions_rows):
    counts = {}
    for s in sessions_rows or []:
        if not (s.get("judgment") and s.get("feedback")):
            continue
        phone = (s.get("phone") or "").strip()
        if phone:
            counts[phone] = counts.get(phone, 0) + 1
    return counts


def _groups_in_use_from_snapshot(participants, appointments):
    booked = {
        (a.get("phone") or "").strip()
        for a in appointments
        if a.get("status") == "confirmed" and (a.get("phone") or "").strip()
    }
    used = set()
    for p in participants:
        phone = (p.get("phone") or "").strip()
        if phone not in booked:
            continue
        n = _parse_group_number(p.get("group_name"))
        if n > 0:
            used.add(n)
    return used


def _slim_participant_for_admin(p):
    return {k: v for k, v in p.items() if k not in ADMIN_PARTICIPANT_OMIT_FIELDS}


def _slim_questionnaire_for_admin(q):
    row = dict(q)
    ans = row.get("answers_json") or ""
    if isinstance(ans, str) and len(ans) > ADMIN_QUESTIONNAIRE_ANSWER_PREVIEW:
        row["answers_json"] = ans[:ADMIN_QUESTIONNAIRE_ANSWER_PREVIEW] + "…"
        row["answers_truncated"] = True
    return row


def _build_admin_list_context(snapshot):
    """Pre-index snapshot data for fast per-participant admin progress."""
    participants = snapshot["participants"]
    appointments = snapshot["appointments"]
    questionnaires = snapshot["questionnaires"]

    booking_by_phone = {}
    slot_roles = {}
    for appt in appointments:
        if appt.get("status") != "confirmed":
            continue
        ph = (appt.get("phone") or "").strip()
        slot = (appt.get("time_slot") or "").strip()
        if ph:
            booking_by_phone[ph] = appt
        if slot:
            slot_roles.setdefault(slot, set()).add(appt.get("role"))

    qs_phones = {"pre": set(), "post": set()}
    for q in questionnaires:
        ph = (q.get("phone") or "").strip()
        phase = (q.get("phase") or "").strip().lower()
        if ph and phase in qs_phones:
            qs_phones[phase].add(ph)

    return {
        "booking_by_phone": booking_by_phone,
        "slot_roles": slot_roles,
        "qs_phones": qs_phones,
        "blacklist": _blacklist_set_from_meta(snapshot.get("meta")),
        "training_counts": _training_session_counts(snapshot.get("training_sessions")),
    }


def _participant_admin_progress(p, ctx, include_steps=False):
    phone = (p.get("phone") or "").strip()
    booking = ctx["booking_by_phone"].get(phone)
    slot = (booking.get("time_slot") or "").strip() if booking else ""
    roles = ctx["slot_roles"].get(slot, set()) if slot else set()
    matched = "S" in roles and "I" in roles
    training_sessions = ctx["training_counts"].get(phone, 0) if p.get("role") == "I" else 0
    return build_participant_admin_progress(
        p,
        booking=booking,
        matched=matched,
        qs_pre=phone in ctx["qs_phones"]["pre"],
        qs_post=phone in ctx["qs_phones"]["post"],
        training_sessions=training_sessions,
        blacklisted=phone in ctx["blacklist"],
        include_steps=include_steps,
    )


@app.route("/api/admin/results")
@admin_required
def admin_results():
    """Return admin dashboard data (optimized single Excel read)."""
    reconcile = request.args.get("reconcile", "").strip().lower() in ("1", "true", "yes")
    if reconcile:
        _reconcile_group_allocations()

    snapshot = store.get_admin_snapshot()
    participants = snapshot["participants"]
    appointments = snapshot["appointments"]
    ctx = _build_admin_list_context(snapshot)

    enriched = []
    for p in participants:
        row = _slim_participant_for_admin(p)
        row["admin_progress"] = _participant_admin_progress(p, ctx, include_steps=False)
        enriched.append(row)

    used_groups = _groups_in_use_from_snapshot(participants, appointments)
    max_g = max(used_groups) if used_groups else 0
    next_group = None
    for n in range(1, MAX_GROUPS + 1):
        if n not in used_groups:
            next_group = f"{n:03d}"
            break

    questionnaires = [_slim_questionnaire_for_admin(q) for q in snapshot["questionnaires"]]

    return jsonify({
        "participants": enriched,
        "appointments": appointments,
        "interview_questionnaires": questionnaires,
        "questionnaire_overrides": snapshot["overrides"],
        "max_group_number": max_g,
        "next_group_number": next_group,
    })


@app.route("/api/admin/participant-progress", methods=["POST"])
@admin_required
def admin_participant_progress():
    """Load full step checklist for one or more participants (group detail modal)."""
    data = request.get_json() or {}
    raw_phones = data.get("phones")
    if not raw_phones:
        one = (data.get("phone") or "").strip()
        raw_phones = [one] if one else []
    phones = []
    for ph in raw_phones:
        ph = (ph or "").strip()
        if ph and ph not in phones:
            phones.append(ph)
    if not phones:
        return jsonify({"error": "请提供 phone 或 phones"}), 400

    snapshot = store.get_admin_snapshot()
    by_phone = {
        (p.get("phone") or "").strip(): p
        for p in snapshot["participants"]
    }
    ctx = _build_admin_list_context(snapshot)
    progress = {}
    for phone in phones:
        p = by_phone.get(phone)
        if not p:
            progress[phone] = None
            continue
        progress[phone] = _participant_admin_progress(p, ctx, include_steps=True)
    return jsonify({"progress": progress})


@app.route("/api/admin/participants/group", methods=["POST"])
@admin_required
def admin_set_participant_group():
    """Manually set three-digit experiment group; rebuilds full_id. Optionally syncs slot mates."""
    data = request.get_json() or {}
    phone = (data.get("phone") or "").strip()
    if not phone:
        return jsonify({"error": "请提供手机号 phone"}), 400
    try:
        result = admin_apply_participant_group(
            phone,
            data.get("group_name"),
            sync_slot=data.get("sync_slot", True),
        )
        return jsonify({"success": True, **result})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/admin/download-data")
@admin_required
def admin_download_data():
    """Download the full experiment Excel database."""
    ensure_excel_file()
    if not os.path.isfile(EXCEL_FILE):
        return jsonify({"error": "数据文件不存在"}), 404
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    return send_file(
        EXCEL_FILE,
        as_attachment=True,
        download_name=f"experiment_data_{stamp}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/api/admin/purge-unbooked", methods=["POST"])
@admin_required
def admin_purge_unbooked():
    """Manually purge participants registered 24h+ ago without a confirmed booking."""
    purged = _purge_stale_unbooked_participants()
    if purged:
        logger.info(
            "Admin purged %d unbooked participant(s): %s",
            len(purged),
            [x["phone"] for x in purged],
        )
    return jsonify({"success": True, "count": len(purged), "purged": purged})


@app.route("/api/admin/results/<int:pid>", methods=["DELETE"])
@admin_required
def admin_delete_result(pid):
    store.delete_participant(pid)
    return jsonify({"success": True})


@app.route("/api/admin/appointments/<int:aid>", methods=["DELETE"])
@admin_required
def admin_delete_appointment(aid):
    store.delete_appointment_by_id(aid)
    return jsonify({"success": True})


@app.route("/api/admin/appointments/<int:aid>/reschedule", methods=["POST"])
@admin_required
def admin_reschedule_appointment_route(aid):
    data = request.get_json() or {}
    time_slot = (data.get("time_slot") or "").strip()
    try:
        result = admin_reschedule_appointment(aid, time_slot)
        return jsonify(result)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400


@app.route("/api/admin/appointment-slots")
@admin_required
def admin_get_appointment_slots():
    """Return all candidate slots with enabled/booking status for admin UI."""
    role_arg = (request.args.get("role") or "").strip().upper()
    role = role_arg if role_arg in ("S", "I") else None
    disabled = store.get_disabled_slots()
    slot_bookings = store.get_slot_bookings()
    fully_booked = store.get_confirmed_slot_set()
    slots = []
    for s in _candidate_booking_slots(role):
        roles = list(slot_bookings.get(s, set()))
        slots.append({
            "slot": s,
            "enabled": s not in disabled,
            "roles_booked": roles,
            "fully_booked": s in fully_booked,
        })
    return jsonify({
        "slots": slots,
        "role": role,
        "min_hours": BOOKING_MIN_HOURS,
        "max_days": BOOKING_MAX_DAYS,
        "time_windows": _booking_slot_windows_for_api(),
    })


@app.route("/api/admin/appointment-slots/apply-defaults", methods=["POST"])
@admin_required
def admin_apply_default_appointment_slots():
    """Enable all default booking slots within the configured date range (one-click)."""
    candidates = _candidate_booking_slots()
    count = store.enable_all_candidate_slots(candidates)
    return jsonify({
        "success": True,
        "enabled_count": count,
        "max_days": BOOKING_MAX_DAYS,
        "time_windows": _booking_slot_windows_for_api(),
    })


@app.route("/api/admin/appointment-slots", methods=["POST"])
@admin_required
def admin_set_appointment_slot():
    data = request.get_json() or {}
    slot_str = (data.get("slot") or "").strip()
    enabled = data.get("enabled")
    if not slot_str:
        return jsonify({"error": "时间段不能为空"}), 400
    if slot_str not in _candidate_booking_slots():
        return jsonify({"error": "该时间段不在可配置范围内"}), 400
    if enabled is None:
        return jsonify({"error": "请指定 enabled 参数"}), 400
    store.set_slot_enabled(slot_str, bool(enabled))
    return jsonify({"success": True, "slot": slot_str, "enabled": bool(enabled)})


@app.route("/api/admin/questionnaire/override", methods=["POST"])
@admin_required
def admin_set_questionnaire_override():
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    phase = (data.get("phase") or "").strip().lower()
    is_open = bool(data.get("is_open"))

    if not phone:
        return jsonify({"error": "手机号不能为空"}), 400
    if phase not in (PHASE_PRE, PHASE_POST):
        return jsonify({"error": "phase 必须为 pre 或 post"}), 400
    if not store.get_participant(phone):
        return jsonify({"error": "未找到参与者"}), 404

    oid = store.set_questionnaire_override(phone=phone, phase=phase, is_open=is_open)
    return jsonify({"success": True, "override_id": oid, "phone": phone, "phase": phase, "is_open": is_open})


@app.route("/manage")
def manage_page():
    return render_template("manage.html")


@app.route("/api/health")
def health_check():
    """Health check endpoint for monitoring and load balancer."""
    try:
        if os.path.exists(EXCEL_FILE):
            return jsonify({
                "status": "healthy",
                "timestamp": datetime.now().isoformat(),
                "data_file": "ok",
            }), 200
        else:
            return jsonify({"status": "degraded", "data_file": "missing"}), 503
    except Exception as e:
        return jsonify({"status": "unhealthy", "error": str(e)}), 503


# ====== Serious Game Routes ======

@app.route("/api/serious-game/start", methods=["POST"])
def serious_game_start():
    from flask import session
    data = request.get_json()
    phone = (data.get("phone") or "").strip()
    p = store.get_participant(phone)
    if not p:
        return jsonify({"error": "未找到参与者"}), 404

    case = p.get("case_type", "arson")
    guilt = p.get("guilt", "Guilty")
    case_label = "Arson" if case == "arson" else "Theft"

    timeline = build_serious_game_timeline(case_label, guilt)
    serialized = []
    for s in timeline:
        serialized.append({
            "video": s.video,
            "question": s.question,
            "a_label": s.a_label,
            "b_label": s.b_label,
            "next_default": s.next_default,
            "next_if_a": s.next_if_a,
            "next_if_b": s.next_if_b,
            "has_choice": s.has_choice,
        })

    session["sg_participant_id"] = p["full_id"]
    session["sg_phone"] = phone
    session["sg_case"] = case_label
    session["sg_condition"] = guilt
    session["sg_idx"] = 0
    session["sg_timeline"] = serialized
    session["sg_completed"] = False

    store.update_participant(phone, game_completed=0)

    return jsonify({
        "success": True,
        "case": case_label,
        "condition": guilt,
        "total_steps": len(timeline),
    })


@app.route("/api/serious-game/status", methods=["GET"])
def serious_game_status():
    from flask import session
    if "sg_participant_id" not in session:
        return jsonify({"started": False})
    return jsonify({
        "started": True,
        "idx": session.get("sg_idx", 0),
        "completed": session.get("sg_completed", False),
        "total": len(session.get("sg_timeline", [])),
    })


@app.route("/api/serious-game/step")
def serious_game_step():
    from flask import session
    if "sg_participant_id" not in session:
        return jsonify({"error": "Game not started"}), 400

    idx = session.get("sg_idx", 0)
    timeline = session.get("sg_timeline", [])

    if idx < 0 or idx >= len(timeline):
        return jsonify({"done": True})

    step = timeline[idx]
    video_url = serious_game_video_url(step["video"])
    if not video_url:
        return jsonify({
            "error": f"视频文件未找到: {step['video']}（请放入 static/videos/serious-game/）",
        }), 500

    return jsonify({
        "done": False,
        "idx": idx,
        "total": len(timeline),
        "video_url": video_url,
        "has_choice": step["has_choice"],
        "question": step.get("question"),
        "a_label": step.get("a_label"),
        "b_label": step.get("b_label"),
    })


@app.route("/api/serious-game/choice", methods=["POST"])
def serious_game_choice():
    from flask import session
    if "sg_participant_id" not in session:
        return jsonify({"error": "Game not started"}), 400

    idx = session.get("sg_idx", 0)
    timeline = session.get("sg_timeline", [])

    if idx < 0 or idx >= len(timeline):
        return jsonify({"done": True})

    step = timeline[idx]
    if not step["has_choice"]:
        return jsonify({"error": "This step does not accept a choice"}), 400

    if request.is_json:
        c = (request.get_json().get("choice") or "").strip().upper()
    else:
        c = (request.form.get("choice") or "").strip().upper()
    if c not in ("A", "B"):
        return jsonify({"error": "Invalid choice"}), 400

    next_idx = step["next_if_a"] if c == "A" else step["next_if_b"]
    if next_idx is None:
        return jsonify({"error": "Choice routing not configured"}), 400

    log_serious_game_choice(
        phone=session["sg_phone"],
        pid=session["sg_participant_id"],
        case=session["sg_case"],
        condition=session["sg_condition"],
        step_index=idx,
        video=step["video"],
        choice=c,
    )

    session["sg_idx"] = int(next_idx)
    return jsonify({"success": True, "next_idx": int(next_idx)})


@app.route("/api/serious-game/next", methods=["POST"])
def serious_game_next():
    from flask import session
    if "sg_participant_id" not in session:
        return jsonify({"error": "Game not started"}), 400

    idx = session.get("sg_idx", 0)
    timeline = session.get("sg_timeline", [])

    if idx < 0 or idx >= len(timeline):
        return jsonify({"done": True})

    step = timeline[idx]
    if step["has_choice"]:
        return jsonify({"error": "This step requires a choice"}), 400

    nxt = step.get("next_default")
    session["sg_idx"] = int(nxt) if nxt is not None else (idx + 1)
    return jsonify({"success": True, "next_idx": session["sg_idx"]})


@app.route("/api/serious-game/complete", methods=["POST"])
def serious_game_complete():
    from flask import session
    if "sg_participant_id" not in session:
        return jsonify({"error": "Game not started"}), 400

    phone = session.get("sg_phone", "")
    store.update_participant(phone, game_completed=1)
    session["sg_completed"] = True

    return jsonify({
        "success": True,
        "full_id": session["sg_participant_id"],
        "message": f"编号 {session['sg_participant_id']} 模拟行动游戏已完成。请截图此页面。",
    })


def _parse_participant_created_at(created_at_str):
    if not created_at_str:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M"):
        try:
            return datetime.strptime(str(created_at_str).strip(), fmt)
        except ValueError:
            continue
    return None


def _is_unbooked_past_deadline(participant):
    phone = (participant.get("phone") or "").strip()
    if not phone or store.has_booking(phone):
        return False
    created = _parse_participant_created_at(participant.get("created_at"))
    if not created:
        return False
    return datetime.now() - created >= timedelta(hours=UNBOOKED_PURGE_HOURS)


def _block_if_unbooked_timeout(participant):
    """On login: if past booking deadline without appointment, delete, blacklist, deny access."""
    if not _is_unbooked_past_deadline(participant):
        return None
    phone = participant.get("phone")
    store.blacklist_phone(phone, reason="unbooked_timeout")
    store.delete_participant(participant["id"])
    return jsonify({
        "error": "您注册已超过 24 小时仍未预约访谈时间，账号已注销且无法继续登录。",
        "blacklisted": True,
        "unbooked_timeout": True,
    }), 403


def _purge_stale_unbooked_participants():
    """Delete and blacklist participants with no booking UNBOOKED_PURGE_HOURS after registration."""
    now = datetime.now()
    cutoff = timedelta(hours=UNBOOKED_PURGE_HOURS)
    purged = []
    for p in store.get_all_participants():
        phone = (p.get("phone") or "").strip()
        if not phone:
            continue
        if store.has_booking(phone):
            continue
        created = _parse_participant_created_at(p.get("created_at"))
        if not created or now - created < cutoff:
            continue
        store.blacklist_phone(phone, reason="unbooked_timeout")
        store.delete_participant(p["id"])
        purged.append({
            "id": p["id"],
            "phone": phone,
            "role": p.get("role", ""),
            "group_name": p.get("group_name", ""),
            "created_at": p.get("created_at", ""),
        })
    return purged


def initialize_app():
    """Run all startup initialization. Called at import time for Gunicorn workers."""
    global AVATAR_FEEDBACK_SYSTEM_PROMPT
    init_excel()
    migrate_old_data()
    migrate_legacy_exports()
    setup_materials_dir()
    AVATAR_FEEDBACK_SYSTEM_PROMPT = _load_avatar_feedback_system_prompt()
    setup_liveavatar_voices()
    logger.info("Application initialized successfully")


# Initialize immediately on module import (required for Gunicorn)
initialize_app()


if __name__ == "__main__":
    hostname = socket.gethostname()
    local_ip = socket.gethostbyname(hostname)
    print(f"\n{'='*50}")
    print(f"  审讯实验系统已启动")
    print(f"  PC上访问: http://localhost:5000")
    print(f"  管理页面: http://localhost:5000/manage")
    print(f"  局域网访问: http://{local_ip}:5000")
    print(f"  数据存储: {EXCEL_FILE}")
    print(f"{'='*50}\n")

    cert_file = os.path.join(BASE_DIR, "cert.pem")
    key_file = os.path.join(BASE_DIR, "key.pem")
    use_https = os.path.isfile(cert_file) and os.path.isfile(key_file)

    if use_https:
        print(f"  HTTPS 模式: https://{local_ip}:5000")
        app.run(debug=DEBUG_MODE, host="0.0.0.0", port=5000, ssl_context=(cert_file, key_file))
    else:
        app.run(debug=DEBUG_MODE, host="0.0.0.0", port=5000)
