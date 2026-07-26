import ast
import os
import re
from datetime import datetime, timezone
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = os.path.dirname(os.path.abspath(__file__))
APP = os.path.join(BASE, "app.py")
REF = r"E:\Working Documents\【Important】My Projects\24 Specific avatar\嫌疑人Avatar背景提示词.docx"
OUT = os.path.join(BASE, "嫌疑人Avatar提示词_程序与参考文档对比_修订版.docx")


def load_constants(path):
    text = open(path, "r", encoding="utf-8").read()
    tree = ast.parse(text)
    out = {}
    for node in tree.body:
        if isinstance(node, ast.Assign) and len(node.targets) == 1:
            name = node.targets[0].id if isinstance(node.targets[0], ast.Name) else None
            if name in {"ARSON_GUILTY_CONTEXT", "ARSON_INNOCENT_CONTEXT", "THEFT_GUILTY_CONTEXT", "THEFT_INNOCENT_CONTEXT"}:
                out[name] = ast.literal_eval(node.value)
    return out


def extract_reference():
    doc = Document(REF)
    paras = [p.text.strip() for p in doc.paragraphs]
    sections = {}
    starts = [("arson_guilty", "一、纵火案 — 有罪版"), ("arson_innocent", "二、纵火案 — 无罪版"),
              ("theft_guilty", "三、盗窃案 — 有罪版"), ("theft_innocent", "四、盗窃案 — 无罪版")]
    indices = []
    for key, heading in starts:
        idx = next((i for i, x in enumerate(paras) if x == heading), None)
        if idx is not None:
            indices.append((key, idx))
    for n, (key, idx) in enumerate(indices):
        end = indices[n + 1][1] if n + 1 < len(indices) else len(paras)
        sections[key] = [x for x in paras[idx:end] if x and x != "---"]
    return paras, sections


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_track_revisions(doc):
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.insert(0, OxmlElement("w:trackRevisions"))


def revision_run(text, kind):
    wrapper = OxmlElement("w:ins" if kind == "ins" else "w:del")
    wrapper.set(qn("w:id"), str(revision_run.counter))
    revision_run.counter += 1
    wrapper.set(qn("w:author"), "Codex")
    wrapper.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "C00000" if kind == "ins" else "808080")
    rPr.append(color)
    if kind == "del":
        strike = OxmlElement("w:strike")
        rPr.append(strike)
    r.append(rPr)
    t = OxmlElement("w:delText" if kind == "del" else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    wrapper.append(r)
    return wrapper


revision_run.counter = 1


def add_revision_paragraph(doc, text, kind, style=None):
    p = doc.add_paragraph(style=style)
    p._p.append(revision_run(text, kind))
    p.paragraph_format.space_after = Pt(3)
    return p


def add_source_paragraph(doc, text, color="404040", italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(color)
    r.italic = italic
    p.paragraph_format.space_after = Pt(3)
    return p


def add_label(doc, text, fill):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(fill)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    return p


def main():
    constants = load_constants(APP)
    _, ref_sections = extract_reference()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    set_track_revisions(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("嫌疑人 Avatar 提示词\n程序版本与参考 Word 对比（修订版）")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("生成日期：2026-07-17  |  修订作者：Codex")
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.add_run("结论：").bold = True
    p.add_run("当前运行代码（app.py）实际使用 4 套案件/罪责上下文：纵火案有罪、纵火案无罪、盗窃案有罪、盗窃案无罪。generic avatar 配置只提供开场白，不包含独立背景提示词。参考 Word 同样包含这 4 个版本，但每个版本都扩展为完整的逐字稿式提示词，新增了角色约束、应答策略、示例问答和规则总结。")

    doc.add_heading("一、程序当前实际使用的提示词", level=1)
    doc.add_paragraph("以下内容直接提取自 app.py 中的 ARSON_*_CONTEXT / THEFT_*_CONTEXT 常量；它们在训练和审讯会话中被按案件类型与罪责状态注入 Avatar。")
    labels = [("ARSON_GUILTY_CONTEXT", "纵火案 — 有罪版"), ("ARSON_INNOCENT_CONTEXT", "纵火案 — 无罪版"),
              ("THEFT_GUILTY_CONTEXT", "盗窃案 — 有罪版"), ("THEFT_INNOCENT_CONTEXT", "盗窃案 — 无罪版")]
    for name, title_text in labels:
        doc.add_heading(title_text, level=2)
        add_source_paragraph(doc, constants[name], "404040")

    doc.add_heading("二、参考 Word 相对程序版本的修订对比", level=1)
    doc.add_paragraph("下面每个小节均以程序当前上下文为基线：灰色删除线表示程序中的简短上下文被参考文档版本替换；红色内容表示参考 Word 新增或扩展的内容。文档已启用 Word Track Changes（修订）。")
    mapping = [("arson_guilty", "ARSON_GUILTY_CONTEXT", "纵火案 — 有罪版"),
               ("arson_innocent", "ARSON_INNOCENT_CONTEXT", "纵火案 — 无罪版"),
               ("theft_guilty", "THEFT_GUILTY_CONTEXT", "盗窃案 — 有罪版"),
               ("theft_innocent", "THEFT_INNOCENT_CONTEXT", "盗窃案 — 无罪版")]
    for key, const_name, title_text in mapping:
        doc.add_heading(title_text, level=2)
        add_label(doc, "程序原始上下文（删除/被替换）", "808080")
        add_revision_paragraph(doc, constants[const_name], "del")
        add_label(doc, "参考 Word 新增与扩展内容（插入）", "C00000")
        for line in ref_sections.get(key, []):
            if line.startswith(("一、", "二、", "三、", "四、")):
                continue
            add_revision_paragraph(doc, line, "ins")

    doc.add_heading("三、差异摘要", level=1)
    bullets = [
        "参考 Word 新增了统一的角色边界：明确 Avatar 不是助手、不得使用普通助手式问候语，并要求始终保持嫌疑人角色。",
        "纵火案与盗窃案各自新增了完整的内部案件记忆、时间线、证据解释和罪责状态下的行为目标。",
        "参考 Word 新增了按证据强弱、具体程度逐步承认或否认的应答策略，并配套大量审讯者/嫌疑人示例对话。",
        "参考 Word 新增了每个版本的规则总结；这些内容在 app.py 的四个 *_CONTEXT 常量中没有对应的完整文本。",
        "没有发现参考 Word 中存在而程序完全没有案件对应关系的第五个案件版本；文档标题提到的 Charlie/液体炸弹模板是格式来源，不是当前 app.py 的独立运行时背景提示词。",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b)

    doc.add_heading("四、来源", level=1)
    add_source_paragraph(doc, "程序来源：app.py（ARSON_GUILTY_CONTEXT、ARSON_INNOCENT_CONTEXT、THEFT_GUILTY_CONTEXT、THEFT_INNOCENT_CONTEXT）。", "404040")
    add_source_paragraph(doc, "参考来源：E:\\Working Documents\\【Important】My Projects\\24 Specific avatar\\嫌疑人Avatar背景提示词.docx。", "404040")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
